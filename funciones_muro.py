from dataclasses import dataclass
from io import BytesIO
from datetime import datetime
import math

import pandas as pd
from matplotlib.patches import Polygon, FancyArrowPatch
from docx import Document


FT_A_M = 0.3048
KIP_A_TON = 0.45359237
KSF_A_TON_M2 = 4.882427636


@dataclass
class DatosMuro:
    """
    Almacena los datos geométricos, geotécnicos y de materiales del muro.

    Unidades de entrada:
    - Longitudes: m
    - Capacidad admisible del suelo qa: ton/m²
    - Peso unitario del suelo y hormigón: ton/m³
    - Cohesión del suelo: ton/m²
    - Resistencia del hormigón f'c: kg/cm²
    - Fluencia del acero fy: kg/cm²
    - Ángulo de fricción interna: grados
    """
    H: float
    B: float
    hz: float
    puntera: float
    t_base: float
    t_corona: float
    altura_relleno: float
    pendiente_h: float
    pendiente_v: float
    usar_llave: bool
    ancho_llave: float
    profundidad_llave: float
    pos_llave: float
    qa: float
    gamma_suelo: float
    phi: float
    cohesion: float
    mu: float
    fc: float
    fy: float
    gamma_hormigon: float


def validar_datos_muro(datos: DatosMuro) -> list[str]:
    """
    Revisa que los datos ingresados sean geométrica y físicamente coherentes.
    """
    errores = []
    talon = datos.B - datos.puntera - datos.t_base

    if datos.H <= 0:
        errores.append("La altura del fuste debe ser mayor que cero.")
    if datos.B <= 0:
        errores.append("El ancho total de zapata debe ser mayor que cero.")
    if datos.hz <= 0:
        errores.append("El espesor de zapata debe ser mayor que cero.")
    if datos.t_base <= 0:
        errores.append("El espesor del fuste en la base debe ser mayor que cero.")
    if datos.t_corona <= 0:
        errores.append("El espesor del fuste en la corona debe ser mayor que cero.")
    if datos.t_corona > datos.t_base:
        errores.append("El espesor de corona no debería ser mayor que el espesor de base.")
    if datos.puntera < 0:
        errores.append("La puntera no puede ser negativa.")
    if talon <= 0:
        errores.append("La suma puntera + espesor del fuste en la base no puede superar el ancho total de zapata.")
    if datos.altura_relleno < 0:
        errores.append("La altura de relleno no puede ser negativa.")
    if datos.pendiente_h <= 0:
        errores.append("La componente horizontal de la pendiente debe ser mayor que cero.")

    if datos.usar_llave:
        borde_izq_llave = datos.pos_llave - datos.ancho_llave / 2
        borde_der_llave = datos.pos_llave + datos.ancho_llave / 2
        if datos.ancho_llave <= 0:
            errores.append("El ancho de la llave debe ser mayor que cero.")
        if datos.profundidad_llave <= 0:
            errores.append("La profundidad de la llave debe ser mayor que cero.")
        if borde_izq_llave < 0 or borde_der_llave > datos.B:
            errores.append("La llave debe quedar dentro del ancho total de la zapata.")

    if datos.qa <= 0:
        errores.append("La capacidad admisible del suelo qa debe ser mayor que cero.")
    if datos.gamma_suelo <= 0:
        errores.append("El peso unitario del suelo debe ser mayor que cero.")
    if datos.gamma_hormigon <= 0:
        errores.append("El peso unitario del hormigón debe ser mayor que cero.")
    if datos.fc <= 0:
        errores.append("La resistencia del hormigón f'c debe ser mayor que cero.")
    if datos.fy <= 0:
        errores.append("La fluencia del acero fy debe ser mayor que cero.")
    if datos.phi < 0 or datos.phi > 50:
        errores.append("El ángulo de fricción interna debería estar entre 0° y 50°.")
    if datos.cohesion < 0:
        errores.append("La cohesión no puede ser negativa.")
    if datos.mu <= 0:
        errores.append("El coeficiente de fricción debe ser mayor que cero.")

    return errores


def calcular_talon(datos: DatosMuro) -> float:
    """
    Calcula la longitud del talón de la zapata en metros.
    """
    return datos.B - datos.puntera - datos.t_base


def generar_puntos_muro(datos: DatosMuro) -> dict:
    """
    Genera las coordenadas principales del muro para graficarlo.
    """
    x_frente_fuste = datos.puntera
    x_dorso_fuste_base = datos.puntera + datos.t_base
    x_dorso_fuste_corona = datos.puntera + datos.t_corona

    zapata = [
        (0, 0),
        (datos.B, 0),
        (datos.B, -datos.hz),
        (0, -datos.hz),
    ]

    fuste = [
        (x_frente_fuste, 0),
        (x_dorso_fuste_base, 0),
        (x_dorso_fuste_corona, datos.H),
        (x_frente_fuste, datos.H),
    ]

    geometria = {
        "zapata": zapata,
        "fuste": fuste,
        "x_frente_fuste": x_frente_fuste,
        "x_dorso_fuste_base": x_dorso_fuste_base,
        "x_dorso_fuste_corona": x_dorso_fuste_corona,
    }

    if datos.usar_llave:
        x1 = datos.pos_llave - datos.ancho_llave / 2
        x2 = datos.pos_llave + datos.ancho_llave / 2
        geometria["llave"] = [
            (x1, -datos.hz),
            (x2, -datos.hz),
            (x2, -datos.hz - datos.profundidad_llave),
            (x1, -datos.hz - datos.profundidad_llave),
        ]

    return geometria


def calcular_linea_relleno(datos: DatosMuro, geometria: dict) -> list[tuple[float, float]]:
    """
    Calcula la línea superior del relleno detrás del muro.

    La línea ya no se fuerza a pasar por la corona. Parte desde la altura de
    relleno ingresada por el usuario.
    """
    x_inicio = geometria["x_dorso_fuste_corona"]
    y_inicio = datos.altura_relleno

    x_fin = datos.B + max(datos.B * 0.60, 1.00)
    y_fin = calcular_y_terreno(datos, geometria, x_fin)

    return [(x_inicio, y_inicio), (x_fin, y_fin)]



def altura_relleno_en_muro(datos: DatosMuro, geometria: dict) -> float:
    """
    Altura efectiva de relleno contra el fuste desde la base del muro.

    Se limita entre 0 y H para el diseño local del fuste. Si el relleno está
    por encima de la corona, la parte excedente debe tratarse como sobrecarga
    o geometría de terreno superior; para esta app se limita la altura retenida
    directa al alto del fuste.
    """
    return max(0.0, min(datos.altura_relleno, datos.H))


def suelo_sobre_talon(datos: DatosMuro, geometria: dict) -> dict:
    """
    Calcula el área, peso y centroide aproximado del suelo sobre el talón.

    Se usa un trapecio entre:
    - x = puntera + t_base;
    - x = B;
    con alturas dadas por la línea real de terreno.
    """
    talon = calcular_talon(datos)
    if talon <= 0:
        return {"area_m2": 0.0, "W_ton_m": 0.0, "x_centroide_m": datos.B}

    x1 = datos.puntera + datos.t_base
    x2 = datos.B
    h1 = max(calcular_y_terreno(datos, geometria, x1), 0.0)
    h2 = max(calcular_y_terreno(datos, geometria, x2), 0.0)

    area = talon * (h1 + h2) / 2.0
    W = area * datos.gamma_suelo

    # Centroide de una altura lineal h(x) sobre el intervalo [x1, x2].
    if h1 + h2 > 0:
        x_rel = talon * (h1 + 2.0 * h2) / (3.0 * (h1 + h2))
        x_c = x1 + x_rel
    else:
        x_c = x1 + talon / 2.0

    return {"area_m2": area, "W_ton_m": W, "x_centroide_m": x_c, "h_inicio_m": h1, "h_fin_m": h2}



def dibujar_poligono(ax, puntos: list[tuple[float, float]], etiqueta: str):
    """
    Dibuja un polígono cerrado en el eje de Matplotlib.
    """
    poligono = Polygon(
        puntos,
        closed=True,
        fill=False,
        linewidth=2.2,
        label=etiqueta
    )
    ax.add_patch(poligono)


def dibujar_cotas_principales(ax, datos: DatosMuro, tamano_texto: int = 8):
    """
    Dibuja cotas principales del muro.
    """
    talon = calcular_talon(datos)

    y_cota_puntera_talon = -datos.hz - 0.18
    y_texto_puntera_talon = -datos.hz - 0.26

    y_cota_total = -datos.hz - 1.35
    y_texto_total = -datos.hz - 1.48

    ax.annotate(
        "",
        xy=(-0.25, 0),
        xytext=(-0.25, datos.H),
        arrowprops=dict(arrowstyle="<->", linewidth=1.0)
    )
    ax.text(-0.33, datos.H / 2, f"H = {datos.H:.2f} m", rotation=90, va="center", ha="right", fontsize=tamano_texto)

    ax.annotate(
        "",
        xy=(0, y_cota_puntera_talon),
        xytext=(datos.puntera, y_cota_puntera_talon),
        arrowprops=dict(arrowstyle="<->", linewidth=0.9)
    )
    ax.text(datos.puntera / 2, y_texto_puntera_talon, f"Puntera = {datos.puntera:.2f} m", ha="center", va="top", fontsize=tamano_texto)

    x_ini_talon = datos.puntera + datos.t_base
    ax.annotate(
        "",
        xy=(x_ini_talon, y_cota_puntera_talon),
        xytext=(datos.B, y_cota_puntera_talon),
        arrowprops=dict(arrowstyle="<->", linewidth=0.9)
    )
    ax.text(x_ini_talon + talon / 2, y_texto_puntera_talon, f"Talón = {talon:.2f} m", ha="center", va="top", fontsize=tamano_texto)

    ax.annotate(
        "",
        xy=(0, y_cota_total),
        xytext=(datos.B, y_cota_total),
        arrowprops=dict(arrowstyle="<->", linewidth=1.0)
    )
    ax.text(datos.B / 2, y_texto_total, f"B = {datos.B:.2f} m", ha="center", va="top", fontsize=tamano_texto)

    ax.annotate(
        "",
        xy=(datos.B + 0.18, 0),
        xytext=(datos.B + 0.18, -datos.hz),
        arrowprops=dict(arrowstyle="<->", linewidth=0.9)
    )
    ax.text(datos.B + 0.25, -datos.hz / 2, f"hz = {datos.hz:.2f} m", rotation=90, va="center", ha="left", fontsize=tamano_texto)


def configurar_ejes_muro(ax, datos: DatosMuro, mostrar_ejes: bool = True):
    """
    Configura límites, escala y apariencia del gráfico del muro.
    """
    margen_x = max(datos.B * 0.25, 0.80)
    margen_y = max(datos.H * 0.20, 0.80)

    y_min = -datos.hz - (datos.profundidad_llave if datos.usar_llave else 0) - max(margen_y, 2.10)
    y_max = max(datos.H, datos.altura_relleno) + margen_y

    ax.set_xlim(-margen_x, datos.B + margen_x)
    ax.set_ylim(y_min, y_max)
    ax.set_aspect("equal", adjustable="box")

    if mostrar_ejes:
        ax.grid(True, linestyle=":", linewidth=0.7)
        ax.set_xlabel("x [m]")
        ax.set_ylabel("y [m]")
    else:
        ax.axis("off")


def dibujar_muro(
    ax,
    datos: DatosMuro,
    geometria: dict,
    tamano_texto: int = 8,
    mostrar_cotas: bool = True,
    mostrar_ejes: bool = True
):
    """
    Dibuja el muro completo sin leyenda.
    """
    dibujar_poligono(ax, geometria["zapata"], "Zapata")
    dibujar_poligono(ax, geometria["fuste"], "Fuste")

    if datos.usar_llave and "llave" in geometria:
        dibujar_poligono(ax, geometria["llave"], "Llave de corte")

    linea_relleno = calcular_linea_relleno(datos, geometria)
    xs = [p[0] for p in linea_relleno]
    ys = [p[1] for p in linea_relleno]
    ax.plot(xs, ys, linewidth=2)

    ax.plot([0, datos.B], [0, 0], linestyle="--", linewidth=1)

    if mostrar_cotas:
        dibujar_cotas_principales(ax, datos, tamano_texto=tamano_texto)

    ax.set_title("Geometría del muro de contención")
    configurar_ejes_muro(ax, datos, mostrar_ejes=mostrar_ejes)


def flecha(ax, inicio, fin, texto, dx=0.0, dy=0.0, tamano=9):
    """
    Dibuja una flecha con etiqueta.
    """
    arrow = FancyArrowPatch(
        inicio,
        fin,
        arrowstyle="->",
        mutation_scale=14,
        linewidth=1.8
    )
    ax.add_patch(arrow)
    ax.text(fin[0] + dx, fin[1] + dy, texto, fontsize=tamano, ha="center", va="center")


def dibujar_diagrama_fuerzas(ax, datos: DatosMuro, geometria: dict, mostrar_ejes: bool = True):
    """
    Dibuja un diagrama didáctico de las fuerzas principales del PDF:
    ΣW, khΣW, PA/PAE, PP/PPE, V y Rf.

    El objetivo es visual, no reemplaza el cálculo numérico.
    """
    dibujar_muro(
        ax,
        datos,
        geometria,
        tamano_texto=8,
        mostrar_cotas=False,
        mostrar_ejes=mostrar_ejes
    )

    x_centro = datos.B * 0.58
    y_centro = datos.H * 0.45

    # Peso total.
    flecha(ax, (x_centro, y_centro + 0.90), (x_centro, y_centro + 0.25), "ΣW", dx=0.18, dy=0.05)

    # Fuerza sísmica inercial.
    flecha(ax, (x_centro + 0.80, y_centro + 0.80), (x_centro + 0.20, y_centro + 0.80), "khΣW", dx=-0.20, dy=0.18)

    # Empuje activo.
    x_dorso = geometria["x_dorso_fuste_base"]
    flecha(
        ax,
        (datos.B + 0.65, datos.H * 0.38),
        (x_dorso + 0.08, datos.H * 0.32),
        "PA / PAE",
        dx=0.20,
        dy=0.25
    )

    # Empuje pasivo.
    flecha(
        ax,
        (-0.65, -datos.hz * 0.45),
        (0.05, -datos.hz * 0.45),
        "PP / PPE",
        dx=-0.10,
        dy=-0.25
    )

    # Reacción vertical.
    flecha(
        ax,
        (datos.B * 0.42, -datos.hz - 0.90),
        (datos.B * 0.42, -datos.hz - 0.15),
        "V",
        dx=0.15,
        dy=0.00
    )

    # Fricción en base.
    flecha(
        ax,
        (datos.B * 0.45, -datos.hz - 0.10),
        (datos.B * 0.70, -datos.hz - 0.10),
        "Rf",
        dx=0.20,
        dy=0.12
    )

    ax.set_title("Diagrama didáctico de fuerzas")


def convertir_resistencias_a_sistema_interno(datos: DatosMuro) -> pd.DataFrame:
    """
    Convierte resistencias y pesos a unidades alternativas útiles para cálculo.
    """
    fc_ton_m2 = datos.fc * 10.0
    fy_ton_m2 = datos.fy * 10.0
    qa_kg_cm2 = datos.qa / 10.0

    filas = [
        ("f'c", datos.fc, "kg/cm²", fc_ton_m2, "ton/m²"),
        ("fy", datos.fy, "kg/cm²", fy_ton_m2, "ton/m²"),
        ("qa", datos.qa, "ton/m²", qa_kg_cm2, "kg/cm²"),
        ("γ suelo", datos.gamma_suelo, "ton/m³", datos.gamma_suelo, "ton/m³"),
        ("γ hormigón", datos.gamma_hormigon, "ton/m³", datos.gamma_hormigon, "ton/m³"),
        ("c", datos.cohesion, "ton/m²", datos.cohesion, "ton/m²"),
    ]

    return pd.DataFrame(
        filas,
        columns=["Parámetro", "Valor ingresado", "Unidad ingresada", "Valor convertido", "Unidad convertida"]
    )


def resumen_geometria(datos: DatosMuro) -> pd.DataFrame:
    """
    Construye una tabla resumen con geometría, suelo y materiales.
    """
    talon = calcular_talon(datos)

    filas = [
        ("Altura del fuste H", datos.H, "m"),
        ("Ancho total de zapata B", datos.B, "m"),
        ("Espesor de zapata hz", datos.hz, "m"),
        ("Puntera", datos.puntera, "m"),
        ("Talón calculado", talon, "m"),
        ("Espesor fuste base", datos.t_base, "m"),
        ("Espesor fuste corona", datos.t_corona, "m"),
        ("Altura de relleno", datos.altura_relleno, "m"),
        ("Pendiente relleno V:H", f"{datos.pendiente_v:.2f}:{datos.pendiente_h:.2f}", "-"),
        ("Capacidad admisible qa", datos.qa, "ton/m²"),
        ("Peso unitario del suelo γs", datos.gamma_suelo, "ton/m³"),
        ("Ángulo de fricción φ", datos.phi, "grados"),
        ("Cohesión c", datos.cohesion, "ton/m²"),
        ("Coeficiente de fricción μ", datos.mu, "-"),
        ("Resistencia hormigón f'c", datos.fc, "kg/cm²"),
        ("Fluencia acero fy", datos.fy, "kg/cm²"),
        ("Peso unitario hormigón γc", datos.gamma_hormigon, "ton/m³"),
    ]

    if datos.usar_llave:
        filas.extend([
            ("Ancho de llave", datos.ancho_llave, "m"),
            ("Profundidad de llave", datos.profundidad_llave, "m"),
            ("Posición eje de llave", datos.pos_llave, "m"),
        ])
    else:
        filas.append(("Llave de corte", "No incluida", "-"))

    return pd.DataFrame(filas, columns=["Parámetro", "Valor", "Unidad"])



def crear_caso_pdf_caltrans() -> dict:
    """
    Crea el caso de validación del ejemplo Caltrans BDP 11.2.

    IMPORTANTE:
    Este diccionario contiene únicamente INSUMOS del ejemplo:
    geometría equivalente, pesos de partes, brazos de momento, empujes obtenidos
    por el método Trial Wedge reportados en el PDF y factores de carga.

    No contiene resultados finales como x, e, B', q o R/H. Esos valores se
    calculan mediante el motor dinámico calcular_estabilidad_externa_desde_caso().
    """
    return {
        "nombre": "Caltrans BDP 11.2 - RC Retaining Wall",
        "B_ft": 19.0,
        "phi_deg": 34.0,
        "kh": 0.28,
        "delta_A_deg": 9.74,
        "delta_P_deg": 22.67,
        "PA_kip": 17.54,
        "PAE_kip": 31.56,
        "PP_kip": 9.97,
        "PPE_kip": 8.00,
        "qn_extreme_ksf": 33.22,
        "qn_strength_ksf": 57.78,
        "q_overburden_service_ksf": 0.54,
        "bearing_phi_extreme": 0.80,
        "bearing_phi_strength": 0.55,
        "friction_phi": 1.00,
        "passive_phi_strength": 0.50,
        "parts": [
            {"id": "1", "tipo": "concreto", "W_kip": 3.23, "x_ft": 5.97, "y_inercia_ft": 14.00},
            {"id": "2", "tipo": "concreto", "W_kip": 2.48, "x_ft": 6.92, "y_inercia_ft": 10.17},
            {"id": "3", "tipo": "concreto", "W_kip": 7.12, "x_ft": 9.50, "y_inercia_ft": 1.25},
            {"id": "4", "tipo": "concreto", "W_kip": 0.22, "x_ft": 14.00, "y_inercia_ft": 0.38},
            {"id": "5", "tipo": "suelo", "W_kip": 1.82, "x_ft": 7.42, "y_inercia_ft": 17.17},
            {"id": "6", "tipo": "suelo", "W_kip": 29.37, "x_ft": 13.44, "y_inercia_ft": 13.50},
            {"id": "7", "tipo": "suelo", "W_kip": 1.32, "x_ft": 2.75, "y_inercia_ft": 3.50},
            {"id": "8", "tipo": "suelo", "W_kip": 4.45, "x_ft": 14.83, "y_inercia_ft": 26.48},
        ],
        "pdf_resultados": {
            "Extreme Event": {"x_ft": 3.52, "e_ft": 5.98, "Bp_ft": 7.03, "q_ksf": 7.87, "qr_ksf": 26.58},
            "Service": {"x_ft": 8.79, "e_ft": 0.70, "Bp_ft": 17.59, "qnet_ksf": 2.47},
            "Strength Ia": {"x_ft": 8.50, "Bp_ft": 17.00, "q_ksf": 4.16, "qr_ksf": 31.78},
            "Strength Ib": {"x_ft": 7.44, "e_ft": 2.06, "Q_kip": 25.94, "R_kip": 40.50},
        },
    }


def calcular_estabilidad_externa_desde_caso(caso: dict) -> dict:
    """
    Motor dinámico de estabilidad externa.

    Recibe un caso con pesos, brazos, empujes, factores y propiedades, y calcula:
    - x de la resultante;
    - excentricidad;
    - ancho efectivo B';
    - presión de contacto;
    - resistencia de apoyo;
    - resistencia a deslizamiento.

    Esta función NO usa resultados finales del PDF. Por eso sirve como motor
    de validación: si los insumos son los del PDF, los resultados deben coincidir.
    """
    B = caso["B_ft"]
    parts = caso["parts"]
    delta_A = math.radians(caso["delta_A_deg"])
    delta_P = math.radians(caso["delta_P_deg"])
    phi = math.radians(caso["phi_deg"])

    W_sum = sum(p["W_kip"] for p in parts)

    # -------------------------
    # Extreme Event - Seismic
    # -------------------------
    PAE = caso["PAE_kip"]
    PAE_v = PAE * math.sin(delta_A)
    PAE_h = PAE * math.cos(delta_A)

    M_pesos = sum(p["W_kip"] * p["x_ft"] for p in parts)
    M_PAE_v = PAE_v * B
    M_PAE_h = -PAE_h * 10.15
    M_inercia = -sum(caso["kh"] * p["W_kip"] * p["y_inercia_ft"] for p in parts)

    V_ext = W_sum + PAE_v
    M_ext = M_pesos + M_PAE_v + M_PAE_h + M_inercia
    x_ext = M_ext / V_ext
    e_ext = B / 2.0 - x_ext
    Bp_ext = B - 2.0 * e_ext
    q_ext = V_ext / Bp_ext
    qr_ext = caso["bearing_phi_extreme"] * caso["qn_extreme_ksf"]

    # -------------------------
    # Service
    # -------------------------
    PA = caso["PA_kip"]
    PA_v = PA * math.sin(delta_A)
    PA_h = PA * math.cos(delta_A)

    V_serv = W_sum + PA_v
    M_serv = M_pesos + PA_v * B - PA_h * 10.15
    x_serv = M_serv / V_serv
    e_serv = B / 2.0 - x_serv
    Bp_serv = B - 2.0 * e_serv
    qgross_serv = V_serv / Bp_serv
    qnet_serv = qgross_serv - caso["q_overburden_service_ksf"]

    # -------------------------
    # Strength Ia - Bearing
    # -------------------------
    factores_ia = {"concreto": 1.25, "suelo": 1.35}
    cargas_ia = [p["W_kip"] * factores_ia[p["tipo"]] for p in parts]
    M_ia = sum(c * p["x_ft"] for c, p in zip(cargas_ia, parts))
    V_ia = sum(cargas_ia) + 1.50 * PA_v
    M_ia += 1.50 * PA_v * B - 1.50 * PA_h * 10.15
    x_ia = M_ia / V_ia
    Bp_ia = 2.0 * x_ia
    q_ia = V_ia / Bp_ia
    qr_ia = caso["bearing_phi_strength"] * caso["qn_strength_ksf"]

    # -------------------------
    # Strength Ib - Eccentricity and sliding
    # -------------------------
    factores_ib = {"concreto": 0.90, "suelo": 1.00}
    cargas_ib = [p["W_kip"] * factores_ib[p["tipo"]] for p in parts]
    V_ib = sum(cargas_ib) + 1.50 * PA_v
    M_ib = sum(c * p["x_ft"] for c, p in zip(cargas_ib, parts))
    M_ib += 1.50 * PA_v * B - 1.50 * PA_h * 10.15
    x_ib = M_ib / V_ib
    e_ib = B / 2.0 - x_ib

    Q_activo = 1.50 * PA_h
    R_friccion = caso["friction_phi"] * V_ib * math.tan(phi)
    R_pasivo = caso["passive_phi_strength"] * caso["PP_kip"] * math.cos(delta_P)
    R_total = R_friccion + R_pasivo

    return {
        "Extreme Event": {
            "x_ft": x_ext,
            "e_ft": e_ext,
            "Bp_ft": Bp_ext,
            "q_ksf": q_ext,
            "qr_ksf": qr_ext,
        },
        "Service": {
            "x_ft": x_serv,
            "e_ft": e_serv,
            "Bp_ft": Bp_serv,
            "qnet_ksf": qnet_serv,
        },
        "Strength Ia": {
            "x_ft": x_ia,
            "Bp_ft": Bp_ia,
            "q_ksf": q_ia,
            "qr_ksf": qr_ia,
        },
        "Strength Ib": {
            "x_ft": x_ib,
            "e_ft": e_ib,
            "Q_kip": Q_activo,
            "R_kip": R_total,
        },
    }


def calcular_verificacion_caltrans_pdf() -> dict:
    """
    Ejecuta la validación del ejemplo PDF usando el motor dinámico.

    Esta función se mantiene por compatibilidad con app.py, pero ahora ya no
    calcula con resultados finales quemados. Crea el caso PDF como insumo y
    lo resuelve con calcular_estabilidad_externa_desde_caso().
    """
    caso = crear_caso_pdf_caltrans()
    return calcular_estabilidad_externa_desde_caso(caso)


def tabla_fuerzas_pdf() -> pd.DataFrame:
    """
    Devuelve los insumos principales del PDF para la validación.

    Son valores de entrada del ejemplo Caltrans, no resultados finales de la tabla
    de comparación.
    """
    caso = crear_caso_pdf_caltrans()
    filas = [
        ("ΣW", sum(p["W_kip"] for p in caso["parts"]), "kip", "Peso total calculado con partes 1 a 8"),
        ("PA", caso["PA_kip"], "kip", "Empuje activo estático obtenido por Trial Wedge en el PDF"),
        ("δA", caso["delta_A_deg"], "grados", "Ángulo de acción del empuje activo"),
        ("PAE", caso["PAE_kip"], "kip", "Empuje activo sísmico obtenido por Trial Wedge en el PDF"),
        ("PP", caso["PP_kip"], "kip", "Empuje pasivo estático"),
        ("PPE", caso["PPE_kip"], "kip", "Empuje pasivo sísmico"),
        ("kh", caso["kh"], "-", "Coeficiente sísmico usado"),
        ("B", caso["B_ft"], "ft", "Ancho de zapata del ejemplo"),
    ]
    return pd.DataFrame(filas, columns=["Parámetro", "Valor PDF", "Unidad", "Descripción"])


def tabla_insumos_partes_pdf() -> pd.DataFrame:
    """
    Muestra las partes del muro usadas como entrada para el motor dinámico.
    """
    caso = crear_caso_pdf_caltrans()
    filas = []
    for p in caso["parts"]:
        filas.append((
            p["id"],
            p["tipo"],
            p["W_kip"],
            p["x_ft"],
            p["y_inercia_ft"],
            p["W_kip"] * p["x_ft"],
        ))
    return pd.DataFrame(
        filas,
        columns=["Parte", "Tipo", "W [kip]", "Brazo x [ft]", "Brazo inercia y [ft]", "Momento W*x [kip-ft]"],
    )


def tabla_comparacion_pdf(resultados: dict) -> pd.DataFrame:
    """
    Compara resultados dinámicos calculados contra los valores publicados en el PDF.
    """
    caso = crear_caso_pdf_caltrans()
    pdf = caso["pdf_resultados"]
    filas = []

    for estado, resultados_estado in resultados.items():
        valores_pdf_estado = pdf.get(estado, {})
        for variable, valor_pdf in valores_pdf_estado.items():
            valor_calc = resultados_estado.get(variable)
            if valor_calc is None:
                continue

            diferencia = valor_calc - valor_pdf
            error_pct = diferencia / valor_pdf * 100.0 if valor_pdf != 0 else 0.0
            tolerancia = max(0.02, abs(valor_pdf) * 0.03)

            filas.append((
                estado,
                variable,
                round(valor_calc, 3),
                round(valor_pdf, 3),
                round(diferencia, 3),
                round(error_pct, 2),
                "OK" if abs(diferencia) <= tolerancia else "Revisar",
            ))

    return pd.DataFrame(
        filas,
        columns=["Estado límite", "Variable", "Dinámico calculado", "PDF", "Diferencia", "Error %", "Estado"],
    )



def area_poligono(puntos: list[tuple[float, float]]) -> float:
    """
    Calcula el área de un polígono mediante la fórmula del polígono o fórmula de Gauss.

    Se usa para obtener el área de cada cuña de suelo ensayada en el método Trial Wedge.
    El resultado se entrega en m² por metro de longitud del muro.
    """
    area = 0.0
    n = len(puntos)
    for i in range(n):
        x1, y1 = puntos[i]
        x2, y2 = puntos[(i + 1) % n]
        area += x1 * y2 - x2 * y1
    return abs(area) / 2.0


def calcular_y_terreno(datos: DatosMuro, geometria: dict, x: float) -> float:
    """
    Calcula la elevación del terreno de relleno para una coordenada x.

    Corrección importante:
    - antes la línea de relleno siempre partía desde la corona del muro (y = H);
    - ahora parte desde la altura real ingresada `altura_relleno`.

    Por tanto, si cambias la altura de relleno, cambia:
    - el dibujo;
    - la cuña Trial Wedge;
    - PA;
    - momentos y cortantes del fuste;
    - presiones de contacto;
    - diseño de zapata y dentellón.
    """
    x0 = geometria["x_dorso_fuste_corona"]
    y0 = datos.altura_relleno

    if datos.pendiente_v == 0:
        return y0

    pendiente = datos.pendiente_v / datos.pendiente_h
    return y0 + pendiente * max(x - x0, 0.0)



def calcular_trial_wedge_activo(datos: DatosMuro, geometria: dict, numero_cunas: int = 180) -> tuple[pd.DataFrame, dict]:
    """
    Calcula el empuje activo estático mediante el método Trial Wedge.

    Procedimiento implementado según la lógica del PDF:
    1. Se asumen varias superficies de falla con diferentes ángulos alfa.
    2. Para cada superficie se define una cuña de suelo detrás del muro.
    3. Se calcula el peso W de la cuña.
    4. Se calcula el empuje P asociado a esa cuña.
    5. Se selecciona como PA la cuña que produce el mayor empuje.

    Alcance de esta primera versión:
    - suelo sin cohesión;
    - análisis estático;
    - terreno de relleno lineal;
    - cálculo por metro longitudinal de muro;
    - empuje aplicado sobre un plano vertical en el extremo del talón.

    La expresión de equilibrio usada es la forma clásica de cuña de Coulomb
    aplicada por tanteos. Para relleno irregular, el peso de la cuña se obtiene
    geométricamente.
    """
    phi = math.radians(datos.phi)

    # Plano vertical donde se evalúa el empuje activo: extremo posterior de la zapata.
    x_muro = datos.B
    y_base = -datos.hz
    y_superior = calcular_y_terreno(datos, geometria, x_muro)

    # Si no existe relleno por encima de la base de la zapata, no hay cuña activa.
    if y_superior <= y_base:
        vacio = pd.DataFrame(columns=[
            "alfa_grados", "P_ton_m", "W_ton_m", "area_m2", "x_inter_m", "y_inter_m"
        ])
        return vacio, {
            "PA_ton_m": 0.0,
            "alfa_critico_grados": 0.0,
            "delta_grados": 0.0,
            "peso_cuna_ton_m": 0.0,
            "area_cuna_m2": 0.0,
            "x_interseccion_m": None,
            "y_interseccion_m": None,
            "poligono_cuna": None,
        }

    # Pendiente del relleno usada como dirección aproximada del empuje.
    beta = math.atan(datos.pendiente_v / datos.pendiente_h) if datos.pendiente_v > 0 else 0.0

    # Se evita ensayar ángulos menores al ángulo de fricción.
    alfa_min = max(math.degrees(phi) + 1.0, math.degrees(beta) + 2.0)
    alfa_max = 89.0

    filas = []
    mejor = {
        "PA_ton_m": 0.0,
        "alfa_grados": None,
        "delta_grados": math.degrees(beta),
        "peso_cuna_ton_m": 0.0,
        "area_cuna_m2": 0.0,
        "x_interseccion_m": None,
        "y_interseccion_m": None,
        "poligono_cuna": None,
    }

    for i in range(numero_cunas):
        alfa_g = alfa_min + (alfa_max - alfa_min) * i / max(numero_cunas - 1, 1)
        alfa = math.radians(alfa_g)

        tan_alfa = math.tan(alfa)
        tan_beta = math.tan(beta)

        # Intersección entre superficie de falla y terreno.
        denominador = tan_alfa - tan_beta
        if denominador <= 0:
            continue

        x_inter = x_muro + (y_superior - y_base) / denominador
        if x_inter <= x_muro:
            continue

        y_inter = y_base + tan_alfa * (x_inter - x_muro)

        poligono = [
            (x_muro, y_base),
            (x_muro, y_superior),
            (x_inter, y_inter),
        ]

        area = area_poligono(poligono)
        W = datos.gamma_suelo * area

        numerador = math.sin(alfa - phi)
        denominador_p = math.sin(math.pi / 2.0 + beta + phi - alfa)

        if denominador_p <= 0 or numerador <= 0:
            P = 0.0
        else:
            P = W * numerador / denominador_p

        filas.append({
            "α [grados]": alfa_g,
            "Área cuña [m²/m]": area,
            "W [ton/m]": W,
            "δ asumido [grados]": math.degrees(beta),
            "P [ton/m]": P,
            "x intersección [m]": x_inter,
            "y intersección [m]": y_inter,
        })

        if P > mejor["PA_ton_m"]:
            mejor = {
                "PA_ton_m": P,
                "alfa_grados": alfa_g,
                "delta_grados": math.degrees(beta),
                "peso_cuna_ton_m": W,
                "area_cuna_m2": area,
                "x_interseccion_m": x_inter,
                "y_interseccion_m": y_inter,
                "poligono_cuna": poligono,
            }

    return pd.DataFrame(filas), mejor


def dibujar_trial_wedge_critico(ax, datos: DatosMuro, geometria: dict, resultado_trial: dict, mostrar_ejes: bool = True):
    """
    Dibuja la cuña crítica obtenida con el método Trial Wedge.

    Se muestra la geometría del muro, la superficie del terreno y la superficie
    de falla que generó el mayor empuje PA.
    """
    dibujar_muro(
        ax,
        datos,
        geometria,
        tamano_texto=8,
        mostrar_cotas=False,
        mostrar_ejes=mostrar_ejes
    )

    poligono = resultado_trial.get("poligono_cuna")
    if poligono:
        cuna = Polygon(
            poligono,
            closed=True,
            fill=False,
            linewidth=2.0,
            linestyle="--"
        )
        ax.add_patch(cuna)

        x0, y0 = poligono[0]
        x2, y2 = poligono[2]
        ax.plot([x0, x2], [y0, y2], linewidth=2.0, linestyle="--")
        ax.text(
            (x0 + x2) / 2,
            (y0 + y2) / 2,
            f"α = {resultado_trial['alfa_grados']:.1f}°",
            fontsize=9,
            ha="center",
            va="bottom"
        )

    ax.set_title("Cuña crítica por Trial Wedge")


def graficar_curva_trial_wedge(ax, tabla_trial: pd.DataFrame):
    """
    Grafica la relación entre el ángulo de falla α y el empuje P.

    El máximo de esta curva corresponde al empuje activo PA adoptado.
    """
    ax.plot(tabla_trial["α [grados]"], tabla_trial["P [ton/m]"], linewidth=2)
    ax.set_xlabel("Ángulo de superficie de falla α [grados]")
    ax.set_ylabel("Empuje P [ton/m]")
    ax.set_title("Búsqueda de PA mediante Trial Wedge")
    ax.grid(True, linestyle=":", linewidth=0.7)



def tabla_momentos_estabilidad(presiones: dict) -> pd.DataFrame:
    """
    Tabla de momentos estabilizantes y desestabilizantes usados para la resultante.

    Muestra explícitamente cómo la posición del dentellón modifica su aporte:
    M_dentellón = W_dentellón * x_dentellón.
    """
    filas = [
        ("ESTABILIZANTES", "", "", ""),
        ("Peso zapata", presiones.get("W_zapata_ton_m", 0.0), presiones.get("x_zapata_m", 0.0), presiones.get("M_est_zapata_ton_m_m", 0.0)),
        ("Peso fuste/pantalla", presiones.get("W_fuste_ton_m", 0.0), presiones.get("x_fuste_m", 0.0), presiones.get("M_est_fuste_ton_m_m", 0.0)),
        ("Peso dentellón", presiones.get("W_dentellon_ton_m", 0.0), presiones.get("x_dentellon_m", 0.0), presiones.get("M_est_dentellon_ton_m_m", 0.0)),
        ("Peso suelo sobre talón", presiones.get("W_suelo_talon_ton_m", 0.0), presiones.get("x_suelo_talon_m", 0.0), presiones.get("M_est_suelo_talon_ton_m_m", 0.0)),
        ("Peso adicional por pendiente", presiones.get("W_pendiente_ton_m", 0.0), presiones.get("x_pendiente_m", 0.0), presiones.get("M_est_pendiente_ton_m_m", 0.0)),
        ("Componente vertical PA", presiones.get("PA_v_ton_m", 0.0), None, presiones.get("M_est_PA_v_ton_m_m", 0.0)),
        ("Total estabilizante", "", "", presiones.get("M_est_ton_m_m", 0.0)),
        ("DESESTABILIZANTES", "", "", ""),
        ("Componente horizontal PAh", presiones.get("PA_h_ton_m", 0.0), presiones.get("brazo_PA_h_m", 0.0), presiones.get("M_volc_PA_h_ton_m_m", 0.0)),
        ("Total desestabilizante", "", "", presiones.get("M_volc_ton_m_m", 0.0)),
        ("RESULTANTE", "", "", ""),
        ("M neto = M_est - M_volc", "", "", presiones.get("M_est_ton_m_m", 0.0) - presiones.get("M_volc_ton_m_m", 0.0)),
        ("x resultante", "", "", presiones.get("x_resultante_m", 0.0)),
        ("e", "", "", presiones.get("e_m", 0.0)),
    ]

    return pd.DataFrame(filas, columns=["Concepto", "Fuerza [ton/m]", "Brazo x o y [m]", "Momento [ton·m/m]"])


def recomendar_posicion_dentellon(datos: DatosMuro, modo: str) -> float:
    """
    Devuelve una posición sugerida para el dentellón según el modo seleccionado.
    """
    if modo == "Bajo pantalla":
        return datos.puntera + datos.t_base / 2.0
    if modo == "Según PDF / hacia talón":
        # En el ejemplo del PDF, el dentellón se ubica hacia el talón.
        return min(max(datos.puntera + datos.t_base + 0.70 * calcular_talon(datos), datos.ancho_llave / 2.0), datos.B - datos.ancho_llave / 2.0)
    return datos.pos_llave


def agregar_tabla_dataframe(documento: Document, tabla: pd.DataFrame):
    """
    Inserta un DataFrame como tabla Word.
    """
    tabla_word = documento.add_table(rows=1, cols=len(tabla.columns))
    tabla_word.style = "Table Grid"

    encabezados = tabla_word.rows[0].cells
    for i, columna in enumerate(tabla.columns):
        encabezados[i].text = str(columna)

    for _, fila in tabla.iterrows():
        celdas = tabla_word.add_row().cells
        for i, valor in enumerate(fila):
            if isinstance(valor, float):
                celdas[i].text = f"{valor:.3f}"
            else:
                celdas[i].text = str(valor)


def generar_memoria_word(datos: DatosMuro, incluir_pdf: bool = True) -> bytes:
    """
    Genera una memoria preliminar de cálculo en formato Word.
    """
    documento = Document()

    documento.add_heading("Memoria de cálculo - Muro de contención de hormigón armado", level=0)
    documento.add_paragraph("Documento generado automáticamente desde la aplicación de diseño de muro de contención.")
    documento.add_paragraph(f"Fecha de generación: {datetime.now().strftime('%d/%m/%Y %H:%M')}")

    documento.add_heading("1. Unidades adoptadas", level=1)
    documento.add_paragraph("Longitudes: m")
    documento.add_paragraph("Capacidad admisible del suelo y cohesión: ton/m²")
    documento.add_paragraph("Pesos unitarios: ton/m³")
    documento.add_paragraph("Resistencia del hormigón f'c y fluencia del acero fy: kg/cm²")

    documento.add_heading("2. Datos de entrada", level=1)
    agregar_tabla_dataframe(documento, resumen_geometria(datos))

    documento.add_heading("3. Conversiones internas", level=1)
    agregar_tabla_dataframe(documento, convertir_resistencias_a_sistema_interno(datos))

    if incluir_pdf:
        documento.add_heading("4. Verificación con ejemplo Caltrans BDP 11.2", level=1)
        documento.add_paragraph(
            "Se incluye una tabla de comparación con los resultados principales del ejemplo de muro de contención "
            "de hormigón armado del PDF usado como base de cálculo."
        )
        agregar_tabla_dataframe(documento, tabla_comparacion_pdf(calcular_verificacion_caltrans_pdf()))

    documento.add_heading("5. Método Trial Wedge", level=1)
    documento.add_paragraph(
        "El método Trial Wedge ensaya varias superficies de falla. Para cada cuña se calcula "
        "su peso y el empuje correspondiente. El mayor valor se adopta como empuje activo PA."
    )

    documento.add_heading("6. Alcance actual", level=1)
    documento.add_paragraph(
        "Esta versión incorpora geometría, diagrama de fuerzas y verificación de estabilidad externa "
        "contra los resultados del ejemplo del PDF. En las siguientes etapas se ampliará el cálculo "
        "para diseño interno del fuste, zapata, puntera, talón y llave de corte."
    )

    buffer = BytesIO()
    documento.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()




def area_barra_cm2(diametro_mm: float) -> float:
    """
    Calcula el área de una barra circular de acero en cm².

    El diámetro se ingresa en milímetros porque es la forma usual de especificar
    barras en Ecuador. Internamente se convierte a centímetros.
    """
    diametro_cm = diametro_mm / 10.0
    return math.pi * diametro_cm ** 2 / 4.0


def resolver_as_flexion_rectangular(Mu_ton_m: float, b_cm: float, d_cm: float, fc_kg_cm2: float, fy_kg_cm2: float, phi_flexion: float = 0.90) -> float:
    """
    Calcula el acero requerido As para flexión de una sección rectangular por metro de muro.

    Se resuelve la ecuación:
    phi * As * fy * (d - a/2) >= Mu
    con:
    a = As * fy / (0.85 * f'c * b)

    Unidades:
    - Mu en ton·m por metro de muro
    - b y d en cm
    - f'c y fy en kg/cm²
    - As resultante en cm²/m
    """
    Mu_kg_cm = Mu_ton_m * 100000.0

    A = phi_flexion * fy_kg_cm2 ** 2 / (2.0 * 0.85 * fc_kg_cm2 * b_cm)
    B = -phi_flexion * fy_kg_cm2 * d_cm
    C = Mu_kg_cm

    discriminante = B ** 2 - 4.0 * A * C

    if discriminante < 0:
        return float("nan")

    raiz1 = (-B - math.sqrt(discriminante)) / (2.0 * A)
    raiz2 = (-B + math.sqrt(discriminante)) / (2.0 * A)

    candidatos = [r for r in [raiz1, raiz2] if r > 0]
    return min(candidatos) if candidatos else float("nan")


def seleccionar_separacion(area_barra: float, As_req_cm2_m: float, separacion_max_cm: float = 30.0) -> tuple[float, float]:
    """
    Selecciona una separación práctica para barras de refuerzo.

    La separación teórica se obtiene con:
    s = Ab * 100 / As_req

    Luego se redondea hacia abajo a separaciones comerciales para garantizar
    que el acero provisto sea mayor o igual al requerido.
    """
    if As_req_cm2_m <= 0 or math.isnan(As_req_cm2_m):
        return float("nan"), float("nan")

    separacion_teorica = area_barra * 100.0 / As_req_cm2_m
    separaciones_comerciales = [40, 35, 30, 25, 20, 18, 15, 12.5, 10, 7.5]

    separaciones_validas = [s for s in separaciones_comerciales if s <= separacion_teorica and s <= separacion_max_cm]
    if not separaciones_validas:
        separacion = min(separaciones_comerciales)
    else:
        separacion = max(separaciones_validas)

    As_prov = area_barra * 100.0 / separacion
    return separacion, As_prov


def calcular_diseno_fuste_dinamico(
    datos: DatosMuro,
    numero_cunas: int = 180,
    recubrimiento_cm: float = 7.5,
    diametro_vertical_mm: float = 16.0,
    diametro_horizontal_mm: float = 12.0,
    separacion_max_cm: float = 30.0
) -> dict:
    """
    Diseña dinámicamente el armado preliminar del fuste.

    La demanda se calcula con el procedimiento del PDF para el fuste:
    - el empuje activo se obtiene mediante Trial Wedge;
    - el empuje se aplica a H/3 desde la base del fuste;
    - el momento factorizado se toma como gamma_p * PA * cos(delta) * H/3;
    - el cortante factorizado se toma como gamma_p * PA * cos(delta).

    Esta función recalcula todo cuando cambian geometría, suelo, altura, pendiente,
    f'c o fy.
    """
    geometria = generar_puntos_muro(datos)
    tabla_trial, resultado_trial = calcular_trial_wedge_activo(datos, geometria, numero_cunas=numero_cunas)

    PA = resultado_trial["PA_ton_m"]
    delta = math.radians(resultado_trial["delta_grados"])
    H = datos.H
    gamma_p = 1.50

    q_base_ton_m2 = 2.0 * PA / H if H > 0 else 0.0
    Pu_horizontal = PA * math.cos(delta)

    Mu_ton_m = gamma_p * Pu_horizontal * H / 3.0
    Vu_ton = gamma_p * Pu_horizontal

    b_cm = 100.0
    h_cm = datos.t_base * 100.0
    diametro_vertical_cm = diametro_vertical_mm / 10.0
    d_cm = h_cm - recubrimiento_cm - diametro_vertical_cm / 2.0

    As_flexion = resolver_as_flexion_rectangular(
        Mu_ton_m=Mu_ton_m,
        b_cm=b_cm,
        d_cm=d_cm,
        fc_kg_cm2=datos.fc,
        fy_kg_cm2=datos.fy,
        phi_flexion=0.90
    )

    # Criterio del PDF para temperatura y retracción: As = 0.0018*b*h total.
    # Para una cara se usa la mitad.
    As_temp_total = 0.0018 * b_cm * h_cm
    As_temp_cara = As_temp_total / 2.0

    As_vertical_req = max(As_flexion, As_temp_cara)

    Ab_vertical = area_barra_cm2(diametro_vertical_mm)
    sep_vertical_cm, As_vertical_prov = seleccionar_separacion(
        area_barra=Ab_vertical,
        As_req_cm2_m=As_vertical_req,
        separacion_max_cm=separacion_max_cm
    )

    Ab_horizontal = area_barra_cm2(diametro_horizontal_mm)
    As_horizontal_req = As_temp_cara
    sep_horizontal_cm, As_horizontal_prov = seleccionar_separacion(
        area_barra=Ab_horizontal,
        As_req_cm2_m=As_horizontal_req,
        separacion_max_cm=separacion_max_cm
    )

    # Chequeo preliminar de cortante de concreto, en unidades kgf.
    # Vc = 0.53*sqrt(fc)*b*d, por metro de muro. Se deja como chequeo preliminar.
    phi_cortante = 0.75
    Vc_kgf = 0.53 * math.sqrt(datos.fc) * b_cm * d_cm
    phi_Vc_ton = phi_cortante * Vc_kgf / 1000.0

    estado_cortante = "OK" if Vu_ton <= phi_Vc_ton else "Revisar cortante/refuerzo transversal"

    return {
        "PA_ton_m": PA,
        "delta_grados": resultado_trial["delta_grados"],
        "alfa_critico_grados": resultado_trial["alfa_grados"],
        "q_base_ton_m2": q_base_ton_m2,
        "Mu_ton_m_m": Mu_ton_m,
        "Vu_ton_m": Vu_ton,
        "b_cm": b_cm,
        "h_cm": h_cm,
        "d_cm": d_cm,
        "As_flexion_cm2_m": As_flexion,
        "As_temp_total_cm2_m": As_temp_total,
        "As_temp_cara_cm2_m": As_temp_cara,
        "As_vertical_req_cm2_m": As_vertical_req,
        "diametro_vertical_mm": diametro_vertical_mm,
        "separacion_vertical_cm": sep_vertical_cm,
        "As_vertical_prov_cm2_m": As_vertical_prov,
        "diametro_horizontal_mm": diametro_horizontal_mm,
        "As_horizontal_req_cm2_m": As_horizontal_req,
        "separacion_horizontal_cm": sep_horizontal_cm,
        "As_horizontal_prov_cm2_m": As_horizontal_prov,
        "phi_Vc_ton_m": phi_Vc_ton,
        "estado_cortante": estado_cortante,
        "tabla_trial": tabla_trial,
        "resultado_trial": resultado_trial,
    }


def tabla_diseno_fuste(resultado: dict) -> pd.DataFrame:
    """
    Convierte el resultado del diseño del fuste en una tabla para mostrar en Streamlit.
    """
    filas = [
        ("PA dinámico", resultado["PA_ton_m"], "ton/m"),
        ("δ asumido", resultado["delta_grados"], "grados"),
        ("α crítico", resultado["alfa_critico_grados"], "grados"),
        ("q base", resultado["q_base_ton_m2"], "ton/m²"),
        ("Altura activa de relleno", resultado.get("altura_activa_relleno_m", 0.0), "m"),
        ("Brazo activo", resultado.get("brazo_activo_m", 0.0), "m"),
        ("Mu fuste", resultado["Mu_ton_m_m"], "ton·m/m"),
        ("Vu fuste", resultado["Vu_ton_m"], "ton/m"),
        ("Peralte efectivo d", resultado["d_cm"], "cm"),
        ("As por flexión", resultado["As_flexion_cm2_m"], "cm²/m"),
        ("As temperatura por cara", resultado["As_temp_cara_cm2_m"], "cm²/m"),
        ("As vertical requerido", resultado["As_vertical_req_cm2_m"], "cm²/m"),
        ("As vertical provisto", resultado["As_vertical_prov_cm2_m"], "cm²/m"),
        ("Separación vertical", resultado["separacion_vertical_cm"], "cm"),
        ("As horizontal requerido", resultado["As_horizontal_req_cm2_m"], "cm²/m"),
        ("As horizontal provisto", resultado["As_horizontal_prov_cm2_m"], "cm²/m"),
        ("Separación horizontal", resultado["separacion_horizontal_cm"], "cm"),
        ("φVc preliminar", resultado["phi_Vc_ton_m"], "ton/m"),
        ("Estado cortante", resultado["estado_cortante"], "-"),
    ]
    return pd.DataFrame(filas, columns=["Parámetro", "Valor", "Unidad"])


def dibujar_armado_fuste(ax, datos: DatosMuro, geometria: dict, resultado: dict, mostrar_ejes: bool = True):
    """
    Dibuja un esquema didáctico del armado del fuste.

    Se muestra:
    - acero vertical principal en la cara posterior del fuste;
    - acero horizontal por temperatura/retracción;
    - etiquetas con diámetros y separaciones calculadas.
    """
    dibujar_muro(
        ax,
        datos,
        geometria,
        tamano_texto=8,
        mostrar_cotas=False,
        mostrar_ejes=mostrar_ejes
    )

    x_back_base = geometria["x_dorso_fuste_base"]
    x_back_top = geometria["x_dorso_fuste_corona"]
    x_bar = x_back_base - 0.08

    # Barras verticales principales en cara de relleno.
    n_barras = 6
    for i in range(n_barras):
        t = i / max(n_barras - 1, 1)
        y1 = 0.15 + t * (datos.H - 0.30)
        y2 = min(y1 + datos.H / 5.0, datos.H - 0.05)
        ax.plot([x_bar, x_bar], [y1, y2], linewidth=2.0)

    # Barras horizontales esquemáticas.
    x1 = geometria["x_frente_fuste"] + 0.05
    x2 = geometria["x_dorso_fuste_base"] - 0.05
    niveles = [datos.H * 0.25, datos.H * 0.45, datos.H * 0.65, datos.H * 0.85]
    for y in niveles:
        ax.plot([x1, x2], [y, y], linewidth=1.5, linestyle="--")

    txt_vertical = f"Vertical principal: Ø{resultado['diametro_vertical_mm']:.0f} @ {resultado['separacion_vertical_cm']:.1f} cm"
    txt_horizontal = f"Horizontal temp.: Ø{resultado['diametro_horizontal_mm']:.0f} @ {resultado['separacion_horizontal_cm']:.1f} cm"

    ax.text(
        datos.B * 0.50,
        datos.H * 0.62,
        txt_vertical,
        fontsize=9,
        ha="center",
        va="center"
    )
    ax.text(
        datos.B * 0.50,
        datos.H * 0.52,
        txt_horizontal,
        fontsize=9,
        ha="center",
        va="center"
    )

    ax.set_title("Armado preliminar dinámico del fuste")




def calcular_presiones_contacto_servicio(datos: DatosMuro, numero_cunas: int = 180) -> dict:
    """
    Calcula presiones de contacto dinámicas bajo la zapata en estado de servicio.

    Se sigue la lógica del PDF para equilibrio externo:
    - se calcula el empuje activo PA;
    - se descompone en componente horizontal y vertical;
    - se suman pesos propios simplificados del muro y suelo sobre el talón;
    - se calcula la posición de la resultante x desde la puntera;
    - se obtiene la excentricidad e = B/2 - x;
    - se calculan presiones lineales qmax y qmin.

    Esta función es dinámica y cambia con geometría, suelo y materiales.
    """
    geometria = generar_puntos_muro(datos)
    tabla_trial, resultado_trial = calcular_trial_wedge_activo(datos, geometria, numero_cunas=numero_cunas)

    PA = resultado_trial["PA_ton_m"]
    delta = math.radians(resultado_trial["delta_grados"])
    PA_v = PA * math.sin(delta)
    PA_h = PA * math.cos(delta)

    # Pesos por metro longitudinal.
    # Se toman momentos alrededor de la puntera inferior O, con signo:
    # + estabilizante horario por cargas verticales hacia abajo ubicadas a x>0
    # - desestabilizante antihorario por empujes horizontales.
    area_zapata = datos.B * datos.hz
    W_zapata = area_zapata * datos.gamma_hormigon
    x_zapata = datos.B / 2.0

    area_fuste = (datos.t_base + datos.t_corona) / 2.0 * datos.H
    W_fuste = area_fuste * datos.gamma_hormigon
    x_fuste = datos.puntera + (datos.t_base + datos.t_corona) / 4.0

    # Peso propio del dentellón. Su posición sí modifica el momento estabilizante.
    if datos.usar_llave:
        W_dentellon = datos.ancho_llave * datos.profundidad_llave * datos.gamma_hormigon
        x_dentellon = datos.pos_llave
    else:
        W_dentellon = 0.0
        x_dentellon = 0.0

    talon = calcular_talon(datos)

    # Suelo real sobre el talón según altura_relleno y pendiente.
    # Antes se usaba max(altura_relleno, H), lo cual hacía que bajar el relleno
    # no redujera el peso ni los momentos estabilizantes.
    suelo_talon = suelo_sobre_talon(datos, geometria)
    area_suelo_talon = suelo_talon["area_m2"]
    W_suelo_talon = suelo_talon["W_ton_m"]
    x_suelo_talon = suelo_talon["x_centroide_m"]

    # La pendiente ya está incorporada en el trapecio de suelo_sobre_talon.
    W_pendiente = 0.0
    x_pendiente = x_suelo_talon

    # Momento estabilizante por cargas verticales.
    M_est_zapata = W_zapata * x_zapata
    M_est_fuste = W_fuste * x_fuste
    M_est_dentellon = W_dentellon * x_dentellon
    M_est_suelo_talon = W_suelo_talon * x_suelo_talon
    M_est_pendiente = W_pendiente * x_pendiente
    M_est_PA_v = PA_v * datos.B

    V = W_zapata + W_fuste + W_dentellon + W_suelo_talon + W_pendiente + PA_v
    M_est = (
        M_est_zapata
        + M_est_fuste
        + M_est_dentellon
        + M_est_suelo_talon
        + M_est_pendiente
        + M_est_PA_v
    )

    # Momento volcador por componente horizontal del empuje.
    brazo_PA_h = datos.H / 3.0
    M_volc_PA_h = PA_h * brazo_PA_h
    M_volc = M_volc_PA_h

    M_resultante = M_est - M_volc
    x_resultante = M_resultante / V if V > 0 else float("nan")
    e = datos.B / 2.0 - x_resultante

    if abs(e) <= datos.B / 6.0:
        q_prom = V / datos.B
        q_max = q_prom * (1.0 + 6.0 * e / datos.B)
        q_min = q_prom * (1.0 - 6.0 * e / datos.B)
    else:
        B_efectivo = 3.0 * (datos.B / 2.0 - abs(e))
        q_max = 2.0 * V / B_efectivo if B_efectivo > 0 else float("nan")
        q_min = 0.0

    return {
        "PA_ton_m": PA,
        "PA_h_ton_m": PA_h,
        "PA_v_ton_m": PA_v,
        "W_zapata_ton_m": W_zapata,
        "x_zapata_m": x_zapata,
        "M_est_zapata_ton_m_m": M_est_zapata,
        "W_fuste_ton_m": W_fuste,
        "x_fuste_m": x_fuste,
        "M_est_fuste_ton_m_m": M_est_fuste,
        "W_dentellon_ton_m": W_dentellon,
        "x_dentellon_m": x_dentellon,
        "M_est_dentellon_ton_m_m": M_est_dentellon,
        "W_suelo_talon_ton_m": W_suelo_talon,
        "x_suelo_talon_m": x_suelo_talon,
        "area_suelo_talon_m2": area_suelo_talon,
        "h_suelo_talon_inicio_m": suelo_talon.get("h_inicio_m", 0.0),
        "h_suelo_talon_fin_m": suelo_talon.get("h_fin_m", 0.0),
        "M_est_suelo_talon_ton_m_m": M_est_suelo_talon,
        "W_pendiente_ton_m": W_pendiente,
        "x_pendiente_m": x_pendiente,
        "M_est_pendiente_ton_m_m": M_est_pendiente,
        "M_est_PA_v_ton_m_m": M_est_PA_v,
        "V_total_ton_m": V,
        "M_est_ton_m_m": M_est,
        "brazo_PA_h_m": brazo_PA_h,
        "M_volc_PA_h_ton_m_m": M_volc_PA_h,
        "M_volc_ton_m_m": M_volc,
        "x_resultante_m": x_resultante,
        "e_m": e,
        "qmax_ton_m2": q_max,
        "qmin_ton_m2": q_min,
        "q_adm_ton_m2": datos.qa,
        "estado_q": "OK" if q_max <= datos.qa and q_min >= 0 else "Revisar",
        "resultado_trial": resultado_trial,
    }


def presion_lineal_en_x(datos: DatosMuro, presiones: dict, x: float) -> float:
    """
    Interpola la presión de contacto bajo la zapata en una posición x.

    Convención:
    - x = 0 en la puntera.
    - x = B en el extremo posterior.
    """
    qmax = presiones["qmax_ton_m2"]
    qmin = presiones["qmin_ton_m2"]
    # Se asume qmax en la puntera cuando e positivo. Si e negativo, se invierte.
    if presiones["e_m"] >= 0:
        return qmax + (qmin - qmax) * (x / datos.B)
    return qmin + (qmax - qmin) * (x / datos.B)


def calcular_diseno_zapata_dinamico(
    datos: DatosMuro,
    numero_cunas: int = 180,
    recubrimiento_cm: float = 7.5,
    diametro_puntera_mm: float = 16.0,
    diametro_talon_mm: float = 16.0,
    separacion_max_cm: float = 30.0
) -> dict:
    """
    Diseña dinámicamente el armado preliminar de puntera y talón.

    Enfoque:
    - Se calculan presiones de contacto dinámicas bajo la zapata.
    - La puntera se modela como voladizo desde la cara frontal del fuste.
    - El talón se modela como voladizo desde la cara posterior del fuste.
    - Se obtienen momentos por metro y se calcula As por flexión.

    Este módulo es preliminar, pero cambia automáticamente con geometría, suelo,
    materiales y empuje calculado por Trial Wedge.
    """
    presiones = calcular_presiones_contacto_servicio(datos, numero_cunas=numero_cunas)

    # Factor de carga para pasar a una demanda conservadora preliminar.
    gamma_u = 1.50

    # Puntera: carga hacia arriba por presión de suelo.
    Lp = datos.puntera
    q0 = presion_lineal_en_x(datos, presiones, 0.0)
    q1 = presion_lineal_en_x(datos, presiones, datos.puntera)
    q_prom_p = (q0 + q1) / 2.0
    Mu_puntera = gamma_u * q_prom_p * Lp ** 2 / 2.0
    Vu_puntera = gamma_u * q_prom_p * Lp

    # Talón: peso de suelo sobre talón menos presión de contacto hacia arriba.
    Lt = calcular_talon(datos)
    x_t0 = datos.puntera + datos.t_base
    x_t1 = datos.B
    q_t0 = presion_lineal_en_x(datos, presiones, x_t0)
    q_t1 = presion_lineal_en_x(datos, presiones, x_t1)
    q_prom_t = (q_t0 + q_t1) / 2.0

    altura_suelo = max(datos.altura_relleno, datos.H)
    w_suelo_talon = datos.gamma_suelo * altura_suelo
    if datos.pendiente_v > 0:
        w_suelo_talon += datos.gamma_suelo * (datos.pendiente_v / datos.pendiente_h) * Lt / 2.0

    w_neto_talon = max(w_suelo_talon - q_prom_t, 0.0)
    Mu_talon = gamma_u * w_neto_talon * Lt ** 2 / 2.0
    Vu_talon = gamma_u * w_neto_talon * Lt

    b_cm = 100.0
    h_cm = datos.hz * 100.0

    d_puntera_cm = h_cm - recubrimiento_cm - (diametro_puntera_mm / 10.0) / 2.0
    d_talon_cm = h_cm - recubrimiento_cm - (diametro_talon_mm / 10.0) / 2.0

    As_puntera = resolver_as_flexion_rectangular(
        Mu_ton_m=Mu_puntera,
        b_cm=b_cm,
        d_cm=d_puntera_cm,
        fc_kg_cm2=datos.fc,
        fy_kg_cm2=datos.fy
    )
    As_talon = resolver_as_flexion_rectangular(
        Mu_ton_m=Mu_talon,
        b_cm=b_cm,
        d_cm=d_talon_cm,
        fc_kg_cm2=datos.fc,
        fy_kg_cm2=datos.fy
    )

    As_temp_zapata = 0.0018 * b_cm * h_cm / 2.0

    As_puntera_req = max(As_puntera, As_temp_zapata)
    As_talon_req = max(As_talon, As_temp_zapata)

    sep_puntera_cm, As_puntera_prov = seleccionar_separacion(
        area_barra=area_barra_cm2(diametro_puntera_mm),
        As_req_cm2_m=As_puntera_req,
        separacion_max_cm=separacion_max_cm
    )
    sep_talon_cm, As_talon_prov = seleccionar_separacion(
        area_barra=area_barra_cm2(diametro_talon_mm),
        As_req_cm2_m=As_talon_req,
        separacion_max_cm=separacion_max_cm
    )

    return {
        "presiones": presiones,
        "Lp_m": Lp,
        "Lt_m": Lt,
        "q_puntera_prom_ton_m2": q_prom_p,
        "q_talon_prom_ton_m2": q_prom_t,
        "w_suelo_talon_ton_m2": w_suelo_talon,
        "w_neto_talon_ton_m2": w_neto_talon,
        "Mu_puntera_ton_m_m": Mu_puntera,
        "Vu_puntera_ton_m": Vu_puntera,
        "Mu_talon_ton_m_m": Mu_talon,
        "Vu_talon_ton_m": Vu_talon,
        "As_temp_zapata_cm2_m": As_temp_zapata,
        "As_puntera_flexion_cm2_m": As_puntera,
        "As_talon_flexion_cm2_m": As_talon,
        "As_puntera_req_cm2_m": As_puntera_req,
        "As_talon_req_cm2_m": As_talon_req,
        "diametro_puntera_mm": diametro_puntera_mm,
        "diametro_talon_mm": diametro_talon_mm,
        "sep_puntera_cm": sep_puntera_cm,
        "sep_talon_cm": sep_talon_cm,
        "As_puntera_prov_cm2_m": As_puntera_prov,
        "As_talon_prov_cm2_m": As_talon_prov,
    }


def tabla_presiones_contacto(presiones: dict) -> pd.DataFrame:
    """
    Convierte el resultado de presiones de contacto en una tabla.
    """
    filas = [
        ("PA", presiones["PA_ton_m"], "ton/m"),
        ("PA horizontal", presiones["PA_h_ton_m"], "ton/m"),
        ("PA vertical", presiones["PA_v_ton_m"], "ton/m"),
        ("V total", presiones["V_total_ton_m"], "ton/m"),
        ("M estabilizador", presiones["M_est_ton_m_m"], "ton·m/m"),
        ("M volcador", presiones["M_volc_ton_m_m"], "ton·m/m"),
        ("x resultante", presiones["x_resultante_m"], "m"),
        ("excentricidad e", presiones["e_m"], "m"),
        ("qmax", presiones["qmax_ton_m2"], "ton/m²"),
        ("qmin", presiones["qmin_ton_m2"], "ton/m²"),
        ("qa", presiones["q_adm_ton_m2"], "ton/m²"),
        ("Estado", presiones["estado_q"], "-"),
    ]
    return pd.DataFrame(filas, columns=["Parámetro", "Valor", "Unidad"])


def tabla_diseno_zapata(resultado: dict) -> pd.DataFrame:
    """
    Convierte el diseño preliminar de puntera y talón en tabla.
    """
    filas = [
        ("Longitud puntera", resultado["Lp_m"], "m"),
        ("Longitud talón", resultado["Lt_m"], "m"),
        ("q promedio puntera", resultado["q_puntera_prom_ton_m2"], "ton/m²"),
        ("q promedio talón", resultado["q_talon_prom_ton_m2"], "ton/m²"),
        ("w suelo talón", resultado["w_suelo_talon_ton_m2"], "ton/m²"),
        ("w neto talón", resultado["w_neto_talon_ton_m2"], "ton/m²"),
        ("Mu puntera", resultado["Mu_puntera_ton_m_m"], "ton·m/m"),
        ("Mu talón", resultado["Mu_talon_ton_m_m"], "ton·m/m"),
        ("As puntera requerido", resultado["As_puntera_req_cm2_m"], "cm²/m"),
        ("As puntera provisto", resultado["As_puntera_prov_cm2_m"], "cm²/m"),
        ("Separación puntera", resultado["sep_puntera_cm"], "cm"),
        ("As talón requerido", resultado["As_talon_req_cm2_m"], "cm²/m"),
        ("As talón provisto", resultado["As_talon_prov_cm2_m"], "cm²/m"),
        ("Separación talón", resultado["sep_talon_cm"], "cm"),
    ]
    return pd.DataFrame(filas, columns=["Parámetro", "Valor", "Unidad"])


def dibujar_presiones_contacto(ax, datos: DatosMuro, geometria: dict, presiones: dict, mostrar_ejes: bool = True):
    """
    Dibuja la distribución lineal de presiones de contacto bajo la zapata.
    """
    dibujar_muro(ax, datos, geometria, tamano_texto=8, mostrar_cotas=False, mostrar_ejes=mostrar_ejes)

    qmax = presiones["qmax_ton_m2"]
    qmin = presiones["qmin_ton_m2"]
    escala = max(datos.hz * 1.80 / max(qmax, qmin, 1e-6), 0.02)

    if presiones["e_m"] >= 0:
        y0 = -datos.hz - qmax * escala
        y1 = -datos.hz - qmin * escala
    else:
        y0 = -datos.hz - qmin * escala
        y1 = -datos.hz - qmax * escala

    pol = Polygon(
        [(0, -datos.hz), (datos.B, -datos.hz), (datos.B, y1), (0, y0)],
        closed=True,
        fill=False,
        linewidth=2.0,
        linestyle="--"
    )
    ax.add_patch(pol)

    ax.text(0, y0, f"q = {qmax:.2f}", fontsize=8, ha="left", va="top")
    ax.text(datos.B, y1, f"q = {qmin:.2f}", fontsize=8, ha="right", va="top")
    ax.set_title("Presiones de contacto bajo la zapata")


def dibujar_armado_zapata(ax, datos: DatosMuro, geometria: dict, resultado: dict, mostrar_ejes: bool = True):
    """
    Dibuja un esquema didáctico del armado de puntera y talón.
    """
    dibujar_muro(ax, datos, geometria, tamano_texto=8, mostrar_cotas=False, mostrar_ejes=mostrar_ejes)

    y_inf = -datos.hz + 0.08
    y_sup = -0.08

    # Acero inferior puntera.
    ax.plot([0.10, datos.puntera - 0.08], [y_inf, y_inf], linewidth=2.5)
    ax.text(
        datos.puntera / 2,
        y_inf - 0.18,
        f"Puntera inf.: Ø{resultado['diametro_puntera_mm']:.0f} @ {resultado['sep_puntera_cm']:.1f} cm",
        fontsize=8,
        ha="center",
        va="top"
    )

    # Acero superior talón.
    x0 = datos.puntera + datos.t_base + 0.08
    x1 = datos.B - 0.10
    ax.plot([x0, x1], [y_sup, y_sup], linewidth=2.5)
    ax.text(
        (x0 + x1) / 2,
        y_sup + 0.18,
        f"Talón sup.: Ø{resultado['diametro_talon_mm']:.0f} @ {resultado['sep_talon_cm']:.1f} cm",
        fontsize=8,
        ha="center",
        va="bottom"
    )

    ax.set_title("Armado preliminar dinámico de zapata")



def longitud_desarrollo_basica_cm(diametro_mm: float, fy_kg_cm2: float, fc_kg_cm2: float, factor: float = 0.075) -> float:
    """
    Estima una longitud de desarrollo preliminar en cm.

    Fórmula práctica preliminar:
    ld = factor * fy * db / sqrt(f'c)

    donde:
    - db está en cm,
    - fy y f'c en kg/cm².

    Se usa como control constructivo inicial para verificar si las barras principales
    tienen longitud disponible suficiente dentro de la zapata.
    """
    db_cm = diametro_mm / 10.0
    if fc_kg_cm2 <= 0:
        return float("nan")
    return factor * fy_kg_cm2 * db_cm / math.sqrt(fc_kg_cm2)


def verificar_cortante_rectangular(Vu_ton_m: float, b_cm: float, d_cm: float, fc_kg_cm2: float, phi_cortante: float = 0.75) -> dict:
    """
    Verifica el cortante unidireccional para una franja de 1 m de ancho.

    Se usa una resistencia preliminar:
    Vc = 0.53 * sqrt(f'c) * b * d

    con Vc en kgf, b y d en cm, f'c en kg/cm².
    """
    Vc_kgf = 0.53 * math.sqrt(fc_kg_cm2) * b_cm * d_cm
    phi_Vc_ton = phi_cortante * Vc_kgf / 1000.0

    return {
        "Vu_ton_m": Vu_ton_m,
        "phi_Vc_ton_m": phi_Vc_ton,
        "relacion": Vu_ton_m / phi_Vc_ton if phi_Vc_ton > 0 else float("nan"),
        "estado": "OK" if Vu_ton_m <= phi_Vc_ton else "No cumple"
    }


def calcular_diseno_zapata_definitivo(
    datos: DatosMuro,
    numero_cunas: int = 180,
    recubrimiento_cm: float = 7.5,
    diametro_puntera_mm: float = 16.0,
    diametro_talon_mm: float = 16.0,
    separacion_max_cm: float = 30.0
) -> dict:
    """
    Amplía el diseño dinámico de zapata con verificaciones más completas.

    Incluye:
    - flexión de puntera;
    - flexión de talón;
    - cortante de puntera a la cara del fuste;
    - cortante de talón a la cara posterior del fuste;
    - acero mínimo por temperatura/retracción;
    - separación práctica;
    - longitud de desarrollo preliminar;
    - estado global OK / Revisar.

    Nota: sigue siendo un módulo de cálculo programable y editable. Para diseño
    profesional final se debe contrastar con la norma local aplicable y el criterio
    geotécnico del proyecto.
    """
    base_resultado = calcular_diseno_zapata_dinamico(
        datos=datos,
        numero_cunas=numero_cunas,
        recubrimiento_cm=recubrimiento_cm,
        diametro_puntera_mm=diametro_puntera_mm,
        diametro_talon_mm=diametro_talon_mm,
        separacion_max_cm=separacion_max_cm
    )

    presiones = base_resultado["presiones"]

    b_cm = 100.0
    h_cm = datos.hz * 100.0
    d_puntera_cm = h_cm - recubrimiento_cm - (diametro_puntera_mm / 10.0) / 2.0
    d_talon_cm = h_cm - recubrimiento_cm - (diametro_talon_mm / 10.0) / 2.0

    d_puntera_m = d_puntera_cm / 100.0
    d_talon_m = d_talon_cm / 100.0

    gamma_u = 1.50

    # Cortante en puntera: sección crítica a d desde la cara del fuste hacia la puntera.
    Lp_crit = max(datos.puntera - d_puntera_m, 0.0)
    q0 = presion_lineal_en_x(datos, presiones, 0.0)
    qcrit_p = presion_lineal_en_x(datos, presiones, Lp_crit)
    q_prom_crit_p = (q0 + qcrit_p) / 2.0
    Vu_puntera_crit = gamma_u * q_prom_crit_p * Lp_crit

    cortante_puntera = verificar_cortante_rectangular(
        Vu_ton_m=Vu_puntera_crit,
        b_cm=b_cm,
        d_cm=d_puntera_cm,
        fc_kg_cm2=datos.fc
    )

    # Cortante en talón: sección crítica a d desde la cara posterior del fuste hacia el talón.
    Lt = calcular_talon(datos)
    Lt_crit = max(Lt - d_talon_m, 0.0)
    x_inicio_talon_crit = datos.puntera + datos.t_base + d_talon_m
    x_fin_talon = datos.B

    q_tcrit = presion_lineal_en_x(datos, presiones, x_inicio_talon_crit)
    q_tfin = presion_lineal_en_x(datos, presiones, x_fin_talon)
    q_prom_tcrit = (q_tcrit + q_tfin) / 2.0

    altura_suelo = max(datos.altura_relleno, datos.H)
    w_suelo_talon = datos.gamma_suelo * altura_suelo
    if datos.pendiente_v > 0:
        w_suelo_talon += datos.gamma_suelo * (datos.pendiente_v / datos.pendiente_h) * Lt / 2.0

    w_neto_talon_crit = max(w_suelo_talon - q_prom_tcrit, 0.0)
    Vu_talon_crit = gamma_u * w_neto_talon_crit * Lt_crit

    cortante_talon = verificar_cortante_rectangular(
        Vu_ton_m=Vu_talon_crit,
        b_cm=b_cm,
        d_cm=d_talon_cm,
        fc_kg_cm2=datos.fc
    )

    # Longitudes de desarrollo preliminares.
    ld_puntera_cm = longitud_desarrollo_basica_cm(diametro_puntera_mm, datos.fy, datos.fc)
    ld_talon_cm = longitud_desarrollo_basica_cm(diametro_talon_mm, datos.fy, datos.fc)

    longitud_disponible_puntera_cm = max(datos.puntera * 100.0 - recubrimiento_cm, 0.0)
    longitud_disponible_talon_cm = max(Lt * 100.0 - recubrimiento_cm, 0.0)

    estado_ld_puntera = "OK" if longitud_disponible_puntera_cm >= ld_puntera_cm else "Revisar anclaje/gancho"
    estado_ld_talon = "OK" if longitud_disponible_talon_cm >= ld_talon_cm else "Revisar anclaje/gancho"

    estados = [
        presiones["estado_q"],
        cortante_puntera["estado"],
        cortante_talon["estado"],
        estado_ld_puntera,
        estado_ld_talon
    ]
    estado_global = "OK" if all(e == "OK" for e in estados) else "Revisar"

    base_resultado.update({
        "d_puntera_cm": d_puntera_cm,
        "d_talon_cm": d_talon_cm,
        "Lp_crit_m": Lp_crit,
        "Lt_crit_m": Lt_crit,
        "Vu_puntera_crit_ton_m": Vu_puntera_crit,
        "Vu_talon_crit_ton_m": Vu_talon_crit,
        "cortante_puntera": cortante_puntera,
        "cortante_talon": cortante_talon,
        "ld_puntera_cm": ld_puntera_cm,
        "ld_talon_cm": ld_talon_cm,
        "longitud_disponible_puntera_cm": longitud_disponible_puntera_cm,
        "longitud_disponible_talon_cm": longitud_disponible_talon_cm,
        "estado_ld_puntera": estado_ld_puntera,
        "estado_ld_talon": estado_ld_talon,
        "estado_global_zapata": estado_global,
    })

    return base_resultado


def tabla_diseno_zapata_definitivo(resultado: dict) -> pd.DataFrame:
    """
    Tabla ampliada del diseño de zapata con flexión, cortante y anclaje.
    """
    filas = [
        ("Estado global zapata", resultado["estado_global_zapata"], "-"),

        ("qmax", resultado["presiones"]["qmax_ton_m2"], "ton/m²"),
        ("qmin", resultado["presiones"]["qmin_ton_m2"], "ton/m²"),
        ("qa", resultado["presiones"]["q_adm_ton_m2"], "ton/m²"),
        ("Estado presión admisible", resultado["presiones"]["estado_q"], "-"),

        ("Mu puntera", resultado["Mu_puntera_ton_m_m"], "ton·m/m"),
        ("As puntera requerido", resultado["As_puntera_req_cm2_m"], "cm²/m"),
        ("As puntera provisto", resultado["As_puntera_prov_cm2_m"], "cm²/m"),
        ("Armado puntera", f"Ø{resultado['diametro_puntera_mm']:.0f} @ {resultado['sep_puntera_cm']:.1f}", "mm @ cm"),

        ("Mu talón", resultado["Mu_talon_ton_m_m"], "ton·m/m"),
        ("As talón requerido", resultado["As_talon_req_cm2_m"], "cm²/m"),
        ("As talón provisto", resultado["As_talon_prov_cm2_m"], "cm²/m"),
        ("Armado talón", f"Ø{resultado['diametro_talon_mm']:.0f} @ {resultado['sep_talon_cm']:.1f}", "mm @ cm"),

        ("Vu puntera crítico", resultado["Vu_puntera_crit_ton_m"], "ton/m"),
        ("φVc puntera", resultado["cortante_puntera"]["phi_Vc_ton_m"], "ton/m"),
        ("Relación Vu/φVc puntera", resultado["cortante_puntera"]["relacion"], "-"),
        ("Estado cortante puntera", resultado["cortante_puntera"]["estado"], "-"),

        ("Vu talón crítico", resultado["Vu_talon_crit_ton_m"], "ton/m"),
        ("φVc talón", resultado["cortante_talon"]["phi_Vc_ton_m"], "ton/m"),
        ("Relación Vu/φVc talón", resultado["cortante_talon"]["relacion"], "-"),
        ("Estado cortante talón", resultado["cortante_talon"]["estado"], "-"),

        ("ld puntera estimada", resultado["ld_puntera_cm"], "cm"),
        ("Longitud disponible puntera", resultado["longitud_disponible_puntera_cm"], "cm"),
        ("Estado anclaje puntera", resultado["estado_ld_puntera"], "-"),

        ("ld talón estimada", resultado["ld_talon_cm"], "cm"),
        ("Longitud disponible talón", resultado["longitud_disponible_talon_cm"], "cm"),
        ("Estado anclaje talón", resultado["estado_ld_talon"], "-"),
    ]
    return pd.DataFrame(filas, columns=["Verificación", "Valor", "Unidad"])



def tabla_resumen_armado_dentellon(resultado: dict) -> pd.DataFrame:
    """
    Genera una tabla corta solo con el resumen del armado del dentellón.

    Esta tabla reemplaza a la imagen de la pestaña de dentellón, para que la
    revisión sea más limpia. Si el dentellón es pequeño, se informa que no
    requiere armado independiente y se trata monolítico con la zapata.
    """
    if resultado.get("requiere_detalle_viga", False):
        filas = [
            ("Criterio", resultado["criterio_armado_dentellon"], "-"),
            ("As longitudinal requerido", resultado["As_long_dentellon_req_cm2_m"], "cm²/m"),
            ("As longitudinal provisto", resultado["As_long_dentellon_prov_cm2_m"], "cm²/m"),
            ("Barras longitudinales", f"{resultado['n_barras_long_dentellon']}Ø{resultado['diametro_llave_mm']:.0f}", "barras"),
            ("Estribos", f"Ø{resultado['diametro_estribo_mm']:.0f} @ {resultado['sep_estribo_dentellon_cm']:.1f}", "mm @ cm"),
            ("Estado cortante", resultado["cortante_dentellon"]["estado"], "-"),
            ("Estado estribos", resultado["estado_estribos_dentellon"], "-"),
            ("Estado armado", resultado["estado_armado_dentellon"], "-"),
        ]
    else:
        filas = [
            ("Criterio", resultado["criterio_armado_dentellon"], "-"),
            ("Requiere armado independiente", "No", "-"),
            ("Estado", resultado["estado_armado_dentellon"], "-"),
            ("Nota", "Se considera monolítico con la zapata.", "-"),
        ]
    return pd.DataFrame(filas, columns=["Verificación", "Valor", "Unidad"])


def dibujar_detalle_zapata_definitivo(ax, datos: DatosMuro, geometria: dict, resultado: dict, mostrar_ejes: bool = True):
    """
    Dibuja un esquema más completo del armado y puntos de chequeo de cortante/anclaje.
    """
    dibujar_armado_zapata(ax, datos, geometria, resultado, mostrar_ejes=mostrar_ejes)

    d_p = resultado["d_puntera_cm"] / 100.0
    d_t = resultado["d_talon_cm"] / 100.0

    # Sección crítica de cortante en puntera, a d de la cara del fuste.
    x_cp = max(datos.puntera - d_p, 0.0)
    ax.plot([x_cp, x_cp], [-datos.hz, 0], linewidth=1.6, linestyle=":")
    ax.text(x_cp, -datos.hz - 0.15, "Vu puntera @ d", fontsize=8, ha="center", va="top")

    # Sección crítica de cortante en talón, a d de la cara posterior del fuste.
    x_ct = datos.puntera + datos.t_base + d_t
    ax.plot([x_ct, x_ct], [-datos.hz, 0], linewidth=1.6, linestyle=":")
    ax.text(x_ct, -datos.hz - 0.15, "Vu talón @ d", fontsize=8, ha="center", va="top")

    # Longitud de desarrollo esquemática.
    ax.annotate(
        "",
        xy=(0.08, -datos.hz + 0.18),
        xytext=(min(datos.puntera, resultado["ld_puntera_cm"] / 100.0), -datos.hz + 0.18),
        arrowprops=dict(arrowstyle="<->", linewidth=1.0)
    )
    ax.text(
        min(datos.puntera / 2, resultado["ld_puntera_cm"] / 200.0),
        -datos.hz + 0.28,
        "ld puntera",
        fontsize=8,
        ha="center",
        va="bottom"
    )

    x0 = datos.puntera + datos.t_base
    x1 = min(datos.B - 0.08, x0 + resultado["ld_talon_cm"] / 100.0)
    ax.annotate(
        "",
        xy=(x0, -0.18),
        xytext=(x1, -0.18),
        arrowprops=dict(arrowstyle="<->", linewidth=1.0)
    )
    ax.text((x0 + x1) / 2, -0.08, "ld talón", fontsize=8, ha="center", va="bottom")

    ax.set_title("Detalle dinámico de zapata: flexión, cortante y anclaje")



def coeficiente_pasivo_rankine(phi_grados: float) -> float:
    """
    Calcula el coeficiente pasivo Rankine:
    Kp = tan²(45° + φ/2)

    Se usa como aproximación dinámica para estimar la resistencia pasiva frente
    a la zapata y llave de corte.
    """
    phi = math.radians(phi_grados)
    return math.tan(math.radians(45.0) + phi / 2.0) ** 2





def calcular_refuerzo_lateral_dentellon(profundidad_llave_m: float, diametro_lateral_mm: float = 10.0, espaciamiento_objetivo_cm: float = 25.0) -> dict:
    """
    Genera una nota de criterio para dentellones profundos.

    Criterio gráfico/didáctico:
    - si profundidad <= 0.80 m: no se muestra ni reporta acero lateral distribuido;
    - si profundidad > 0.80 m: NO se dibujan puntos laterales para no confundir el detalle.
      Solo se indica una nota: incluir acero adicional en las caras laterales para
      control de agrietamiento, según criterio/norma aplicable.
    """
    if profundidad_llave_m <= 0.80:
        return {
            "requiere": False,
            "diametro_mm": diametro_lateral_mm,
            "n_barras_por_cara": 0,
            "espaciamiento_cm": 0.0,
            "nota": "No aplica: profundidad del dentellón ≤ 0.80 m."
        }

    return {
        "requiere": True,
        "diametro_mm": diametro_lateral_mm,
        "n_barras_por_cara": 0,
        "espaciamiento_cm": 0.0,
        "nota": "Profundidad > 0.80 m: incluir acero adicional distribuido en caras laterales para control de agrietamiento."
    }


def as_min_longitudinal_aci_cm2(b_cm: float, d_cm: float, fc_kg_cm2: float, fy_kg_cm2: float) -> float:
    """
    Acero longitudinal mínimo tipo ACI para una sección rectangular.

    Se usa como mínimo constructivo del dentellón armado.
    As_min = max(0.25*sqrt(fc')/fy, 1.4/fy) * b * d
    con fc' y fy en MPa, b y d en mm. Resultado en cm².
    """
    fc_mpa = fc_kg_cm2 * 0.0980665
    fy_mpa = fy_kg_cm2 * 0.0980665
    b_mm = b_cm * 10.0
    d_mm = d_cm * 10.0
    if fc_mpa <= 0 or fy_mpa <= 0 or b_mm <= 0 or d_mm <= 0:
        return 0.0
    rho_min = max(0.25 * math.sqrt(fc_mpa) / fy_mpa, 1.4 / fy_mpa)
    return rho_min * b_mm * d_mm / 100.0


def diseno_cortante_dentellon_aci(
    Vu_ton_m: float,
    b_cm: float,
    d_cm: float,
    fc_kg_cm2: float,
    fy_kg_cm2: float,
    diametro_estribo_mm: float,
    separacion_max_usuario_cm: float,
    separacion_max_dentellon_cm: float = 10.0,
    phi_cortante: float = 0.75,
    lambda_concreto: float = 1.0
) -> dict:
    """
    Diseño preliminar de estribos del dentellón por cortante.

    Vc = 0.17 lambda sqrt(fc') bw d
    Vs_req = Vu/phi - Vc
    Av/s = Vs/(fy d)
    """
    fc_mpa = fc_kg_cm2 * 0.0980665
    fy_mpa = fy_kg_cm2 * 0.0980665
    b_mm = b_cm * 10.0
    d_mm = d_cm * 10.0
    Av_estribo_cm2 = 2.0 * area_barra_cm2(diametro_estribo_mm)

    if fc_mpa <= 0 or fy_mpa <= 0 or b_mm <= 0 or d_mm <= 0:
        return {
            "Vc_ton_m": 0.0, "phi_Vc_ton_m": 0.0, "Vs_req_ton_m": 0.0,
            "Av_s_req_cm2_cm": 0.0, "Av_estribo_cm2": Av_estribo_cm2,
            "s_calc_cm": 0.0, "s_adoptado_cm": 0.0,
            "s_max_dentellon_cm": separacion_max_dentellon_cm,
            "estado": "Revisar datos"
        }

    Vc_N = 0.17 * lambda_concreto * math.sqrt(fc_mpa) * b_mm * d_mm
    Vc_ton = Vc_N / 9806.65
    phi_Vc_ton = phi_cortante * Vc_ton

    if Vu_ton_m <= phi_Vc_ton:
        Vs_req_ton = 0.0
        Av_s_req_cm2_cm = 0.0
        s_calc_cm = float("inf")
        s_adoptado = min(separacion_max_usuario_cm, separacion_max_dentellon_cm)
        estado = "OK: estribo constructivo"
    else:
        Vs_req_ton = max(Vu_ton_m / phi_cortante - Vc_ton, 0.0)
        Vs_req_N = Vs_req_ton * 9806.65
        Av_s_req_mm2_mm = Vs_req_N / (fy_mpa * d_mm) if fy_mpa * d_mm > 0 else float("inf")
        Av_s_req_cm2_cm = Av_s_req_mm2_mm
        s_calc_mm = (Av_estribo_cm2 * 100.0) / Av_s_req_mm2_mm if Av_s_req_mm2_mm > 0 else float("inf")
        s_calc_cm = s_calc_mm / 10.0
        s_adoptado = min(s_calc_cm, separacion_max_usuario_cm, separacion_max_dentellon_cm)
        estado = "OK" if s_adoptado <= s_calc_cm + 1e-9 else "Revisar"

    return {
        "Vc_ton_m": Vc_ton,
        "phi_Vc_ton_m": phi_Vc_ton,
        "Vs_req_ton_m": Vs_req_ton,
        "Av_s_req_cm2_cm": Av_s_req_cm2_cm,
        "Av_estribo_cm2": Av_estribo_cm2,
        "s_calc_cm": s_calc_cm,
        "s_adoptado_cm": s_adoptado,
        "s_max_dentellon_cm": separacion_max_dentellon_cm,
        "estado": estado
    }



def calcular_deslizamiento_y_llave(
    datos: DatosMuro,
    numero_cunas: int = 180,
    recubrimiento_cm: float = 7.5,
    diametro_llave_mm: float = 12.0,
    separacion_max_cm: float = 30.0,
    factor_resistencia_friccion: float = 1.00,
    factor_resistencia_pasiva: float = 0.50,
    factor_carga_empuje: float = 1.50,
    diametro_estribo_mm: float = 8.0
) -> dict:
    """
    Calcula deslizamiento, resistencia pasiva y diseño del dentellón.

    Criterio:
    - Dentellón pequeño: monolítico con zapata, sin armado independiente.
    - Dentellón armado: llave de corte. Longitudinal mínimo ACI y estribos por cortante.
    """
    presiones = calcular_presiones_contacto_servicio(datos, numero_cunas=numero_cunas)

    PA_h = presiones["PA_h_ton_m"]
    V_total = presiones["V_total_ton_m"]
    H_actuante = factor_carga_empuje * PA_h

    R_friccion_nominal = datos.mu * V_total
    R_friccion_diseno = factor_resistencia_friccion * R_friccion_nominal

    Kp = coeficiente_pasivo_rankine(datos.phi)
    altura_pasiva_zapata = datos.hz
    altura_pasiva_total = datos.hz + (datos.profundidad_llave if datos.usar_llave else 0.0)

    PP_zapata_nominal = 0.5 * Kp * datos.gamma_suelo * altura_pasiva_zapata ** 2
    PP_total_nominal = 0.5 * Kp * datos.gamma_suelo * altura_pasiva_total ** 2
    PP_dentellon_extra_nominal = max(PP_total_nominal - PP_zapata_nominal, 0.0)
    PP_diseno = factor_resistencia_pasiva * PP_total_nominal

    R_total = R_friccion_diseno + PP_diseno
    FS_deslizamiento = R_total / H_actuante if H_actuante > 0 else float("inf")
    estado_deslizamiento = "OK" if R_total >= H_actuante else "No cumple"

    umbral_dentellon_armado_m = 0.35
    requiere_detalle_viga = bool(datos.usar_llave and datos.profundidad_llave >= umbral_dentellon_armado_m)

    if datos.usar_llave and datos.profundidad_llave > 0:
        h_key = datos.profundidad_llave
        p_base_dentellon = Kp * datos.gamma_suelo * h_key
        P_dentellon = 0.5 * p_base_dentellon * h_key
        Vu_dentellon = factor_carga_empuje * P_dentellon
        Mu_dentellon = 0.0

        b_cm = 100.0
        h_cm = datos.ancho_llave * 100.0
        d_cm = max(h_cm - recubrimiento_cm - (diametro_llave_mm / 10.0) / 2.0, 1.0)

        if requiere_detalle_viga:
            As_long_req = as_min_longitudinal_aci_cm2(b_cm, d_cm, datos.fc, datos.fy)
            Ab_long = area_barra_cm2(diametro_llave_mm)
            n_barras_long = max(4, int(math.ceil(As_long_req / Ab_long)))
            As_long_prov = n_barras_long * Ab_long

            cortante = diseno_cortante_dentellon_aci(
                Vu_ton_m=Vu_dentellon,
                b_cm=b_cm,
                d_cm=d_cm,
                fc_kg_cm2=datos.fc,
                fy_kg_cm2=datos.fy,
                diametro_estribo_mm=diametro_estribo_mm,
                separacion_max_usuario_cm=separacion_max_cm,
                separacion_max_dentellon_cm=10.0,
                phi_cortante=0.75
            )
            sep_estribo_cm = cortante["s_adoptado_cm"]
            estado_estribos = cortante["estado"]
            estado_armado = "OK" if estado_estribos.startswith("OK") else "Revisar"
            criterio_armado = "Dentellón armado: llave de corte con longitudinales mínimos ACI y estribos por cortante."
            lateral = calcular_refuerzo_lateral_dentellon(datos.profundidad_llave, diametro_lateral_mm=diametro_llave_mm)
            ld_long_cm = longitud_desarrollo_basica_cm(diametro_llave_mm, datos.fy, datos.fc)
            estado_anclaje = "Revisar en extremos/traslapes longitudinales"
        else:
            As_long_req = 0.0
            n_barras_long = 0
            As_long_prov = 0.0
            cortante = {
                "Vc_ton_m": 0.0, "phi_Vc_ton_m": 0.0, "Vs_req_ton_m": 0.0,
                "Av_s_req_cm2_cm": 0.0, "Av_estribo_cm2": 0.0,
                "s_calc_cm": 0.0, "s_adoptado_cm": 0.0, "s_max_dentellon_cm": 10.0,
                "estado": "No aplica: dentellón pequeño monolítico"
            }
            sep_estribo_cm = 0.0
            estado_estribos = "No aplica"
            estado_armado = "No requiere diseño independiente"
            criterio_armado = "Dentellón pequeño: tratar monolítico con zapata, como en el PDF."
            lateral = calcular_refuerzo_lateral_dentellon(datos.profundidad_llave, diametro_lateral_mm=diametro_llave_mm)
            ld_long_cm = 0.0
            estado_anclaje = "No aplica"
    else:
        h_key = 0.0
        p_base_dentellon = 0.0
        P_dentellon = 0.0
        Vu_dentellon = 0.0
        Mu_dentellon = 0.0
        d_cm = 0.0
        As_long_req = 0.0
        n_barras_long = 0
        As_long_prov = 0.0
        cortante = {
            "Vc_ton_m": 0.0, "phi_Vc_ton_m": 0.0, "Vs_req_ton_m": 0.0,
            "Av_s_req_cm2_cm": 0.0, "Av_estribo_cm2": 0.0,
            "s_calc_cm": 0.0, "s_adoptado_cm": 0.0, "s_max_dentellon_cm": 10.0,
            "estado": "No aplica"
        }
        sep_estribo_cm = 0.0
        estado_estribos = "No aplica"
        estado_armado = "No aplica"
        criterio_armado = "No existe dentellón."
        lateral = {'requiere': False, 'diametro_mm': diametro_llave_mm, 'n_barras_por_cara': 0, 'espaciamiento_cm': 0.0, 'nota': 'No aplica'}
        ld_long_cm = 0.0
        estado_anclaje = "No aplica"

    estado_global = "OK" if estado_deslizamiento == "OK" and estado_armado in [
        "OK", "No requiere diseño independiente", "No aplica"
    ] else "Revisar"

    return {
        "PA_h_ton_m": PA_h,
        "V_total_ton_m": V_total,
        "H_actuante_ton_m": H_actuante,
        "mu": datos.mu,
        "R_friccion_nominal_ton_m": R_friccion_nominal,
        "R_friccion_diseno_ton_m": R_friccion_diseno,
        "Kp": Kp,
        "altura_pasiva_zapata_m": altura_pasiva_zapata,
        "altura_pasiva_total_m": altura_pasiva_total,
        "PP_zapata_nominal_ton_m": PP_zapata_nominal,
        "PP_total_nominal_ton_m": PP_total_nominal,
        "PP_llave_extra_nominal_ton_m": PP_dentellon_extra_nominal,
        "PP_dentellon_extra_nominal_ton_m": PP_dentellon_extra_nominal,
        "PP_diseno_ton_m": PP_diseno,
        "R_total_ton_m": R_total,
        "FS_deslizamiento": FS_deslizamiento,
        "estado_deslizamiento": estado_deslizamiento,
        "h_key_m": h_key,
        "p_base_llave_ton_m2": p_base_dentellon,
        "p_base_dentellon_ton_m2": p_base_dentellon,
        "P_llave_ton_m": P_dentellon,
        "P_dentellon_ton_m": P_dentellon,
        "Vu_llave_ton_m": Vu_dentellon,
        "Vu_dentellon_ton_m": Vu_dentellon,
        "Mu_llave_ton_m_m": Mu_dentellon,
        "Mu_dentellon_ton_m_m": Mu_dentellon,
        "d_llave_cm": d_cm,
        "d_dentellon_cm": d_cm,
        "As_llave_flexion_cm2_m": 0.0,
        "As_dentellon_flexion_cm2_m": 0.0,
        "As_min_llave_cm2_m": As_long_req,
        "As_min_dentellon_cm2_m": As_long_req,
        "As_llave_req_cm2_m": As_long_req,
        "As_long_dentellon_req_cm2_m": As_long_req,
        "diametro_llave_mm": diametro_llave_mm,
        "diametro_estribo_mm": diametro_estribo_mm,
        "n_barras_long_dentellon": n_barras_long,
        "sep_llave_cm": sep_estribo_cm,
        "sep_estribo_dentellon_cm": sep_estribo_cm,
        "As_llave_prov_cm2_m": As_long_prov,
        "As_long_dentellon_prov_cm2_m": As_long_prov,
        "cortante_llave": {
            "Vu_ton_m": Vu_dentellon,
            "phi_Vc_ton_m": cortante["phi_Vc_ton_m"],
            "relacion": Vu_dentellon / cortante["phi_Vc_ton_m"] if cortante["phi_Vc_ton_m"] > 0 else 0.0,
            "estado": cortante["estado"]
        },
        "cortante_dentellon": {
            "Vu_ton_m": Vu_dentellon,
            "Vc_ton_m": cortante["Vc_ton_m"],
            "phi_Vc_ton_m": cortante["phi_Vc_ton_m"],
            "Vs_req_ton_m": cortante["Vs_req_ton_m"],
            "Av_s_req_cm2_cm": cortante["Av_s_req_cm2_cm"],
            "Av_estribo_cm2": cortante["Av_estribo_cm2"],
            "s_calc_cm": cortante["s_calc_cm"],
            "s_adoptado_cm": cortante["s_adoptado_cm"],
            "s_max_dentellon_cm": cortante["s_max_dentellon_cm"],
            "estado": cortante["estado"]
        },
        "ld_llave_cm": ld_long_cm,
        "ld_longitudinal_dentellon_cm": ld_long_cm,
        "longitud_disponible_llave_cm": 0.0,
        "estado_ld_llave": estado_anclaje,
        "estado_anclaje_dentellon": estado_anclaje,
        "estado_estribos_dentellon": estado_estribos,
        "refuerzo_lateral_dentellon": lateral,
        "requiere_refuerzo_lateral_dentellon": lateral.get("requiere", False),
        "estado_armado_llave": estado_armado,
        "estado_armado_dentellon": estado_armado,
        "criterio_armado_dentellon": criterio_armado,
        "requiere_detalle_viga": requiere_detalle_viga,
        "estado_global": estado_global,
        "presiones": presiones,
    }



def tabla_deslizamiento_llave(resultado: dict) -> pd.DataFrame:
    """
    Tabla resumen de deslizamiento, resistencia pasiva y dentellón.
    """
    filas = [
        ("Estado global", resultado["estado_global"], "-"),
        ("1. DESLIZAMIENTO", "", ""),
        ("Empuje horizontal PAh", resultado["PA_h_ton_m"], "ton/m"),
        ("Empuje horizontal factorizado", resultado["H_actuante_ton_m"], "ton/m"),
        ("Vertical total V", resultado["V_total_ton_m"], "ton/m"),
        ("Coeficiente fricción μ", resultado["mu"], "-"),
        ("Resistencia por fricción Rf", resultado["R_friccion_diseno_ton_m"], "ton/m"),
        ("Resistencia total R", resultado["R_total_ton_m"], "ton/m"),
        ("FS o relación R/H", resultado["FS_deslizamiento"], "-"),
        ("Estado deslizamiento", resultado["estado_deslizamiento"], "-"),
        ("2. RESISTENCIA PASIVA", "", ""),
        ("Kp Rankine", resultado["Kp"], "-"),
        ("Altura pasiva zapata", resultado["altura_pasiva_zapata_m"], "m"),
        ("Altura pasiva total con dentellón", resultado["altura_pasiva_total_m"], "m"),
        ("PP zapata nominal", resultado["PP_zapata_nominal_ton_m"], "ton/m"),
        ("PP total nominal", resultado["PP_total_nominal_ton_m"], "ton/m"),
        ("Incremento por dentellón", resultado["PP_dentellon_extra_nominal_ton_m"], "ton/m"),
        ("PP de diseño", resultado["PP_diseno_ton_m"], "ton/m"),
        ("3. DENTELLÓN COMO LLAVE DE CORTE", "", ""),
        ("Criterio", resultado["criterio_armado_dentellon"], "-"),
        ("Requiere armado independiente", "Sí" if resultado["requiere_detalle_viga"] else "No", "-"),
        ("Profundidad dentellón", resultado["h_key_m"], "m"),
        ("Presión pasiva base dentellón", resultado["p_base_dentellon_ton_m2"], "ton/m²"),
        ("Fuerza pasiva sobre dentellón", resultado["P_dentellon_ton_m"], "ton/m"),
        ("Vu dentellón", resultado["Vu_dentellon_ton_m"], "ton/m"),
    ]

    if resultado["requiere_detalle_viga"]:
        cd = resultado["cortante_dentellon"]
        s_calc = cd["s_calc_cm"] if math.isfinite(cd["s_calc_cm"]) else "No requiere Vs"
        filas.extend([
            ("4. ARMADO DEL DENTELLÓN", "", ""),
            ("d dentellón", resultado["d_dentellon_cm"], "cm"),
            ("As longitudinal mínimo ACI", resultado["As_min_dentellon_cm2_m"], "cm²/m"),
            ("As longitudinal provisto", resultado["As_long_dentellon_prov_cm2_m"], "cm²/m"),
            ("Longitudinales mínimos", f"{resultado['n_barras_long_dentellon']}Ø{resultado['diametro_llave_mm']:.0f}", "barras"),
            ("Vc concreto", cd["Vc_ton_m"], "ton/m"),
            ("φVc", cd["phi_Vc_ton_m"], "ton/m"),
            ("Vs requerido", cd["Vs_req_ton_m"], "ton/m"),
            ("Av/s requerido", cd["Av_s_req_cm2_cm"], "cm²/cm"),
            ("Av estribo 2 ramas", cd["Av_estribo_cm2"], "cm²"),
            ("s calculado", s_calc, "cm"),
            ("s máximo dentellón", cd["s_max_dentellon_cm"], "cm"),
            ("Estribos adoptados", f"Ø{resultado['diametro_estribo_mm']:.0f} @ {resultado['sep_estribo_dentellon_cm']:.1f}", "mm @ cm"),
            ("Estado cortante/estribos", cd["estado"], "-"),
            ("Nota anclaje longitudinal", resultado["estado_anclaje_dentellon"], "-"),
        ])
    else:
        filas.extend([
            ("4. ARMADO DEL DENTELLÓN", "", ""),
            ("Estado armado dentellón", resultado["estado_armado_dentellon"], "-"),
            ("Nota", "No se diseña acero independiente; se considera monolítico con la zapata.", "-"),
        ])

    return pd.DataFrame(filas, columns=["Verificación", "Valor", "Unidad"])


def dibujar_deslizamiento_pasivo_llave(ax, datos: DatosMuro, geometria: dict, resultado: dict, mostrar_ejes: bool = True):
    """
    Dibuja deslizamiento, resistencia pasiva y detalle corregido del dentellón.

    El dentellón se representa como viga corrida bajo la zapata cuando requiere
    detalle: longitudinales acostadas a lo largo del muro + estribos cerrados.
    """
    dibujar_muro(ax, datos, geometria, tamano_texto=8, mostrar_cotas=False, mostrar_ejes=mostrar_ejes)

    # Empuje activo horizontal.
    y_emp = datos.H * 0.35
    flecha(ax, (datos.B + 0.85, y_emp), (datos.B + 0.10, y_emp), "PAh", dx=0.10, dy=0.22)

    # Fricción basal.
    y_base = -datos.hz - 0.08
    flecha(ax, (datos.B * 0.35, y_base), (datos.B * 0.78, y_base), "Rf = μV", dx=0.05, dy=0.18)

    # Resistencia pasiva frontal.
    flecha(ax, (-0.80, -datos.hz * 0.45), (0.03, -datos.hz * 0.45), "PP", dx=-0.12, dy=-0.22)

    if datos.usar_llave and "llave" in geometria:
        x1 = datos.pos_llave - datos.ancho_llave / 2
        x2 = datos.pos_llave + datos.ancho_llave / 2
        xmid = (x1 + x2) / 2

        flecha(
            ax,
            (xmid - 0.85, -datos.hz - datos.profundidad_llave * 0.55),
            (x1 + 0.02, -datos.hz - datos.profundidad_llave * 0.55),
            "PP dentellón",
            dx=-0.05,
            dy=-0.22
        )

        # Estribos esquemáticos en corte transversal.
        n_est = 3
        for i in range(n_est):
            yy = -datos.hz - datos.profundidad_llave + 0.06 + i * max(datos.profundidad_llave - 0.12, 0.02) / max(n_est - 1, 1)
            ax.plot([x1 + 0.05, x2 - 0.05], [yy, yy], linewidth=1.0, linestyle="--")

        if resultado.get("requiere_detalle_viga", False):
            ax.text(
                xmid,
                -datos.hz - datos.profundidad_llave - 0.18,
                f"Dentellón: {resultado['n_barras_long_dentellon']}Ø{resultado['diametro_llave_mm']:.0f} long. + Estr. Ø{resultado['diametro_estribo_mm']:.0f}@{resultado['sep_estribo_dentellon_cm']:.1f}",
                fontsize=8,
                ha="center",
                va="top"
            )

            # Detalle ampliado tipo viga corrida
            det_x0 = datos.B + 0.65
            det_y0 = max(datos.H * 0.35, 1.30)
            det_b = max(datos.ancho_llave * 1.65, 0.90)
            det_h = max(datos.profundidad_llave * 2.2, 1.00)

            rect = [(det_x0, det_y0), (det_x0 + det_b, det_y0), (det_x0 + det_b, det_y0 - det_h), (det_x0, det_y0 - det_h)]
            ax.add_patch(Polygon(rect, closed=True, fill=False, linewidth=1.8))

            # Estribo cerrado ampliado
            rec = 0.10
            estribo = [
                (det_x0 + rec, det_y0 - rec),
                (det_x0 + det_b - rec, det_y0 - rec),
                (det_x0 + det_b - rec, det_y0 - det_h + rec),
                (det_x0 + rec, det_y0 - det_h + rec),
            ]
            ax.add_patch(Polygon(estribo, closed=True, fill=False, linewidth=1.3, linestyle="--"))

            # Longitudinales como puntos en sección de viga
            xs = [det_x0 + det_b * 0.30, det_x0 + det_b * 0.70]
            ys = [det_y0 - det_h * 0.25, det_y0 - det_h * 0.75]
            for xx in xs:
                for yy in ys:
                    ax.scatter([xx], [yy], s=24)

            ax.text(det_x0 + det_b / 2, det_y0 + 0.13, "Dentellón tipo viga corrida", fontsize=9, ha="center", va="bottom")
            ax.text(det_x0 + det_b / 2, det_y0 - det_h - 0.10, f"Longitudinales: {resultado['n_barras_long_dentellon']}Ø{resultado['diametro_llave_mm']:.0f}", fontsize=8, ha="center", va="top")
            ax.text(det_x0 + det_b / 2, det_y0 - det_h - 0.24, f"Estribos cerrados: Ø{resultado['diametro_estribo_mm']:.0f} @ {resultado['sep_estribo_dentellon_cm']:.1f} cm", fontsize=8, ha="center", va="top")
        else:
            ax.text(
                xmid,
                -datos.hz - datos.profundidad_llave - 0.18,
                "Dentellón pequeño: monolítico con zapata\nsin armado independiente",
                fontsize=8,
                ha="center",
                va="top"
            )

    ax.set_title("Deslizamiento, resistencia pasiva y dentellón tipo viga corrida")



def dibujar_detalle_general_armado(ax, datos: DatosMuro, geometria: dict, resultado_fuste: dict, resultado_zapata: dict, resultado_dentellon: dict, mostrar_ejes: bool = True):
    """
    Dibuja un detalle general de armado del muro de retención.

    El esquema muestra en una sola imagen:
    - armado de la pantalla o fuste (barras verticales principales y horizontales);
    - armado de la zapata (puntera y talón);
    - armado del dentellón, incluyendo barras longitudinales y estribos;
    - un pequeño detalle ampliado del dentellón para visualizar mejor el armado.

    Todo el texto se alimenta de los resultados dinámicos calculados en la app.
    """
    dibujar_muro(ax, datos, geometria, tamano_texto=8, mostrar_cotas=False, mostrar_ejes=mostrar_ejes)

    # -----------------------------
    # 1) ARMADO DE LA PANTALLA
    # -----------------------------
    x_back = geometria["x_dorso_fuste_base"] - 0.06
    x_front = geometria["x_frente_fuste"] + 0.06

    # Barras verticales principales en la cara del relleno.
    for x in [x_back - 0.03, x_back]:
        ax.plot([x, x], [0.12, datos.H - 0.10], linewidth=2.0)

    # Barras horizontales de distribución.
    for y in [datos.H * 0.18, datos.H * 0.35, datos.H * 0.52, datos.H * 0.69, datos.H * 0.86]:
        ax.plot([x_front, x_back - 0.05], [y, y], linewidth=1.2, linestyle="--")

    ax.text(
        datos.B * 0.18,
        datos.H * 0.72,
        "Pantalla / fuste",
        fontsize=10,
        ha="center",
        va="center"
    )
    ax.text(
        datos.B * 0.18,
        datos.H * 0.62,
        f"Vertical: Ø{resultado_fuste['diametro_vertical_mm']:.0f} @ {resultado_fuste['separacion_vertical_cm']:.1f} cm",
        fontsize=8,
        ha="center",
        va="center"
    )
    ax.text(
        datos.B * 0.18,
        datos.H * 0.54,
        f"Horizontal: Ø{resultado_fuste['diametro_horizontal_mm']:.0f} @ {resultado_fuste['separacion_horizontal_cm']:.1f} cm",
        fontsize=8,
        ha="center",
        va="center"
    )

    # Flecha indicadora hacia el fuste.
    flecha(ax, (datos.B * 0.28, datos.H * 0.60), (x_back - 0.02, datos.H * 0.60), "", dx=0.0, dy=0.0, tamano=8)

    # -----------------------------
    # 2) ARMADO DE LA ZAPATA
    # -----------------------------
    y_inf = -datos.hz + 0.08
    y_sup = -0.08

    # Puntera: acero inferior.
    ax.plot([0.10, max(datos.puntera - 0.08, 0.10)], [y_inf, y_inf], linewidth=2.6)
    ax.text(
        datos.puntera * 0.55,
        y_inf - 0.18,
        f"Puntera inf.: Ø{resultado_zapata['diametro_puntera_mm']:.0f} @ {resultado_zapata['sep_puntera_cm']:.1f} cm",
        fontsize=8,
        ha="center",
        va="top"
    )

    # Talón: acero superior.
    x_ini_talon = datos.puntera + datos.t_base + 0.08
    x_fin_talon = datos.B - 0.10
    ax.plot([x_ini_talon, x_fin_talon], [y_sup, y_sup], linewidth=2.6)
    ax.text(
        (x_ini_talon + x_fin_talon) / 2,
        y_sup + 0.16,
        f"Talón sup.: Ø{resultado_zapata['diametro_talon_mm']:.0f} @ {resultado_zapata['sep_talon_cm']:.1f} cm",
        fontsize=8,
        ha="center",
        va="bottom"
    )

    # -----------------------------
    # 3) DENTELLÓN EN EL CORTE DEL MURO
    # -----------------------------
    if datos.usar_llave and "llave" in geometria:
        x1 = datos.pos_llave - datos.ancho_llave / 2
        x2 = datos.pos_llave + datos.ancho_llave / 2
        y1 = -datos.hz
        y2 = -datos.hz - datos.profundidad_llave

        # Barras longitudinales esquemáticas dentro del dentellón.
        xb1 = x1 + datos.ancho_llave * 0.30
        xb2 = x1 + datos.ancho_llave * 0.70
        ax.plot([xb1, xb1], [y2 + 0.05, y1 - 0.05], linewidth=2.0)
        ax.plot([xb2, xb2], [y2 + 0.05, y1 - 0.05], linewidth=2.0)

        # Estribos esquemáticos.
        n_est = max(2, int(datos.profundidad_llave / 0.12))
        for i in range(n_est):
            yy = y2 + 0.07 + i * (max(datos.profundidad_llave - 0.14, 0.05) / max(n_est - 1, 1))
            ax.plot([x1 + 0.06, x2 - 0.06], [yy, yy], linewidth=1.0)

        ax.text(
            (x1 + x2) / 2,
            y2 - 0.18,
            (f"Dentellón: {resultado_dentellon['n_barras_long_dentellon']}Ø{resultado_dentellon['diametro_llave_mm']:.0f} long." if resultado_dentellon.get("requiere_detalle_viga", False) else "Dentellón pequeño: sin armado independiente"),
            fontsize=8,
            ha="center",
            va="top"
        )
        ax.text(
            (x1 + x2) / 2,
            y2 - 0.30,
            (f"Estribos Ø{resultado_dentellon['diametro_estribo_mm']:.0f} @ {resultado_dentellon['sep_estribo_dentellon_cm']:.1f} cm" if resultado_dentellon.get("requiere_detalle_viga", False) else "Monolítico con zapata"),
            fontsize=8,
            ha="center",
            va="top"
        )

        # -----------------------------
        # 4) DETALLE AMPLIADO DEL DENTELLÓN
        # -----------------------------
        # Ubicación del detalle ampliado a la derecha.
        det_x0 = datos.B + 0.70
        det_y0 = max(datos.H * 0.25, 0.90)
        det_b = max(datos.ancho_llave * 1.5, 0.80)
        det_h = max(datos.profundidad_llave * 2.0, 1.10)

        # Marco del dentellón ampliado.
        rect = [
            (det_x0, det_y0),
            (det_x0 + det_b, det_y0),
            (det_x0 + det_b, det_y0 - det_h),
            (det_x0, det_y0 - det_h),
        ]
        ax.add_patch(Polygon(rect, closed=True, fill=False, linewidth=1.8))

        # Estribo interior ampliado.
        rec = 0.10
        estribo = [
            (det_x0 + rec, det_y0 - rec),
            (det_x0 + det_b - rec, det_y0 - rec),
            (det_x0 + det_b - rec, det_y0 - det_h + rec),
            (det_x0 + rec, det_y0 - det_h + rec),
        ]
        ax.add_patch(Polygon(estribo, closed=True, fill=False, linewidth=1.3, linestyle="--"))

        # Barras longitudinales (sección transversal del dentellón).
        xs = [det_x0 + det_b * 0.32, det_x0 + det_b * 0.68]
        ys = [det_y0 - det_h * 0.25, det_y0 - det_h * 0.75]
        for xx in xs:
            for yy in ys:
                ax.scatter([xx], [yy], s=22)

        ax.text(
            det_x0 + det_b / 2,
            det_y0 + 0.12,
            "Detalle ampliado del dentellón",
            fontsize=9,
            ha="center",
            va="bottom"
        )
        ax.text(
            det_x0 + det_b / 2,
            det_y0 - det_h - 0.10,
            f"Longitudinales Ø{resultado_dentellon['diametro_llave_mm']:.0f}",
            fontsize=8,
            ha="center",
            va="top"
        )
        ax.text(
            det_x0 + det_b / 2,
            det_y0 - det_h - 0.24,
            "Estribos cerrados",
            fontsize=8,
            ha="center",
            va="top"
        )

        # Flecha del dentellón real al detalle ampliado.
        flecha(
            ax,
            ((x1 + x2) / 2 + 0.10, y2 + datos.profundidad_llave * 0.40),
            (det_x0 - 0.10, det_y0 - det_h * 0.45),
            "",
            dx=0.0,
            dy=0.0,
            tamano=8
        )

    ax.set_title("Detalle general de armado del muro de retención")



def tabla_resumen_armado_dentellon(resultado: dict) -> pd.DataFrame:
    """
    Tabla corta de armado del dentellón.
    """
    lateral = resultado.get("refuerzo_lateral_dentellon", {})
    if resultado.get("requiere_detalle_viga", False):
        cd = resultado["cortante_dentellon"]
        s_calc = cd["s_calc_cm"] if math.isfinite(cd["s_calc_cm"]) else "No requiere Vs"
        filas = [
            ("Criterio", resultado["criterio_armado_dentellon"], "-"),
            ("As longitudinal mínimo ACI", resultado["As_min_dentellon_cm2_m"], "cm²/m"),
            ("As longitudinal provisto", resultado["As_long_dentellon_prov_cm2_m"], "cm²/m"),
            ("Longitudinales mínimos", f"{resultado['n_barras_long_dentellon']}Ø{resultado['diametro_llave_mm']:.0f}", "barras"),
            ("Estribos adoptados", f"Ø{resultado['diametro_estribo_mm']:.0f} @ {resultado['sep_estribo_dentellon_cm']:.1f}", "mm @ cm"),
            ("Estribo llega hasta top de zapata", "Sí", "-"),
            ("Longitudinales llegan hasta top de zapata", "Sí", "-"),
            ("Vu", resultado["Vu_dentellon_ton_m"], "ton/m"),
            ("φVc", cd["phi_Vc_ton_m"], "ton/m"),
            ("Vs requerido", cd["Vs_req_ton_m"], "ton/m"),
            ("Av/s requerido", cd["Av_s_req_cm2_cm"], "cm²/cm"),
            ("s calculado", s_calc, "cm"),
        ]
        if lateral.get("requiere", False):
            filas.append(("Nota acero lateral", lateral.get("nota", ""), "-"))
        filas.append(("Estado", cd["estado"], "-"))
    else:
        filas = [
            ("Criterio", resultado["criterio_armado_dentellon"], "-"),
            ("Requiere armado independiente", "No", "-"),
            ("Estado", resultado["estado_armado_dentellon"], "-"),
            ("Nota", "Se considera monolítico con la zapata.", "-"),
        ]
    return pd.DataFrame(filas, columns=["Verificación", "Valor", "Unidad"])



def _preparar_eje_detalle(ax, titulo: str, xlim: tuple[float, float], ylim: tuple[float, float]):
    """
    Configura un eje de dibujo para un detalle de armado.

    Se usa en los esquemas didácticos de la última pestaña para mantener una
    presentación limpia, con grilla suave y aspecto geométrico adecuado.
    """
    ax.set_title(titulo)
    ax.set_xlim(*xlim)
    ax.set_ylim(*ylim)
    ax.set_aspect("equal", adjustable="box")
    ax.grid(True, linestyle=":", linewidth=0.7)
    ax.set_xlabel("x [m]")
    ax.set_ylabel("y [m]")


def dibujar_detalle_pantalla_frontal(ax, datos: DatosMuro, resultado_fuste: dict):
    """
    Dibuja la vista frontal de la pantalla del muro.

    Se representa una franja didáctica de 1.00 m de ancho del muro para mostrar:
    - barras verticales principales;
    - barras horizontales de distribución;
    - texto del armado adoptado.
    """
    ancho_panel = 1.00
    ax.add_patch(Polygon([(0, 0), (ancho_panel, 0), (ancho_panel, datos.H), (0, datos.H)],
                         closed=True, fill=False, linewidth=2.0))

    # Barras verticales
    for x in [0.12, 0.32, 0.52, 0.72, 0.88]:
        ax.plot([x, x], [0.08, datos.H - 0.08], linewidth=2.0)

    # Barras horizontales
    n_h = 7
    for i in range(1, n_h + 1):
        y = datos.H * i / (n_h + 1)
        ax.plot([0.06, ancho_panel - 0.06], [y, y], linewidth=1.0, linestyle="--")

    ax.text(
        ancho_panel / 2,
        -0.22,
        f"Vertical principal: Ø{resultado_fuste['diametro_vertical_mm']:.0f} @ {resultado_fuste['separacion_vertical_cm']:.1f} cm",
        ha="center", va="top", fontsize=8
    )
    ax.text(
        ancho_panel / 2,
        -0.40,
        f"Horizontal distribución: Ø{resultado_fuste['diametro_horizontal_mm']:.0f} @ {resultado_fuste['separacion_horizontal_cm']:.1f} cm",
        ha="center", va="top", fontsize=8
    )

    _preparar_eje_detalle(ax, "Pantalla - vista frontal", (-0.10, 1.10), (-0.55, max(datos.H + 0.20, 1.0)))



def dibujar_detalle_pantalla_corte(ax, datos: DatosMuro, geometria: dict, resultado_fuste: dict):
    """
    Dibuja el corte transversal de la pantalla del muro con acero en ambas caras.
    """
    dibujar_poligono(ax, geometria["fuste"], "Pantalla")

    x_relleno = geometria["x_dorso_fuste_base"] - 0.06
    x_frontal = geometria["x_frente_fuste"] + 0.06

    ax.plot([x_relleno, x_relleno], [0.10, datos.H - 0.10], linewidth=2.2)
    for y in [datos.H * 0.20, datos.H * 0.42, datos.H * 0.64, datos.H * 0.86]:
        ax.scatter([x_relleno], [y], s=22)

    ax.plot([x_frontal, x_frontal], [0.10, datos.H - 0.10], linewidth=2.2)
    for y in [datos.H * 0.20, datos.H * 0.42, datos.H * 0.64, datos.H * 0.86]:
        ax.scatter([x_frontal], [y], s=22)

    for y in [datos.H * 0.25, datos.H * 0.45, datos.H * 0.65, datos.H * 0.85]:
        ax.plot([x_frontal, x_relleno], [y, y], linewidth=1.1, linestyle="--")

    ax.text(
        x_relleno + 0.35,
        datos.H * 0.60,
        f"Cara posterior / relleno\nØ{resultado_fuste['diametro_vertical_mm']:.0f} @ {resultado_fuste['separacion_vertical_cm']:.1f} cm",
        fontsize=8, ha="left", va="center"
    )
    ax.text(
        x_frontal - 0.35,
        datos.H * 0.36,
        f"Cara frontal\nØ{resultado_fuste['diametro_vertical_mm']:.0f} @ {resultado_fuste['separacion_vertical_cm']:.1f} cm",
        fontsize=8, ha="right", va="center"
    )
    ax.text(
        (x_frontal + x_relleno) / 2,
        datos.H * 0.12,
        f"Horizontal/distribución\nØ{resultado_fuste['diametro_horizontal_mm']:.0f} @ {resultado_fuste['separacion_horizontal_cm']:.1f} cm",
        fontsize=8, ha="center", va="center"
    )

    _preparar_eje_detalle(
        ax,
        "Pantalla - corte",
        (datos.puntera - 0.45, datos.puntera + datos.t_base + 1.25),
        (-0.20, max(datos.H + 0.30, 1.0))
    )


def dibujar_detalle_zapata_planta(ax, datos: DatosMuro, resultado_zapata: dict):
    """
    Dibuja una vista superior didáctica de la zapata.

    Se representa una franja longitudinal del muro donde se identifican la
    puntera y el talón, además de la dirección del acero principal.
    """
    largo = 3.00
    ax.add_patch(Polygon([(0, 0), (datos.B, 0), (datos.B, largo), (0, largo)],
                         closed=True, fill=False, linewidth=2.0))
    x_fuste1 = datos.puntera
    x_fuste2 = datos.puntera + datos.t_base
    ax.add_patch(Polygon([(x_fuste1, 0.85), (x_fuste2, 0.85), (x_fuste2, 2.15), (x_fuste1, 2.15)],
                         closed=True, fill=False, linewidth=1.8))

    # Barras puntera (sentido longitudinal)
    for y in [0.35, 0.75, 1.15, 1.55, 1.95, 2.35, 2.75]:
        ax.plot([0.15, max(datos.puntera - 0.10, 0.15)], [y, y], linewidth=1.6)

    # Barras talón (sentido longitudinal)
    x_t0 = datos.puntera + datos.t_base + 0.10
    x_t1 = datos.B - 0.15
    for y in [0.35, 0.75, 1.15, 1.55, 1.95, 2.35, 2.75]:
        ax.plot([x_t0, x_t1], [y, y], linewidth=1.6)

    ax.text(datos.puntera / 2, 3.12,
            f"Puntera: Ø{resultado_zapata['diametro_puntera_mm']:.0f} @ {resultado_zapata['sep_puntera_cm']:.1f} cm",
            ha="center", va="bottom", fontsize=8)
    ax.text((x_t0 + x_t1) / 2, 3.12,
            f"Talón: Ø{resultado_zapata['diametro_talon_mm']:.0f} @ {resultado_zapata['sep_talon_cm']:.1f} cm",
            ha="center", va="bottom", fontsize=8)
    ax.text((x_fuste1 + x_fuste2)/2, 1.50, "Pantalla", ha="center", va="center", fontsize=8)

    _preparar_eje_detalle(ax, "Zapata - vista superior", (-0.20, datos.B + 0.20), (-0.15, 3.45))


def dibujar_detalle_zapata_corte(ax, datos: DatosMuro, geometria: dict, resultado_zapata: dict):
    """
    Dibuja el corte de la zapata con el acero principal de puntera y talón.
    """
    dibujar_poligono(ax, geometria["zapata"], "Zapata")
    dibujar_poligono(ax, geometria["fuste"], "Fuste")

    y_inf = -datos.hz + 0.08
    y_sup = -0.08

    ax.plot([0.08, max(datos.puntera - 0.08, 0.08)], [y_inf, y_inf], linewidth=2.4)
    ax.plot([datos.puntera + datos.t_base + 0.08, datos.B - 0.08], [y_sup, y_sup], linewidth=2.4)

    ax.text(datos.puntera / 2, y_inf - 0.14,
            f"Puntera inf.\nØ{resultado_zapata['diametro_puntera_mm']:.0f} @ {resultado_zapata['sep_puntera_cm']:.1f} cm",
            ha="center", va="top", fontsize=8)
    ax.text((datos.puntera + datos.t_base + datos.B) / 2, y_sup + 0.14,
            f"Talón sup.\nØ{resultado_zapata['diametro_talon_mm']:.0f} @ {resultado_zapata['sep_talon_cm']:.1f} cm",
            ha="center", va="bottom", fontsize=8)

    _preparar_eje_detalle(ax, "Zapata - corte", (-0.20, datos.B + 0.20), (-datos.hz - max(0.60, datos.profundidad_llave + 0.30), max(datos.H * 0.20, 0.60)))


def dibujar_detalle_dentellon_corte(ax, datos: DatosMuro, geometria: dict, resultado_dentellon: dict):
    """
    Dibuja el corte del dentellón.

    Criterios gráficos:
    - acero mínimo longitudinal solo en los extremos de la sección;
    - estribo cerrado hasta el top de la zapata;
    - si h > 0.80 m, se muestra una nota para incluir acero lateral adicional
      por control de agrietamiento, sin dibujarlo como puntos para evitar confusión.
    """
    if not datos.usar_llave or "llave" not in geometria:
        ax.text(0.5, 0.5, "Sin dentellón", ha="center", va="center", fontsize=12)
        ax.set_axis_off()
        return

    x1 = datos.pos_llave - datos.ancho_llave / 2
    x2 = datos.pos_llave + datos.ancho_llave / 2
    y_top = 0.0
    y_foot_bot = -datos.hz
    y_key_bot = -datos.hz - datos.profundidad_llave

    # Zapata local
    ax.add_patch(Polygon([(x1 - 0.60, y_top), (x2 + 0.60, y_top), (x2 + 0.60, y_foot_bot), (x1 - 0.60, y_foot_bot)],
                         closed=True, fill=False, linewidth=2.0))
    # Dentellón
    ax.add_patch(Polygon([(x1, y_foot_bot), (x2, y_foot_bot), (x2, y_key_bot), (x1, y_key_bot)],
                         closed=True, fill=False, linewidth=2.0))

    if resultado_dentellon.get("requiere_detalle_viga", False):
        rec_x = max(datos.ancho_llave * 0.14, 0.04)
        rec_y_top = max(datos.hz * 0.25, 0.06)
        rec_y_bot = max(datos.profundidad_llave * 0.16, 0.06)

        xL = x1 + rec_x
        xR = x2 - rec_x
        y_st_top = y_top - rec_y_top
        y_st_bot = y_key_bot + rec_y_bot

        # Estribo cerrado: sube hasta el top de la zapata.
        ax.add_patch(Polygon(
            [(xL, y_st_top), (xR, y_st_top), (xR, y_st_bot), (xL, y_st_bot)],
            closed=True, fill=False, linewidth=1.25, linestyle="--"
        ))

        # Acero mínimo longitudinal solo en extremos: 2 arriba dentro de zapata y 2 abajo.
        x_bar_1 = x1 + datos.ancho_llave * 0.24
        x_bar_2 = x2 - datos.ancho_llave * 0.24
        y_bar_top = y_top - max(datos.hz * 0.28, 0.07)
        y_bar_bot = y_key_bot + max(datos.profundidad_llave * 0.18, 0.06)

        for xx, yy in [(x_bar_1, y_bar_top), (x_bar_2, y_bar_top), (x_bar_1, y_bar_bot), (x_bar_2, y_bar_bot)]:
            ax.scatter([xx], [yy], s=30)

        txt1 = f"Longitudinales mín.: {resultado_dentellon['n_barras_long_dentellon']}Ø{resultado_dentellon['diametro_llave_mm']:.0f} en extremos"
        txt2 = f"Estribos: Ø{resultado_dentellon['diametro_estribo_mm']:.0f} @ {resultado_dentellon['sep_estribo_dentellon_cm']:.1f} cm hasta top zapata"

        lateral = resultado_dentellon.get("refuerzo_lateral_dentellon", {})
        txt3 = lateral.get("nota", "") if lateral.get("requiere", False) else ""
    else:
        txt1 = "Dentellón pequeño"
        txt2 = "Monolítico con zapata"
        txt3 = ""

    ax.text((x1 + x2) / 2, y_key_bot - 0.12, txt1, ha="center", va="top", fontsize=8)
    ax.text((x1 + x2) / 2, y_key_bot - 0.28, txt2, ha="center", va="top", fontsize=8)
    if txt3:
        ax.text((x1 + x2) / 2, y_key_bot - 0.44, txt3, ha="center", va="top", fontsize=7.5)

    _preparar_eje_detalle(ax, "Dentellón - corte", (x1 - 0.80, x2 + 0.80), (y_key_bot - 0.60, 0.30))


# Lista explícita de nombres que app.py puede importar desde este módulo.
__all__ = [
    "DatosMuro",
    "validar_datos_muro",
    "generar_puntos_muro",
    "dibujar_muro",
    "dibujar_diagrama_fuerzas",
    "calcular_trial_wedge_activo",
    "dibujar_trial_wedge_critico",
    "graficar_curva_trial_wedge",
    "calcular_verificacion_caltrans_pdf",
    "tabla_comparacion_pdf",
    "tabla_fuerzas_pdf",
    "tabla_insumos_partes_pdf",
    "calcular_estabilidad_externa_desde_caso",
    "crear_caso_pdf_caltrans",
    "generar_memoria_word",
    "calcular_diseno_fuste_dinamico",
    "tabla_diseno_fuste",
    "dibujar_armado_fuste",
    "calcular_presiones_contacto_servicio",
    "calcular_diseno_zapata_dinamico",
    "tabla_presiones_contacto",
    "tabla_diseno_zapata",
    "dibujar_presiones_contacto",
    "dibujar_armado_zapata",
    "calcular_diseno_zapata_definitivo",
    "tabla_diseno_zapata_definitivo",
    "dibujar_detalle_zapata_definitivo",
    "verificar_cortante_rectangular",
    "longitud_desarrollo_basica_cm",
    "calcular_deslizamiento_y_llave",
    "tabla_deslizamiento_llave",
    "dibujar_deslizamiento_pasivo_llave",
    "coeficiente_pasivo_rankine",
    "dibujar_detalle_general_armado",
    "resumen_geometria",
    "convertir_resistencias_a_sistema_interno",
    "recomendar_posicion_dentellon",
    "tabla_momentos_estabilidad",
    "diseno_cortante_dentellon_aci",
    "as_min_longitudinal_aci_cm2",
    "calcular_refuerzo_lateral_dentellon",
    "dibujar_detalle_dentellon_corte",
    "dibujar_detalle_zapata_corte",
    "dibujar_detalle_zapata_planta",
    "dibujar_detalle_pantalla_corte",
    "dibujar_detalle_pantalla_frontal",
    "tabla_resumen_armado_dentellon",
]


# =============================================================================
# OVERRIDES DE REVISIÓN GENERAL
# Estas funciones corrigen el manejo de altura de relleno, puntera cero,
# diseño de fuste/zapata y reportes solicitados por el usuario.
# =============================================================================

def calcular_y_terreno(datos: DatosMuro, geometria: dict, x: float) -> float:
    """
    Elevación de la superficie del relleno.

    La línea del relleno parte desde `altura_relleno`, no desde la corona del muro.
    Así, cambiar altura_relleno modifica el dibujo y todos los cálculos que usan
    Trial Wedge, pesos de suelo y momentos.
    """
    x0 = geometria["x_dorso_fuste_corona"]
    y0 = datos.altura_relleno
    if datos.pendiente_v == 0:
        return y0
    pendiente = datos.pendiente_v / datos.pendiente_h
    return y0 + pendiente * max(x - x0, 0.0)


def calcular_linea_relleno(datos: DatosMuro, geometria: dict) -> list[tuple[float, float]]:
    """
    Línea del relleno para el dibujo. Parte desde la altura real ingresada.
    """
    x_inicio = geometria["x_dorso_fuste_corona"]
    y_inicio = datos.altura_relleno
    x_fin = datos.B + max(datos.B * 0.60, 1.00)
    y_fin = calcular_y_terreno(datos, geometria, x_fin)
    return [(x_inicio, y_inicio), (x_fin, y_fin)]


def altura_relleno_en_muro(datos: DatosMuro, geometria: dict) -> float:
    """
    Altura de suelo retenido directamente contra el fuste.
    Para el diseño local del fuste se limita al alto del fuste.
    """
    return max(0.0, min(datos.altura_relleno, datos.H))


def suelo_sobre_talon(datos: DatosMuro, geometria: dict) -> dict:
    """
    Peso de suelo sobre el talón usando la línea real del terreno.
    """
    talon = calcular_talon(datos)
    if talon <= 0:
        return {"area_m2": 0.0, "W_ton_m": 0.0, "x_centroide_m": datos.B, "h_inicio_m": 0.0, "h_fin_m": 0.0}

    x1 = datos.puntera + datos.t_base
    x2 = datos.B
    h1 = max(calcular_y_terreno(datos, geometria, x1), 0.0)
    h2 = max(calcular_y_terreno(datos, geometria, x2), 0.0)

    area = talon * (h1 + h2) / 2.0
    W = area * datos.gamma_suelo

    if h1 + h2 > 0:
        x_rel = talon * (h1 + 2.0 * h2) / (3.0 * (h1 + h2))
        x_c = x1 + x_rel
    else:
        x_c = x1 + talon / 2.0

    return {"area_m2": area, "W_ton_m": W, "x_centroide_m": x_c, "h_inicio_m": h1, "h_fin_m": h2}


def resolver_as_flexion_rectangular(Mu_ton_m: float, b_cm: float, d_cm: float, fc_kg_cm2: float, fy_kg_cm2: float, phi_flexion: float = 0.90) -> float:
    """
    Calcula As por flexión. Si la sección no puede resistir el momento con acero
    simple razonable, devuelve inf en vez de nan para evitar salidas confusas.
    """
    if Mu_ton_m <= 0:
        return 0.0
    if b_cm <= 0 or d_cm <= 0 or fc_kg_cm2 <= 0 or fy_kg_cm2 <= 0:
        return float("inf")

    Mu_kg_cm = Mu_ton_m * 100000.0
    A = phi_flexion * fy_kg_cm2 ** 2 / (2.0 * 0.85 * fc_kg_cm2 * b_cm)
    Bq = -phi_flexion * fy_kg_cm2 * d_cm
    C = Mu_kg_cm
    disc = Bq ** 2 - 4.0 * A * C
    if disc < 0:
        return float("inf")
    r1 = (-Bq - math.sqrt(disc)) / (2.0 * A)
    r2 = (-Bq + math.sqrt(disc)) / (2.0 * A)
    candidatos = [r for r in [r1, r2] if r > 0 and math.isfinite(r)]
    return min(candidatos) if candidatos else float("inf")


def seleccionar_separacion(area_barra: float, As_req_cm2_m: float, separacion_max_cm: float = 30.0) -> tuple[float, float]:
    """
    Selecciona separación comercial. Si As_req es infinito, devuelve inf y obliga revisión.
    """
    if As_req_cm2_m <= 0:
        return float("nan"), 0.0
    if not math.isfinite(As_req_cm2_m):
        return float("inf"), 0.0

    separacion_teorica = area_barra * 100.0 / As_req_cm2_m
    separaciones_comerciales = [40, 35, 30, 25, 20, 18, 15, 12.5, 10, 7.5]
    separaciones_validas = [s for s in separaciones_comerciales if s <= separacion_teorica and s <= separacion_max_cm]
    if not separaciones_validas:
        separacion = min(separaciones_comerciales)
    else:
        separacion = max(separaciones_validas)
    return separacion, area_barra * 100.0 / separacion


def _estado_acero(As_req: float, As_prov: float) -> str:
    if not math.isfinite(As_req):
        return "No cumple: aumentar sección/refuerzo"
    if As_req <= 0:
        return "No aplica"
    return "OK" if As_prov >= As_req else "No cumple"


def calcular_diseno_fuste_dinamico(
    datos: DatosMuro,
    numero_cunas: int = 180,
    recubrimiento_cm: float = 7.5,
    diametro_vertical_mm: float = 16.0,
    diametro_horizontal_mm: float = 12.0,
    separacion_max_cm: float = 30.0
) -> dict:
    """
    Diseño dinámico del fuste corregido.

    La altura de relleno ingresada controla PA, Mu y Vu. Si altura_relleno < H,
    el empuje local del fuste no se aplica con brazo H/3 sino con h_activa/3.
    """
    geometria = generar_puntos_muro(datos)
    tabla_trial, resultado_trial = calcular_trial_wedge_activo(datos, geometria, numero_cunas=numero_cunas)

    PA = resultado_trial["PA_ton_m"]
    delta_rad = math.radians(resultado_trial["delta_grados"])
    PA_h = PA * math.cos(delta_rad)

    h_activa = altura_relleno_en_muro(datos, geometria)
    brazo = h_activa / 3.0 if h_activa > 0 else 0.0

    gamma_p = 1.50
    Mu_ton_m = gamma_p * PA_h * brazo
    Vu_ton = gamma_p * PA_h

    b_cm = 100.0
    h_cm = datos.t_base * 100.0
    d_cm = max(h_cm - recubrimiento_cm - (diametro_vertical_mm / 10.0) / 2.0, 1.0)

    As_flexion = resolver_as_flexion_rectangular(Mu_ton_m, b_cm, d_cm, datos.fc, datos.fy)
    As_temp_total = 0.0018 * b_cm * h_cm
    As_temp_cara = As_temp_total / 2.0

    if math.isfinite(As_flexion):
        As_vertical_req = max(As_flexion, As_temp_cara)
    else:
        As_vertical_req = float("inf")

    sep_vertical_cm, As_vertical_prov = seleccionar_separacion(
        area_barra_cm2(diametro_vertical_mm), As_vertical_req, separacion_max_cm
    )

    As_horizontal_req = As_temp_cara
    sep_horizontal_cm, As_horizontal_prov = seleccionar_separacion(
        area_barra_cm2(diametro_horizontal_mm), As_horizontal_req, separacion_max_cm
    )

    # Cortante del fuste: si el espesor aumenta, d aumenta y φVc cambia.
    cort = verificar_cortante_rectangular(Vu_ton, b_cm, d_cm, datos.fc)
    estado_cortante = cort["estado"]

    q_base_ton_m2 = PA_h / h_activa if h_activa > 0 else 0.0
    estado_flexion = _estado_acero(As_vertical_req, As_vertical_prov)

    return {
        "PA_ton_m": PA,
        "delta_grados": resultado_trial["delta_grados"],
        "alfa_critico_grados": resultado_trial.get("alfa_grados", resultado_trial.get("alfa_critico_grados", 0.0)),
        "q_base_ton_m2": q_base_ton_m2,
        "altura_activa_relleno_m": h_activa,
        "brazo_activo_m": brazo,
        "Mu_ton_m_m": Mu_ton_m,
        "Vu_ton_m": Vu_ton,
        "b_cm": b_cm,
        "h_cm": h_cm,
        "d_cm": d_cm,
        "As_flexion_cm2_m": As_flexion,
        "As_temp_total_cm2_m": As_temp_total,
        "As_temp_cara_cm2_m": As_temp_cara,
        "As_vertical_req_cm2_m": As_vertical_req,
        "diametro_vertical_mm": diametro_vertical_mm,
        "separacion_vertical_cm": sep_vertical_cm,
        "As_vertical_prov_cm2_m": As_vertical_prov,
        "diametro_horizontal_mm": diametro_horizontal_mm,
        "As_horizontal_req_cm2_m": As_horizontal_req,
        "separacion_horizontal_cm": sep_horizontal_cm,
        "As_horizontal_prov_cm2_m": As_horizontal_prov,
        "phi_Vc_ton_m": cort["phi_Vc_ton_m"],
        "relacion_cortante": cort["relacion"],
        "estado_cortante": estado_cortante,
        "estado_flexion": estado_flexion,
        "tabla_trial": tabla_trial,
        "resultado_trial": resultado_trial,
    }


def _formato_armado(diametro: float, sep: float) -> str:
    if not math.isfinite(sep):
        return "No cumple"
    if sep <= 0 or math.isnan(sep):
        return "No aplica"
    return f"Ø{diametro:.0f} @ {sep:.1f} cm"


def calcular_presiones_contacto_servicio(datos: DatosMuro, numero_cunas: int = 180) -> dict:
    """
    Equilibrio externo corregido.

    Se separan los pesos de puntera, zona bajo fuste y talón para que la tabla de
    momentos muestre explícitamente la contribución de la puntera.
    """
    geometria = generar_puntos_muro(datos)
    tabla_trial, resultado_trial = calcular_trial_wedge_activo(datos, geometria, numero_cunas=numero_cunas)

    PA = resultado_trial["PA_ton_m"]
    delta = math.radians(resultado_trial["delta_grados"])
    PA_v = PA * math.sin(delta)
    PA_h = PA * math.cos(delta)

    # Pesos de zapata separados.
    L_puntera = max(datos.puntera, 0.0)
    L_bajo_fuste = max(datos.t_base, 0.0)
    L_talon = max(calcular_talon(datos), 0.0)

    W_zapata_puntera = L_puntera * datos.hz * datos.gamma_hormigon
    x_zapata_puntera = L_puntera / 2.0 if L_puntera > 0 else 0.0

    W_zapata_fuste = L_bajo_fuste * datos.hz * datos.gamma_hormigon
    x_zapata_fuste = datos.puntera + L_bajo_fuste / 2.0

    W_zapata_talon = L_talon * datos.hz * datos.gamma_hormigon
    x_zapata_talon = datos.puntera + datos.t_base + L_talon / 2.0 if L_talon > 0 else datos.B

    W_zapata = W_zapata_puntera + W_zapata_fuste + W_zapata_talon
    x_zapata = (
        W_zapata_puntera * x_zapata_puntera
        + W_zapata_fuste * x_zapata_fuste
        + W_zapata_talon * x_zapata_talon
    ) / W_zapata if W_zapata > 0 else 0.0

    area_fuste = (datos.t_base + datos.t_corona) / 2.0 * datos.H
    W_fuste = area_fuste * datos.gamma_hormigon
    x_fuste = datos.puntera + (datos.t_base + datos.t_corona) / 4.0

    if datos.usar_llave:
        W_dentellon = datos.ancho_llave * datos.profundidad_llave * datos.gamma_hormigon
        x_dentellon = datos.pos_llave
    else:
        W_dentellon = 0.0
        x_dentellon = 0.0

    suelo_talon = suelo_sobre_talon(datos, geometria)
    W_suelo_talon = suelo_talon["W_ton_m"]
    x_suelo_talon = suelo_talon["x_centroide_m"]

    W_pendiente = 0.0
    x_pendiente = x_suelo_talon

    M_est_zapata_puntera = W_zapata_puntera * x_zapata_puntera
    M_est_zapata_fuste = W_zapata_fuste * x_zapata_fuste
    M_est_zapata_talon = W_zapata_talon * x_zapata_talon
    M_est_zapata = M_est_zapata_puntera + M_est_zapata_fuste + M_est_zapata_talon
    M_est_fuste = W_fuste * x_fuste
    M_est_dentellon = W_dentellon * x_dentellon
    M_est_suelo_talon = W_suelo_talon * x_suelo_talon
    M_est_pendiente = 0.0
    M_est_PA_v = PA_v * datos.B

    V = W_zapata + W_fuste + W_dentellon + W_suelo_talon + W_pendiente + PA_v
    M_est = M_est_zapata + M_est_fuste + M_est_dentellon + M_est_suelo_talon + M_est_PA_v

    h_activa = altura_relleno_en_muro(datos, geometria)
    brazo_PA_h = h_activa / 3.0 if h_activa > 0 else 0.0
    M_volc_PA_h = PA_h * brazo_PA_h
    M_volc = M_volc_PA_h

    M_resultante = M_est - M_volc
    x_resultante = M_resultante / V if V > 0 else float("nan")
    e = datos.B / 2.0 - x_resultante

    if abs(e) <= datos.B / 6.0:
        q_prom = V / datos.B
        q_max = q_prom * (1.0 + 6.0 * e / datos.B)
        q_min = q_prom * (1.0 - 6.0 * e / datos.B)
    else:
        B_efectivo = 3.0 * (datos.B / 2.0 - abs(e))
        q_max = 2.0 * V / B_efectivo if B_efectivo > 0 else float("inf")
        q_min = 0.0

    estado_q = "OK" if math.isfinite(q_max) and q_max <= datos.qa else "No cumple"

    return {
        "PA_ton_m": PA,
        "PA_h_ton_m": PA_h,
        "PA_v_ton_m": PA_v,
        "delta_grados": resultado_trial["delta_grados"],
        "V_total_ton_m": V,
        "M_est_ton_m_m": M_est,
        "M_volc_ton_m_m": M_volc,
        "M_resultante_ton_m_m": M_resultante,
        "x_resultante_m": x_resultante,
        "e_m": e,
        "qmax_ton_m2": q_max,
        "qmin_ton_m2": q_min,
        "q_adm_ton_m2": datos.qa,
        "estado_q": estado_q,

        "W_zapata_ton_m": W_zapata,
        "x_zapata_m": x_zapata,
        "M_est_zapata_ton_m_m": M_est_zapata,
        "W_zapata_puntera_ton_m": W_zapata_puntera,
        "x_zapata_puntera_m": x_zapata_puntera,
        "M_est_zapata_puntera_ton_m_m": M_est_zapata_puntera,
        "W_zapata_fuste_ton_m": W_zapata_fuste,
        "x_zapata_fuste_m": x_zapata_fuste,
        "M_est_zapata_fuste_ton_m_m": M_est_zapata_fuste,
        "W_zapata_talon_ton_m": W_zapata_talon,
        "x_zapata_talon_m": x_zapata_talon,
        "M_est_zapata_talon_ton_m_m": M_est_zapata_talon,

        "W_fuste_ton_m": W_fuste,
        "x_fuste_m": x_fuste,
        "M_est_fuste_ton_m_m": M_est_fuste,
        "W_dentellon_ton_m": W_dentellon,
        "x_dentellon_m": x_dentellon,
        "M_est_dentellon_ton_m_m": M_est_dentellon,
        "W_suelo_talon_ton_m": W_suelo_talon,
        "x_suelo_talon_m": x_suelo_talon,
        "M_est_suelo_talon_ton_m_m": M_est_suelo_talon,
        "area_suelo_talon_m2": suelo_talon["area_m2"],
        "h_suelo_talon_inicio_m": suelo_talon["h_inicio_m"],
        "h_suelo_talon_fin_m": suelo_talon["h_fin_m"],
        "W_pendiente_ton_m": W_pendiente,
        "x_pendiente_m": x_pendiente,
        "M_est_pendiente_ton_m_m": M_est_pendiente,
        "M_est_PA_v_ton_m_m": M_est_PA_v,
        "brazo_PA_h_m": brazo_PA_h,
        "M_volc_PA_h_ton_m_m": M_volc_PA_h,
    }


def calcular_diseno_zapata_dinamico(
    datos: DatosMuro,
    numero_cunas: int = 180,
    recubrimiento_cm: float = 7.5,
    diametro_puntera_mm: float = 16.0,
    diametro_talon_mm: float = 16.0,
    separacion_max_cm: float = 30.0
) -> dict:
    """
    Diseño preliminar de zapata corregido.

    El talón se diseña con carga neta firmada: peso de suelo hacia abajo menos
    reacción de suelo hacia arriba. Para flexión se usa el valor absoluto del
    momento para no reportar Mu=0 cuando realmente hay inversión de signo.
    """
    presiones = calcular_presiones_contacto_servicio(datos, numero_cunas=numero_cunas)
    gamma_u = 1.50

    b_cm = 100.0
    h_cm = datos.hz * 100.0

    # Puntera
    Lp = max(datos.puntera, 0.0)
    q0 = presion_lineal_en_x(datos, presiones, 0.0)
    q1 = presion_lineal_en_x(datos, presiones, datos.puntera)
    q_prom_p = (q0 + q1) / 2.0 if Lp > 0 else 0.0
    Mu_puntera = gamma_u * q_prom_p * Lp ** 2 / 2.0
    Vu_puntera = gamma_u * q_prom_p * Lp

    # Talón
    Lt = max(calcular_talon(datos), 0.0)
    x_t0 = datos.puntera + datos.t_base
    x_t1 = datos.B
    q_t0 = presion_lineal_en_x(datos, presiones, x_t0)
    q_t1 = presion_lineal_en_x(datos, presiones, x_t1)
    q_prom_t = (q_t0 + q_t1) / 2.0 if Lt > 0 else 0.0

    geom = generar_puntos_muro(datos)
    soil = suelo_sobre_talon(datos, geom)
    w_suelo_prom = soil["W_ton_m"] / Lt if Lt > 0 else 0.0
    w_neto_talon = w_suelo_prom - q_prom_t

    Mu_talon_firmado = gamma_u * w_neto_talon * Lt ** 2 / 2.0
    Vu_talon_firmado = gamma_u * w_neto_talon * Lt
    Mu_talon = abs(Mu_talon_firmado)
    Vu_talon = abs(Vu_talon_firmado)

    d_puntera_cm = h_cm - recubrimiento_cm - (diametro_puntera_mm / 10.0) / 2.0
    d_talon_cm = h_cm - recubrimiento_cm - (diametro_talon_mm / 10.0) / 2.0

    As_puntera = resolver_as_flexion_rectangular(Mu_puntera, b_cm, d_puntera_cm, datos.fc, datos.fy) if Lp > 0 else 0.0
    As_talon = resolver_as_flexion_rectangular(Mu_talon, b_cm, d_talon_cm, datos.fc, datos.fy) if Lt > 0 else 0.0

    As_temp_zapata = 0.0018 * b_cm * h_cm / 2.0
    As_puntera_req = max(As_puntera, As_temp_zapata) if Lp > 0 else 0.0
    As_talon_req = max(As_talon, As_temp_zapata) if Lt > 0 else 0.0

    sep_puntera_cm, As_puntera_prov = seleccionar_separacion(
        area_barra_cm2(diametro_puntera_mm), As_puntera_req, separacion_max_cm
    )
    sep_talon_cm, As_talon_prov = seleccionar_separacion(
        area_barra_cm2(diametro_talon_mm), As_talon_req, separacion_max_cm
    )

    return {
        "presiones": presiones,
        "Lp_m": Lp,
        "Lt_m": Lt,
        "q_puntera_prom_ton_m2": q_prom_p,
        "q_talon_prom_ton_m2": q_prom_t,
        "w_suelo_talon_ton_m2": w_suelo_prom,
        "w_neto_talon_ton_m2": w_neto_talon,
        "Mu_puntera_ton_m_m": Mu_puntera,
        "Vu_puntera_ton_m": Vu_puntera,
        "Mu_talon_ton_m_m": Mu_talon,
        "Mu_talon_firmado_ton_m_m": Mu_talon_firmado,
        "Vu_talon_ton_m": Vu_talon,
        "Vu_talon_firmado_ton_m": Vu_talon_firmado,
        "As_puntera_req_cm2_m": As_puntera_req,
        "As_puntera_prov_cm2_m": As_puntera_prov,
        "sep_puntera_cm": sep_puntera_cm,
        "As_talon_req_cm2_m": As_talon_req,
        "As_talon_prov_cm2_m": As_talon_prov,
        "sep_talon_cm": sep_talon_cm,
        "diametro_puntera_mm": diametro_puntera_mm,
        "diametro_talon_mm": diametro_talon_mm,
    }


def _cortante_o_no_aplica(Vu: float, b_cm: float, d_cm: float, fc: float, longitud_crit: float, etiqueta: str) -> dict:
    if longitud_crit <= 1e-9:
        return {
            "Vu_ton_m": 0.0,
            "phi_Vc_ton_m": verificar_cortante_rectangular(0.0, b_cm, d_cm, fc)["phi_Vc_ton_m"],
            "relacion": 0.0,
            "estado": f"No aplica ({etiqueta} ≤ d)"
        }
    return verificar_cortante_rectangular(Vu, b_cm, d_cm, fc)


def calcular_diseno_zapata_definitivo(
    datos: DatosMuro,
    numero_cunas: int = 180,
    recubrimiento_cm: float = 7.5,
    diametro_puntera_mm: float = 16.0,
    diametro_talon_mm: float = 16.0,
    separacion_max_cm: float = 30.0
) -> dict:
    """
    Diseño de zapata corregido.

    Se eliminan las verificaciones automáticas de anclaje puntera/talón como
    estado del diseño, porque son detalles de desarrollo que deben resolverse
    con ganchos/continuidad y no deben aparecer como semáforo global.
    """
    res = calcular_diseno_zapata_dinamico(
        datos, numero_cunas, recubrimiento_cm, diametro_puntera_mm, diametro_talon_mm, separacion_max_cm
    )
    presiones = res["presiones"]

    b_cm = 100.0
    h_cm = datos.hz * 100.0
    d_puntera_cm = h_cm - recubrimiento_cm - (diametro_puntera_mm / 10.0) / 2.0
    d_talon_cm = h_cm - recubrimiento_cm - (diametro_talon_mm / 10.0) / 2.0
    d_puntera_m = d_puntera_cm / 100.0
    d_talon_m = d_talon_cm / 100.0
    gamma_u = 1.50

    Lp = max(datos.puntera, 0.0)
    Lp_crit = max(Lp - d_puntera_m, 0.0)
    q0 = presion_lineal_en_x(datos, presiones, 0.0)
    qcrit_p = presion_lineal_en_x(datos, presiones, Lp_crit)
    q_prom_crit_p = (q0 + qcrit_p) / 2.0 if Lp_crit > 0 else 0.0
    Vu_puntera_crit = gamma_u * q_prom_crit_p * Lp_crit
    cortante_puntera = _cortante_o_no_aplica(Vu_puntera_crit, b_cm, d_puntera_cm, datos.fc, Lp_crit, "Lp")

    Lt = max(calcular_talon(datos), 0.0)
    Lt_crit = max(Lt - d_talon_m, 0.0)
    x_inicio_talon_crit = datos.puntera + datos.t_base + d_talon_m
    x_fin_talon = datos.B
    q_tcrit = presion_lineal_en_x(datos, presiones, x_inicio_talon_crit)
    q_tfin = presion_lineal_en_x(datos, presiones, x_fin_talon)
    q_prom_tcrit = (q_tcrit + q_tfin) / 2.0 if Lt_crit > 0 else 0.0

    geom = generar_puntos_muro(datos)
    soil = suelo_sobre_talon(datos, geom)
    w_suelo_prom = soil["W_ton_m"] / Lt if Lt > 0 else 0.0
    w_neto = w_suelo_prom - q_prom_tcrit
    Vu_talon_crit = gamma_u * abs(w_neto) * Lt_crit
    cortante_talon = _cortante_o_no_aplica(Vu_talon_crit, b_cm, d_talon_cm, datos.fc, Lt_crit, "Lt")

    # Estados de acero flexión
    estado_as_puntera = _estado_acero(res["As_puntera_req_cm2_m"], res["As_puntera_prov_cm2_m"]) if Lp > 0 else "No aplica"
    estado_as_talon = _estado_acero(res["As_talon_req_cm2_m"], res["As_talon_prov_cm2_m"]) if Lt > 0 else "No aplica"

    def estado_falla(e):
        return str(e).lower().startswith("no cumple") or "revisar" in str(e).lower()

    estados = [
        presiones["estado_q"],
        cortante_puntera["estado"],
        cortante_talon["estado"],
        estado_as_puntera,
        estado_as_talon,
    ]
    estado_global = "Revisar" if any(estado_falla(e) for e in estados) else "OK"

    res.update({
        "d_puntera_cm": d_puntera_cm,
        "d_talon_cm": d_talon_cm,
        "Lp_crit_m": Lp_crit,
        "Lt_crit_m": Lt_crit,
        "Vu_puntera_crit_ton_m": Vu_puntera_crit,
        "Vu_talon_crit_ton_m": Vu_talon_crit,
        "cortante_puntera": cortante_puntera,
        "cortante_talon": cortante_talon,
        "estado_as_puntera": estado_as_puntera,
        "estado_as_talon": estado_as_talon,
        "ld_puntera_cm": 0.0,
        "ld_talon_cm": 0.0,
        "longitud_disponible_puntera_cm": 0.0,
        "longitud_disponible_talon_cm": 0.0,
        "estado_ld_puntera": "No se evalúa en dashboard",
        "estado_ld_talon": "No se evalúa en dashboard",
        "estado_global_zapata": estado_global,
    })
    return res


def tabla_diseno_zapata_definitivo(resultado: dict) -> pd.DataFrame:
    """
    Tabla de zapata sin cuadros de anclaje automático.
    """
    filas = [
        ("Estado global zapata", resultado["estado_global_zapata"], "-"),
        ("qmax", resultado["presiones"]["qmax_ton_m2"], "ton/m²"),
        ("qmin", resultado["presiones"]["qmin_ton_m2"], "ton/m²"),
        ("qa", resultado["presiones"]["q_adm_ton_m2"], "ton/m²"),
        ("Estado presión admisible", resultado["presiones"]["estado_q"], "-"),

        ("Mu puntera", resultado["Mu_puntera_ton_m_m"], "ton·m/m"),
        ("As puntera requerido", resultado["As_puntera_req_cm2_m"], "cm²/m"),
        ("As puntera provisto", resultado["As_puntera_prov_cm2_m"], "cm²/m"),
        ("Armado puntera", _formato_armado(resultado["diametro_puntera_mm"], resultado["sep_puntera_cm"]), "mm @ cm"),
        ("Estado flexión puntera", resultado.get("estado_as_puntera", "No aplica"), "-"),

        ("Mu talón", resultado["Mu_talon_ton_m_m"], "ton·m/m"),
        ("Mu talón firmado", resultado.get("Mu_talon_firmado_ton_m_m", 0.0), "ton·m/m"),
        ("As talón requerido", resultado["As_talon_req_cm2_m"], "cm²/m"),
        ("As talón provisto", resultado["As_talon_prov_cm2_m"], "cm²/m"),
        ("Armado talón", _formato_armado(resultado["diametro_talon_mm"], resultado["sep_talon_cm"]), "mm @ cm"),
        ("Estado flexión talón", resultado.get("estado_as_talon", "No aplica"), "-"),

        ("Vu puntera crítico", resultado["Vu_puntera_crit_ton_m"], "ton/m"),
        ("φVc puntera", resultado["cortante_puntera"]["phi_Vc_ton_m"], "ton/m"),
        ("Relación Vu/φVc puntera", resultado["cortante_puntera"]["relacion"], "-"),
        ("Estado cortante puntera", resultado["cortante_puntera"]["estado"], "-"),

        ("Vu talón crítico", resultado["Vu_talon_crit_ton_m"], "ton/m"),
        ("φVc talón", resultado["cortante_talon"]["phi_Vc_ton_m"], "ton/m"),
        ("Relación Vu/φVc talón", resultado["cortante_talon"]["relacion"], "-"),
        ("Estado cortante talón", resultado["cortante_talon"]["estado"], "-"),
    ]
    return pd.DataFrame(filas, columns=["Verificación", "Valor", "Unidad"])


def tabla_momentos_estabilidad(presiones: dict) -> pd.DataFrame:
    """
    Tabla explícita de momentos, separando la zapata en puntera, bajo fuste y talón.
    """
    filas = [
        ("ESTABILIZANTES", "", "", ""),
        ("Peso zapata - puntera", presiones.get("W_zapata_puntera_ton_m", 0.0), presiones.get("x_zapata_puntera_m", 0.0), presiones.get("M_est_zapata_puntera_ton_m_m", 0.0)),
        ("Peso zapata - bajo fuste", presiones.get("W_zapata_fuste_ton_m", 0.0), presiones.get("x_zapata_fuste_m", 0.0), presiones.get("M_est_zapata_fuste_ton_m_m", 0.0)),
        ("Peso zapata - talón", presiones.get("W_zapata_talon_ton_m", 0.0), presiones.get("x_zapata_talon_m", 0.0), presiones.get("M_est_zapata_talon_ton_m_m", 0.0)),
        ("Peso zapata total", presiones.get("W_zapata_ton_m", 0.0), presiones.get("x_zapata_m", 0.0), presiones.get("M_est_zapata_ton_m_m", 0.0)),
        ("Peso fuste/pantalla", presiones.get("W_fuste_ton_m", 0.0), presiones.get("x_fuste_m", 0.0), presiones.get("M_est_fuste_ton_m_m", 0.0)),
        ("Peso dentellón", presiones.get("W_dentellon_ton_m", 0.0), presiones.get("x_dentellon_m", 0.0), presiones.get("M_est_dentellon_ton_m_m", 0.0)),
        ("Peso suelo sobre talón", presiones.get("W_suelo_talon_ton_m", 0.0), presiones.get("x_suelo_talon_m", 0.0), presiones.get("M_est_suelo_talon_ton_m_m", 0.0)),
        ("Componente vertical PA", presiones.get("PA_v_ton_m", 0.0), presiones.get("B", ""), presiones.get("M_est_PA_v_ton_m_m", 0.0)),
        ("Total estabilizante", "", "", presiones.get("M_est_ton_m_m", 0.0)),
        ("DESESTABILIZANTES", "", "", ""),
        ("Componente horizontal PAh", presiones.get("PA_h_ton_m", 0.0), presiones.get("brazo_PA_h_m", 0.0), presiones.get("M_volc_PA_h_ton_m_m", 0.0)),
        ("Total desestabilizante", "", "", presiones.get("M_volc_ton_m_m", 0.0)),
        ("RESULTANTE", "", "", ""),
        ("M neto = M_est - M_volc", "", "", presiones.get("M_resultante_ton_m_m", 0.0)),
        ("x resultante", "", "", presiones.get("x_resultante_m", 0.0)),
        ("e", "", "", presiones.get("e_m", 0.0)),
    ]
    return pd.DataFrame(filas, columns=["Concepto", "Fuerza [ton/m]", "Brazo x o y [m]", "Momento [ton·m/m]"])


# =============================================================================
# OVERRIDES - Separación manual de armado
# =============================================================================

def _as_prov_por_separacion(diametro_mm: float, separacion_cm: float) -> float:
    if separacion_cm <= 0 or not math.isfinite(separacion_cm):
        return 0.0
    return area_barra_cm2(diametro_mm) * 100.0 / separacion_cm


def _estado_acero_manual(As_req: float, As_prov: float) -> str:
    if not math.isfinite(As_req):
        return "No cumple: aumentar sección/refuerzo"
    if As_req <= 0:
        return "No aplica"
    return "OK" if As_prov >= As_req else "No cumple"


def calcular_diseno_fuste_dinamico(
    datos: DatosMuro,
    numero_cunas: int = 180,
    recubrimiento_cm: float = 7.5,
    diametro_vertical_mm: float = 16.0,
    diametro_horizontal_mm: float = 12.0,
    separacion_max_cm: float = 30.0,
    sep_vertical_manual_cm: float | None = None,
    sep_horizontal_manual_cm: float | None = None
) -> dict:
    """
    Diseño dinámico del fuste con separación manual.

    Si se ingresa `sep_vertical_manual_cm` o `sep_horizontal_manual_cm`,
    esa separación se usa directamente para calcular As provisto.
    """
    geometria = generar_puntos_muro(datos)
    tabla_trial, resultado_trial = calcular_trial_wedge_activo(datos, geometria, numero_cunas=numero_cunas)

    PA = resultado_trial["PA_ton_m"]
    delta_rad = math.radians(resultado_trial["delta_grados"])
    PA_h = PA * math.cos(delta_rad)

    h_activa = altura_relleno_en_muro(datos, geometria) if "altura_relleno_en_muro" in globals() else max(0.0, min(datos.altura_relleno, datos.H))
    brazo = h_activa / 3.0 if h_activa > 0 else 0.0

    gamma_p = 1.50
    Mu_ton_m = gamma_p * PA_h * brazo
    Vu_ton = gamma_p * PA_h

    b_cm = 100.0
    h_cm = datos.t_base * 100.0
    d_cm = max(h_cm - recubrimiento_cm - (diametro_vertical_mm / 10.0) / 2.0, 1.0)

    As_flexion = resolver_as_flexion_rectangular(Mu_ton_m, b_cm, d_cm, datos.fc, datos.fy)
    As_temp_total = 0.0018 * b_cm * h_cm
    As_temp_cara = As_temp_total / 2.0
    As_vertical_req = max(As_flexion, As_temp_cara) if math.isfinite(As_flexion) else float("inf")

    if sep_vertical_manual_cm is None:
        sep_vertical_cm, As_vertical_prov = seleccionar_separacion(
            area_barra_cm2(diametro_vertical_mm), As_vertical_req, separacion_max_cm
        )
        modo_sep_vertical = "Automática"
    else:
        sep_vertical_cm = float(sep_vertical_manual_cm)
        As_vertical_prov = _as_prov_por_separacion(diametro_vertical_mm, sep_vertical_cm)
        modo_sep_vertical = "Manual"

    As_horizontal_req = As_temp_cara
    if sep_horizontal_manual_cm is None:
        sep_horizontal_cm, As_horizontal_prov = seleccionar_separacion(
            area_barra_cm2(diametro_horizontal_mm), As_horizontal_req, separacion_max_cm
        )
        modo_sep_horizontal = "Automática"
    else:
        sep_horizontal_cm = float(sep_horizontal_manual_cm)
        As_horizontal_prov = _as_prov_por_separacion(diametro_horizontal_mm, sep_horizontal_cm)
        modo_sep_horizontal = "Manual"

    cort = verificar_cortante_rectangular(Vu_ton, b_cm, d_cm, datos.fc)
    q_base_ton_m2 = PA_h / h_activa if h_activa > 0 else 0.0
    estado_flexion = _estado_acero_manual(As_vertical_req, As_vertical_prov)
    estado_horizontal = _estado_acero_manual(As_horizontal_req, As_horizontal_prov)

    return {
        "PA_ton_m": PA,
        "delta_grados": resultado_trial["delta_grados"],
        "alfa_critico_grados": resultado_trial.get("alfa_grados", resultado_trial.get("alfa_critico_grados", 0.0)),
        "q_base_ton_m2": q_base_ton_m2,
        "altura_activa_relleno_m": h_activa,
        "brazo_activo_m": brazo,
        "Mu_ton_m_m": Mu_ton_m,
        "Vu_ton_m": Vu_ton,
        "b_cm": b_cm,
        "h_cm": h_cm,
        "d_cm": d_cm,
        "As_flexion_cm2_m": As_flexion,
        "As_temp_total_cm2_m": As_temp_total,
        "As_temp_cara_cm2_m": As_temp_cara,
        "As_vertical_req_cm2_m": As_vertical_req,
        "diametro_vertical_mm": diametro_vertical_mm,
        "separacion_vertical_cm": sep_vertical_cm,
        "As_vertical_prov_cm2_m": As_vertical_prov,
        "modo_sep_vertical": modo_sep_vertical,
        "estado_flexion": estado_flexion,
        "diametro_horizontal_mm": diametro_horizontal_mm,
        "As_horizontal_req_cm2_m": As_horizontal_req,
        "separacion_horizontal_cm": sep_horizontal_cm,
        "As_horizontal_prov_cm2_m": As_horizontal_prov,
        "modo_sep_horizontal": modo_sep_horizontal,
        "estado_horizontal": estado_horizontal,
        "phi_Vc_ton_m": cort["phi_Vc_ton_m"],
        "relacion_cortante": cort["relacion"],
        "estado_cortante": cort["estado"],
        "tabla_trial": tabla_trial,
        "resultado_trial": resultado_trial,
    }


def tabla_diseno_fuste(resultado: dict) -> pd.DataFrame:
    filas = [
        ("PA dinámico", resultado["PA_ton_m"], "ton/m"),
        ("δ asumido", resultado["delta_grados"], "grados"),
        ("α crítico", resultado["alfa_critico_grados"], "grados"),
        ("q base", resultado["q_base_ton_m2"], "ton/m²"),
        ("Altura activa de relleno", resultado.get("altura_activa_relleno_m", 0.0), "m"),
        ("Brazo activo", resultado.get("brazo_activo_m", 0.0), "m"),
        ("Mu fuste", resultado["Mu_ton_m_m"], "ton·m/m"),
        ("Vu fuste", resultado["Vu_ton_m"], "ton/m"),
        ("Peralte efectivo d", resultado["d_cm"], "cm"),
        ("As por flexión", resultado["As_flexion_cm2_m"], "cm²/m"),
        ("As temperatura por cara", resultado["As_temp_cara_cm2_m"], "cm²/m"),
        ("As vertical requerido", resultado["As_vertical_req_cm2_m"], "cm²/m"),
        ("As vertical provisto", resultado["As_vertical_prov_cm2_m"], "cm²/m"),
        ("Separación vertical", resultado["separacion_vertical_cm"], "cm"),
        ("Modo separación vertical", resultado.get("modo_sep_vertical", "Automática"), "-"),
        ("Estado flexión vertical", resultado.get("estado_flexion", "-"), "-"),
        ("As horizontal requerido", resultado["As_horizontal_req_cm2_m"], "cm²/m"),
        ("As horizontal provisto", resultado["As_horizontal_prov_cm2_m"], "cm²/m"),
        ("Separación horizontal", resultado["separacion_horizontal_cm"], "cm"),
        ("Modo separación horizontal", resultado.get("modo_sep_horizontal", "Automática"), "-"),
        ("Estado acero horizontal", resultado.get("estado_horizontal", "-"), "-"),
        ("φVc fuste", resultado["phi_Vc_ton_m"], "ton/m"),
        ("Relación Vu/φVc", resultado.get("relacion_cortante", 0.0), "-"),
        ("Estado cortante", resultado["estado_cortante"], "-"),
    ]
    return pd.DataFrame(filas, columns=["Parámetro", "Valor", "Unidad"])


def calcular_diseno_zapata_dinamico(
    datos: DatosMuro,
    numero_cunas: int = 180,
    recubrimiento_cm: float = 7.5,
    diametro_puntera_mm: float = 16.0,
    diametro_talon_mm: float = 16.0,
    separacion_max_cm: float = 30.0,
    sep_puntera_manual_cm: float | None = None,
    sep_talon_manual_cm: float | None = None
) -> dict:
    """
    Diseño preliminar de zapata con separación manual para puntera y talón.
    """
    presiones = calcular_presiones_contacto_servicio(datos, numero_cunas=numero_cunas)
    gamma_u = 1.50
    b_cm = 100.0
    h_cm = datos.hz * 100.0

    Lp = max(datos.puntera, 0.0)
    q0 = presion_lineal_en_x(datos, presiones, 0.0)
    q1 = presion_lineal_en_x(datos, presiones, datos.puntera)
    q_prom_p = (q0 + q1) / 2.0 if Lp > 0 else 0.0
    Mu_puntera = gamma_u * q_prom_p * Lp ** 2 / 2.0
    Vu_puntera = gamma_u * q_prom_p * Lp

    Lt = max(calcular_talon(datos), 0.0)
    x_t0 = datos.puntera + datos.t_base
    x_t1 = datos.B
    q_t0 = presion_lineal_en_x(datos, presiones, x_t0)
    q_t1 = presion_lineal_en_x(datos, presiones, x_t1)
    q_prom_t = (q_t0 + q_t1) / 2.0 if Lt > 0 else 0.0

    geom = generar_puntos_muro(datos)
    soil = suelo_sobre_talon(datos, geom) if "suelo_sobre_talon" in globals() else {"W_ton_m": datos.gamma_suelo * datos.altura_relleno * Lt}
    w_suelo_prom = soil["W_ton_m"] / Lt if Lt > 0 else 0.0
    w_neto_talon = w_suelo_prom - q_prom_t

    Mu_talon_firmado = gamma_u * w_neto_talon * Lt ** 2 / 2.0
    Vu_talon_firmado = gamma_u * w_neto_talon * Lt
    Mu_talon = abs(Mu_talon_firmado)
    Vu_talon = abs(Vu_talon_firmado)

    d_puntera_cm = h_cm - recubrimiento_cm - (diametro_puntera_mm / 10.0) / 2.0
    d_talon_cm = h_cm - recubrimiento_cm - (diametro_talon_mm / 10.0) / 2.0

    As_puntera = resolver_as_flexion_rectangular(Mu_puntera, b_cm, d_puntera_cm, datos.fc, datos.fy) if Lp > 0 else 0.0
    As_talon = resolver_as_flexion_rectangular(Mu_talon, b_cm, d_talon_cm, datos.fc, datos.fy) if Lt > 0 else 0.0

    As_temp_zapata = 0.0018 * b_cm * h_cm / 2.0
    As_puntera_req = max(As_puntera, As_temp_zapata) if Lp > 0 else 0.0
    As_talon_req = max(As_talon, As_temp_zapata) if Lt > 0 else 0.0

    if sep_puntera_manual_cm is None:
        sep_puntera_cm_calc, As_puntera_prov = seleccionar_separacion(area_barra_cm2(diametro_puntera_mm), As_puntera_req, separacion_max_cm)
        modo_sep_puntera = "Automática"
    else:
        sep_puntera_cm_calc = float(sep_puntera_manual_cm)
        As_puntera_prov = _as_prov_por_separacion(diametro_puntera_mm, sep_puntera_cm_calc) if Lp > 0 else 0.0
        modo_sep_puntera = "Manual"

    if sep_talon_manual_cm is None:
        sep_talon_cm_calc, As_talon_prov = seleccionar_separacion(area_barra_cm2(diametro_talon_mm), As_talon_req, separacion_max_cm)
        modo_sep_talon = "Automática"
    else:
        sep_talon_cm_calc = float(sep_talon_manual_cm)
        As_talon_prov = _as_prov_por_separacion(diametro_talon_mm, sep_talon_cm_calc) if Lt > 0 else 0.0
        modo_sep_talon = "Manual"

    return {
        "presiones": presiones,
        "Lp_m": Lp,
        "Lt_m": Lt,
        "q_puntera_prom_ton_m2": q_prom_p,
        "q_talon_prom_ton_m2": q_prom_t,
        "w_suelo_talon_ton_m2": w_suelo_prom,
        "w_neto_talon_ton_m2": w_neto_talon,
        "Mu_puntera_ton_m_m": Mu_puntera,
        "Vu_puntera_ton_m": Vu_puntera,
        "Mu_talon_ton_m_m": Mu_talon,
        "Mu_talon_firmado_ton_m_m": Mu_talon_firmado,
        "Vu_talon_ton_m": Vu_talon,
        "Vu_talon_firmado_ton_m": Vu_talon_firmado,
        "As_puntera_req_cm2_m": As_puntera_req,
        "As_puntera_prov_cm2_m": As_puntera_prov,
        "sep_puntera_cm": sep_puntera_cm_calc,
        "modo_sep_puntera": modo_sep_puntera,
        "As_talon_req_cm2_m": As_talon_req,
        "As_talon_prov_cm2_m": As_talon_prov,
        "sep_talon_cm": sep_talon_cm_calc,
        "modo_sep_talon": modo_sep_talon,
        "diametro_puntera_mm": diametro_puntera_mm,
        "diametro_talon_mm": diametro_talon_mm,
    }


def calcular_diseno_zapata_definitivo(
    datos: DatosMuro,
    numero_cunas: int = 180,
    recubrimiento_cm: float = 7.5,
    diametro_puntera_mm: float = 16.0,
    diametro_talon_mm: float = 16.0,
    separacion_max_cm: float = 30.0,
    sep_puntera_manual_cm: float | None = None,
    sep_talon_manual_cm: float | None = None
) -> dict:
    """
    Diseño de zapata con separación manual para puntera y talón.
    """
    res = calcular_diseno_zapata_dinamico(
        datos,
        numero_cunas,
        recubrimiento_cm,
        diametro_puntera_mm,
        diametro_talon_mm,
        separacion_max_cm,
        sep_puntera_manual_cm,
        sep_talon_manual_cm
    )
    presiones = res["presiones"]
    b_cm = 100.0
    h_cm = datos.hz * 100.0
    d_puntera_cm = h_cm - recubrimiento_cm - (diametro_puntera_mm / 10.0) / 2.0
    d_talon_cm = h_cm - recubrimiento_cm - (diametro_talon_mm / 10.0) / 2.0
    d_puntera_m = d_puntera_cm / 100.0
    d_talon_m = d_talon_cm / 100.0
    gamma_u = 1.50

    Lp = max(datos.puntera, 0.0)
    Lp_crit = max(Lp - d_puntera_m, 0.0)
    q0 = presion_lineal_en_x(datos, presiones, 0.0)
    qcrit_p = presion_lineal_en_x(datos, presiones, Lp_crit)
    q_prom_crit_p = (q0 + qcrit_p) / 2.0 if Lp_crit > 0 else 0.0
    Vu_puntera_crit = gamma_u * q_prom_crit_p * Lp_crit
    cortante_puntera = _cortante_o_no_aplica(Vu_puntera_crit, b_cm, d_puntera_cm, datos.fc, Lp_crit, "Lp")

    Lt = max(calcular_talon(datos), 0.0)
    Lt_crit = max(Lt - d_talon_m, 0.0)
    x_inicio_talon_crit = datos.puntera + datos.t_base + d_talon_m
    x_fin_talon = datos.B
    q_tcrit = presion_lineal_en_x(datos, presiones, x_inicio_talon_crit)
    q_tfin = presion_lineal_en_x(datos, presiones, x_fin_talon)
    q_prom_tcrit = (q_tcrit + q_tfin) / 2.0 if Lt_crit > 0 else 0.0

    geom = generar_puntos_muro(datos)
    soil = suelo_sobre_talon(datos, geom) if "suelo_sobre_talon" in globals() else {"W_ton_m": datos.gamma_suelo * datos.altura_relleno * Lt}
    w_suelo_prom = soil["W_ton_m"] / Lt if Lt > 0 else 0.0
    w_neto = w_suelo_prom - q_prom_tcrit
    Vu_talon_crit = gamma_u * abs(w_neto) * Lt_crit
    cortante_talon = _cortante_o_no_aplica(Vu_talon_crit, b_cm, d_talon_cm, datos.fc, Lt_crit, "Lt")

    estado_as_puntera = _estado_acero_manual(res["As_puntera_req_cm2_m"], res["As_puntera_prov_cm2_m"]) if Lp > 0 else "No aplica"
    estado_as_talon = _estado_acero_manual(res["As_talon_req_cm2_m"], res["As_talon_prov_cm2_m"]) if Lt > 0 else "No aplica"

    def falla(e):
        txt = str(e).lower()
        return txt.startswith("no cumple") or "revisar" in txt

    estados = [presiones["estado_q"], cortante_puntera["estado"], cortante_talon["estado"], estado_as_puntera, estado_as_talon]
    estado_global = "Revisar" if any(falla(e) for e in estados) else "OK"

    res.update({
        "d_puntera_cm": d_puntera_cm,
        "d_talon_cm": d_talon_cm,
        "Lp_crit_m": Lp_crit,
        "Lt_crit_m": Lt_crit,
        "Vu_puntera_crit_ton_m": Vu_puntera_crit,
        "Vu_talon_crit_ton_m": Vu_talon_crit,
        "cortante_puntera": cortante_puntera,
        "cortante_talon": cortante_talon,
        "estado_as_puntera": estado_as_puntera,
        "estado_as_talon": estado_as_talon,
        "ld_puntera_cm": 0.0,
        "ld_talon_cm": 0.0,
        "longitud_disponible_puntera_cm": 0.0,
        "longitud_disponible_talon_cm": 0.0,
        "estado_ld_puntera": "No se evalúa en dashboard",
        "estado_ld_talon": "No se evalúa en dashboard",
        "estado_global_zapata": estado_global,
    })
    return res


def tabla_diseno_zapata_definitivo(resultado: dict) -> pd.DataFrame:
    filas = [
        ("Estado global zapata", resultado["estado_global_zapata"], "-"),
        ("qmax", resultado["presiones"]["qmax_ton_m2"], "ton/m²"),
        ("qmin", resultado["presiones"]["qmin_ton_m2"], "ton/m²"),
        ("qa", resultado["presiones"]["q_adm_ton_m2"], "ton/m²"),
        ("Estado presión admisible", resultado["presiones"]["estado_q"], "-"),

        ("Mu puntera", resultado["Mu_puntera_ton_m_m"], "ton·m/m"),
        ("As puntera requerido", resultado["As_puntera_req_cm2_m"], "cm²/m"),
        ("As puntera provisto", resultado["As_puntera_prov_cm2_m"], "cm²/m"),
        ("Armado puntera", _formato_armado(resultado["diametro_puntera_mm"], resultado["sep_puntera_cm"]), "mm @ cm"),
        ("Modo separación puntera", resultado.get("modo_sep_puntera", "Automática"), "-"),
        ("Estado flexión puntera", resultado.get("estado_as_puntera", "No aplica"), "-"),

        ("Mu talón", resultado["Mu_talon_ton_m_m"], "ton·m/m"),
        ("Mu talón firmado", resultado.get("Mu_talon_firmado_ton_m_m", 0.0), "ton·m/m"),
        ("As talón requerido", resultado["As_talon_req_cm2_m"], "cm²/m"),
        ("As talón provisto", resultado["As_talon_prov_cm2_m"], "cm²/m"),
        ("Armado talón", _formato_armado(resultado["diametro_talon_mm"], resultado["sep_talon_cm"]), "mm @ cm"),
        ("Modo separación talón", resultado.get("modo_sep_talon", "Automática"), "-"),
        ("Estado flexión talón", resultado.get("estado_as_talon", "No aplica"), "-"),

        ("Vu puntera crítico", resultado["Vu_puntera_crit_ton_m"], "ton/m"),
        ("φVc puntera", resultado["cortante_puntera"]["phi_Vc_ton_m"], "ton/m"),
        ("Relación Vu/φVc puntera", resultado["cortante_puntera"]["relacion"], "-"),
        ("Estado cortante puntera", resultado["cortante_puntera"]["estado"], "-"),

        ("Vu talón crítico", resultado["Vu_talon_crit_ton_m"], "ton/m"),
        ("φVc talón", resultado["cortante_talon"]["phi_Vc_ton_m"], "ton/m"),
        ("Relación Vu/φVc talón", resultado["cortante_talon"]["relacion"], "-"),
        ("Estado cortante talón", resultado["cortante_talon"]["estado"], "-"),
    ]
    return pd.DataFrame(filas, columns=["Verificación", "Valor", "Unidad"])


# =============================================================================
# OVERRIDE NOMENCLATURA ARMADO ZAPATA POR CARAS
# =============================================================================
def tabla_diseno_zapata_definitivo(resultado: dict) -> pd.DataFrame:
    """
    Tabla de zapata usando nomenclatura por caras:
    - acero inferior de zapata;
    - acero superior de zapata.

    Internamente se conserva:
    - puntera = acero inferior;
    - talón = acero superior.
    """
    filas = [
        ("Estado global zapata", resultado["estado_global_zapata"], "-"),
        ("qmax", resultado["presiones"]["qmax_ton_m2"], "ton/m²"),
        ("qmin", resultado["presiones"]["qmin_ton_m2"], "ton/m²"),
        ("qa", resultado["presiones"]["q_adm_ton_m2"], "ton/m²"),
        ("Estado presión admisible", resultado["presiones"]["estado_q"], "-"),

        ("Mu puntera", resultado["Mu_puntera_ton_m_m"], "ton·m/m"),
        ("As inferior requerido", resultado["As_puntera_req_cm2_m"], "cm²/m"),
        ("As inferior provisto", resultado["As_puntera_prov_cm2_m"], "cm²/m"),
        ("Armado inferior zapata", _formato_armado(resultado["diametro_puntera_mm"], resultado["sep_puntera_cm"]), "mm @ cm"),
        ("Modo separación inferior", resultado.get("modo_sep_puntera", "Manual"), "-"),
        ("Estado flexión acero inferior", resultado.get("estado_as_puntera", "No aplica"), "-"),

        ("Mu talón", resultado["Mu_talon_ton_m_m"], "ton·m/m"),
        ("Mu talón firmado", resultado.get("Mu_talon_firmado_ton_m_m", 0.0), "ton·m/m"),
        ("As superior requerido", resultado["As_talon_req_cm2_m"], "cm²/m"),
        ("As superior provisto", resultado["As_talon_prov_cm2_m"], "cm²/m"),
        ("Armado superior zapata", _formato_armado(resultado["diametro_talon_mm"], resultado["sep_talon_cm"]), "mm @ cm"),
        ("Modo separación superior", resultado.get("modo_sep_talon", "Manual"), "-"),
        ("Estado flexión acero superior", resultado.get("estado_as_talon", "No aplica"), "-"),

        ("Vu puntera crítico", resultado["Vu_puntera_crit_ton_m"], "ton/m"),
        ("φVc puntera", resultado["cortante_puntera"]["phi_Vc_ton_m"], "ton/m"),
        ("Relación Vu/φVc puntera", resultado["cortante_puntera"]["relacion"], "-"),
        ("Estado cortante puntera", resultado["cortante_puntera"]["estado"], "-"),

        ("Vu talón crítico", resultado["Vu_talon_crit_ton_m"], "ton/m"),
        ("φVc talón", resultado["cortante_talon"]["phi_Vc_ton_m"], "ton/m"),
        ("Relación Vu/φVc talón", resultado["cortante_talon"]["relacion"], "-"),
        ("Estado cortante talón", resultado["cortante_talon"]["estado"], "-"),
    ]
    return pd.DataFrame(filas, columns=["Verificación", "Valor", "Unidad"])


# =============================================================================
# OVERRIDE FINAL - Zapata diseñada por caras crítica inferior/superior
# =============================================================================

def _as_req_zapata_por_momento(Mu_ton_m: float, b_cm: float, d_cm: float, fc: float, fy: float, As_min: float) -> tuple[float, float]:
    """
    Devuelve As_req y As_flex para una cara de la zapata.
    Si el momento es cero, la cara conserva As mínimo.
    """
    As_flex = resolver_as_flexion_rectangular(Mu_ton_m, b_cm, d_cm, fc, fy) if Mu_ton_m > 0 else 0.0
    As_req = max(As_flex, As_min) if math.isfinite(As_flex) else float("inf")
    return As_req, As_flex


def calcular_diseno_zapata_dinamico(
    datos: DatosMuro,
    numero_cunas: int = 180,
    recubrimiento_cm: float = 7.5,
    diametro_puntera_mm: float = 16.0,
    diametro_talon_mm: float = 16.0,
    separacion_max_cm: float = 30.0,
    sep_puntera_manual_cm: float | None = None,
    sep_talon_manual_cm: float | None = None
) -> dict:
    """
    Diseño preliminar de zapata por caras.

    Concepto:
    - Se calculan por separado los momentos de puntera y talón porque provienen
      de zonas distintas.
    - Pero el refuerzo se reporta por CARAS de la zapata:
        inferior = cara inferior de toda la zapata;
        superior = cara superior de toda la zapata.
    - Si no existe puntera, el acero inferior no se vuelve NaN; queda con As
      mínimo de la zapata.
    """
    presiones = calcular_presiones_contacto_servicio(datos, numero_cunas=numero_cunas)
    gamma_u = 1.50
    b_cm = 100.0
    h_cm = datos.hz * 100.0

    Lp = max(datos.puntera, 0.0)
    Lt = max(calcular_talon(datos), 0.0)

    # Momento de puntera: usualmente controla acero inferior.
    q0 = presion_lineal_en_x(datos, presiones, 0.0)
    q1 = presion_lineal_en_x(datos, presiones, datos.puntera)
    q_prom_p = (q0 + q1) / 2.0 if Lp > 0 else 0.0
    Mu_puntera = gamma_u * q_prom_p * Lp ** 2 / 2.0
    Vu_puntera = gamma_u * q_prom_p * Lp

    # Momento de talón: usualmente controla acero superior.
    x_t0 = datos.puntera + datos.t_base
    x_t1 = datos.B
    q_t0 = presion_lineal_en_x(datos, presiones, x_t0)
    q_t1 = presion_lineal_en_x(datos, presiones, x_t1)
    q_prom_t = (q_t0 + q_t1) / 2.0 if Lt > 0 else 0.0

    geom = generar_puntos_muro(datos)
    if "suelo_sobre_talon" in globals():
        soil = suelo_sobre_talon(datos, geom)
    else:
        soil = {"W_ton_m": datos.gamma_suelo * datos.altura_relleno * Lt}

    w_suelo_prom = soil["W_ton_m"] / Lt if Lt > 0 else 0.0
    w_neto_talon = w_suelo_prom - q_prom_t

    Mu_talon_firmado = gamma_u * w_neto_talon * Lt ** 2 / 2.0
    Vu_talon_firmado = gamma_u * w_neto_talon * Lt
    Mu_talon = abs(Mu_talon_firmado)
    Vu_talon = abs(Vu_talon_firmado)

    d_inferior_cm = h_cm - recubrimiento_cm - (diametro_puntera_mm / 10.0) / 2.0
    d_superior_cm = h_cm - recubrimiento_cm - (diametro_talon_mm / 10.0) / 2.0

    # Mínimo por cara de la zapata. Se aplica aunque la puntera sea cero.
    As_min_cara = 0.0018 * b_cm * h_cm / 2.0

    # Cara inferior: toma el momento inferior crítico. En este modelo, el principal
    # proviene de puntera. Si puntera=0, queda As mínimo.
    As_inferior_req, As_inferior_flex = _as_req_zapata_por_momento(
        Mu_puntera, b_cm, d_inferior_cm, datos.fc, datos.fy, As_min_cara
    )

    # Cara superior: toma el momento superior crítico. En este modelo, el principal
    # proviene del talón.
    As_superior_req, As_superior_flex = _as_req_zapata_por_momento(
        Mu_talon, b_cm, d_superior_cm, datos.fc, datos.fy, As_min_cara
    )

    if sep_puntera_manual_cm is None:
        sep_inferior_cm, As_inferior_prov = seleccionar_separacion(
            area_barra_cm2(diametro_puntera_mm), As_inferior_req, separacion_max_cm
        )
        modo_sep_inferior = "Automática"
    else:
        sep_inferior_cm = float(sep_puntera_manual_cm)
        As_inferior_prov = _as_prov_por_separacion(diametro_puntera_mm, sep_inferior_cm)
        modo_sep_inferior = "Manual"

    if sep_talon_manual_cm is None:
        sep_superior_cm, As_superior_prov = seleccionar_separacion(
            area_barra_cm2(diametro_talon_mm), As_superior_req, separacion_max_cm
        )
        modo_sep_superior = "Automática"
    else:
        sep_superior_cm = float(sep_talon_manual_cm)
        As_superior_prov = _as_prov_por_separacion(diametro_talon_mm, sep_superior_cm)
        modo_sep_superior = "Manual"

    return {
        "presiones": presiones,
        "Lp_m": Lp,
        "Lt_m": Lt,
        "q_puntera_prom_ton_m2": q_prom_p,
        "q_talon_prom_ton_m2": q_prom_t,
        "w_suelo_talon_ton_m2": w_suelo_prom,
        "w_neto_talon_ton_m2": w_neto_talon,

        "Mu_puntera_ton_m_m": Mu_puntera,
        "Vu_puntera_ton_m": Vu_puntera,
        "Mu_talon_ton_m_m": Mu_talon,
        "Mu_talon_firmado_ton_m_m": Mu_talon_firmado,
        "Vu_talon_ton_m": Vu_talon,
        "Vu_talon_firmado_ton_m": Vu_talon_firmado,

        "As_min_zapata_cara_cm2_m": As_min_cara,
        "As_inferior_flexion_cm2_m": As_inferior_flex,
        "As_superior_flexion_cm2_m": As_superior_flex,
        "As_inferior_req_cm2_m": As_inferior_req,
        "As_inferior_prov_cm2_m": As_inferior_prov,
        "sep_inferior_cm": sep_inferior_cm,
        "diametro_inferior_mm": diametro_puntera_mm,
        "modo_sep_inferior": modo_sep_inferior,
        "As_superior_req_cm2_m": As_superior_req,
        "As_superior_prov_cm2_m": As_superior_prov,
        "sep_superior_cm": sep_superior_cm,
        "diametro_superior_mm": diametro_talon_mm,
        "modo_sep_superior": modo_sep_superior,

        # Alias para compatibilidad con app/dibujos existentes.
        "As_puntera_req_cm2_m": As_inferior_req,
        "As_puntera_prov_cm2_m": As_inferior_prov,
        "sep_puntera_cm": sep_inferior_cm,
        "diametro_puntera_mm": diametro_puntera_mm,
        "modo_sep_puntera": modo_sep_inferior,
        "As_talon_req_cm2_m": As_superior_req,
        "As_talon_prov_cm2_m": As_superior_prov,
        "sep_talon_cm": sep_superior_cm,
        "diametro_talon_mm": diametro_talon_mm,
        "modo_sep_talon": modo_sep_superior,
    }


def calcular_diseno_zapata_definitivo(
    datos: DatosMuro,
    numero_cunas: int = 180,
    recubrimiento_cm: float = 7.5,
    diametro_puntera_mm: float = 16.0,
    diametro_talon_mm: float = 16.0,
    separacion_max_cm: float = 30.0,
    sep_puntera_manual_cm: float | None = None,
    sep_talon_manual_cm: float | None = None
) -> dict:
    """
    Diseño de zapata por cara inferior/superior.
    """
    res = calcular_diseno_zapata_dinamico(
        datos,
        numero_cunas,
        recubrimiento_cm,
        diametro_puntera_mm,
        diametro_talon_mm,
        separacion_max_cm,
        sep_puntera_manual_cm,
        sep_talon_manual_cm
    )

    presiones = res["presiones"]
    b_cm = 100.0
    h_cm = datos.hz * 100.0
    d_inferior_cm = h_cm - recubrimiento_cm - (diametro_puntera_mm / 10.0) / 2.0
    d_superior_cm = h_cm - recubrimiento_cm - (diametro_talon_mm / 10.0) / 2.0
    d_puntera_m = d_inferior_cm / 100.0
    d_talon_m = d_superior_cm / 100.0
    gamma_u = 1.50

    Lp = max(datos.puntera, 0.0)
    Lp_crit = max(Lp - d_puntera_m, 0.0)
    q0 = presion_lineal_en_x(datos, presiones, 0.0)
    qcrit_p = presion_lineal_en_x(datos, presiones, Lp_crit)
    q_prom_crit_p = (q0 + qcrit_p) / 2.0 if Lp_crit > 0 else 0.0
    Vu_puntera_crit = gamma_u * q_prom_crit_p * Lp_crit
    cortante_puntera = _cortante_o_no_aplica(Vu_puntera_crit, b_cm, d_inferior_cm, datos.fc, Lp_crit, "Lp")

    Lt = max(calcular_talon(datos), 0.0)
    Lt_crit = max(Lt - d_talon_m, 0.0)
    x_inicio_talon_crit = datos.puntera + datos.t_base + d_talon_m
    x_fin_talon = datos.B
    q_tcrit = presion_lineal_en_x(datos, presiones, x_inicio_talon_crit)
    q_tfin = presion_lineal_en_x(datos, presiones, x_fin_talon)
    q_prom_tcrit = (q_tcrit + q_tfin) / 2.0 if Lt_crit > 0 else 0.0

    geom = generar_puntos_muro(datos)
    if "suelo_sobre_talon" in globals():
        soil = suelo_sobre_talon(datos, geom)
    else:
        soil = {"W_ton_m": datos.gamma_suelo * datos.altura_relleno * Lt}
    w_suelo_prom = soil["W_ton_m"] / Lt if Lt > 0 else 0.0
    w_neto = w_suelo_prom - q_prom_tcrit
    Vu_talon_crit = gamma_u * abs(w_neto) * Lt_crit
    cortante_talon = _cortante_o_no_aplica(Vu_talon_crit, b_cm, d_superior_cm, datos.fc, Lt_crit, "Lt")

    estado_as_inferior = _estado_acero_manual(res["As_inferior_req_cm2_m"], res["As_inferior_prov_cm2_m"])
    estado_as_superior = _estado_acero_manual(res["As_superior_req_cm2_m"], res["As_superior_prov_cm2_m"])

    def falla(e):
        txt = str(e).lower()
        return txt.startswith("no cumple") or "revisar" in txt

    estados = [
        presiones["estado_q"],
        cortante_puntera["estado"],
        cortante_talon["estado"],
        estado_as_inferior,
        estado_as_superior,
    ]
    estado_global = "Revisar" if any(falla(e) for e in estados) else "OK"

    res.update({
        "d_puntera_cm": d_inferior_cm,
        "d_talon_cm": d_superior_cm,
        "d_inferior_cm": d_inferior_cm,
        "d_superior_cm": d_superior_cm,
        "Lp_crit_m": Lp_crit,
        "Lt_crit_m": Lt_crit,
        "Vu_puntera_crit_ton_m": Vu_puntera_crit,
        "Vu_talon_crit_ton_m": Vu_talon_crit,
        "cortante_puntera": cortante_puntera,
        "cortante_talon": cortante_talon,
        "estado_as_inferior": estado_as_inferior,
        "estado_as_superior": estado_as_superior,
        # Alias para compatibilidad.
        "estado_as_puntera": estado_as_inferior,
        "estado_as_talon": estado_as_superior,
        "ld_puntera_cm": 0.0,
        "ld_talon_cm": 0.0,
        "longitud_disponible_puntera_cm": 0.0,
        "longitud_disponible_talon_cm": 0.0,
        "estado_ld_puntera": "No se evalúa en dashboard",
        "estado_ld_talon": "No se evalúa en dashboard",
        "estado_global_zapata": estado_global,
    })
    return res


def tabla_diseno_zapata_definitivo(resultado: dict) -> pd.DataFrame:
    """
    Tabla de zapata por cara inferior y cara superior.
    """
    filas = [
        ("Estado global zapata", resultado["estado_global_zapata"], "-"),
        ("qmax", resultado["presiones"]["qmax_ton_m2"], "ton/m²"),
        ("qmin", resultado["presiones"]["qmin_ton_m2"], "ton/m²"),
        ("qa", resultado["presiones"]["q_adm_ton_m2"], "ton/m²"),
        ("Estado presión admisible", resultado["presiones"]["estado_q"], "-"),

        ("Momento puntera para cara inferior", resultado["Mu_puntera_ton_m_m"], "ton·m/m"),
        ("As flexión inferior", resultado["As_inferior_flexion_cm2_m"], "cm²/m"),
        ("As mínimo por cara", resultado["As_min_zapata_cara_cm2_m"], "cm²/m"),
        ("As inferior requerido", resultado["As_inferior_req_cm2_m"], "cm²/m"),
        ("As inferior provisto", resultado["As_inferior_prov_cm2_m"], "cm²/m"),
        ("Armado inferior zapata", _formato_armado(resultado["diametro_inferior_mm"], resultado["sep_inferior_cm"]), "mm @ cm"),
        ("Modo separación inferior", resultado.get("modo_sep_inferior", "Manual"), "-"),
        ("Estado flexión acero inferior", resultado.get("estado_as_inferior", "No aplica"), "-"),

        ("Momento talón para cara superior", resultado["Mu_talon_ton_m_m"], "ton·m/m"),
        ("Mu talón firmado", resultado.get("Mu_talon_firmado_ton_m_m", 0.0), "ton·m/m"),
        ("As flexión superior", resultado["As_superior_flexion_cm2_m"], "cm²/m"),
        ("As mínimo por cara", resultado["As_min_zapata_cara_cm2_m"], "cm²/m"),
        ("As superior requerido", resultado["As_superior_req_cm2_m"], "cm²/m"),
        ("As superior provisto", resultado["As_superior_prov_cm2_m"], "cm²/m"),
        ("Armado superior zapata", _formato_armado(resultado["diametro_superior_mm"], resultado["sep_superior_cm"]), "mm @ cm"),
        ("Modo separación superior", resultado.get("modo_sep_superior", "Manual"), "-"),
        ("Estado flexión acero superior", resultado.get("estado_as_superior", "No aplica"), "-"),

        ("Vu puntera crítico", resultado["Vu_puntera_crit_ton_m"], "ton/m"),
        ("φVc puntera", resultado["cortante_puntera"]["phi_Vc_ton_m"], "ton/m"),
        ("Relación Vu/φVc puntera", resultado["cortante_puntera"]["relacion"], "-"),
        ("Estado cortante puntera", resultado["cortante_puntera"]["estado"], "-"),

        ("Vu talón crítico", resultado["Vu_talon_crit_ton_m"], "ton/m"),
        ("φVc talón", resultado["cortante_talon"]["phi_Vc_ton_m"], "ton/m"),
        ("Relación Vu/φVc talón", resultado["cortante_talon"]["relacion"], "-"),
        ("Estado cortante talón", resultado["cortante_talon"]["estado"], "-"),
    ]
    return pd.DataFrame(filas, columns=["Verificación", "Valor", "Unidad"])


# =============================================================================
# OVERRIDE FINAL - Zapata por signo de momento + relleno sobre puntera
# =============================================================================

def suelo_sobre_puntera(datos: DatosMuro) -> dict:
    """
    Peso de suelo sobre la puntera. Es una carga vertical hacia abajo que puede
    reducir o invertir el momento de la puntera.
    """
    h = max(float(getattr(datos, "altura_relleno_puntera", 0.0)), 0.0)
    Lp = max(datos.puntera, 0.0)
    area = Lp * h
    W = area * datos.gamma_suelo
    x_c = Lp / 2.0 if Lp > 0 else 0.0
    return {"h_m": h, "area_m2": area, "W_ton_m": W, "x_centroide_m": x_c, "w_ton_m2": h * datos.gamma_suelo}


def calcular_presiones_contacto_servicio(datos: DatosMuro, numero_cunas: int = 180) -> dict:
    """
    Equilibrio externo incluyendo suelo sobre puntera.
    """
    geometria = generar_puntos_muro(datos)
    tabla_trial, resultado_trial = calcular_trial_wedge_activo(datos, geometria, numero_cunas=numero_cunas)

    PA = resultado_trial["PA_ton_m"]
    delta = math.radians(resultado_trial["delta_grados"])
    PA_v = PA * math.sin(delta)
    PA_h = PA * math.cos(delta)

    L_puntera = max(datos.puntera, 0.0)
    L_bajo_fuste = max(datos.t_base, 0.0)
    L_talon = max(calcular_talon(datos), 0.0)

    W_zapata_puntera = L_puntera * datos.hz * datos.gamma_hormigon
    x_zapata_puntera = L_puntera / 2.0 if L_puntera > 0 else 0.0

    W_zapata_fuste = L_bajo_fuste * datos.hz * datos.gamma_hormigon
    x_zapata_fuste = datos.puntera + L_bajo_fuste / 2.0

    W_zapata_talon = L_talon * datos.hz * datos.gamma_hormigon
    x_zapata_talon = datos.puntera + datos.t_base + L_talon / 2.0 if L_talon > 0 else datos.B

    W_zapata = W_zapata_puntera + W_zapata_fuste + W_zapata_talon
    x_zapata = (
        W_zapata_puntera * x_zapata_puntera
        + W_zapata_fuste * x_zapata_fuste
        + W_zapata_talon * x_zapata_talon
    ) / W_zapata if W_zapata > 0 else 0.0

    area_fuste = (datos.t_base + datos.t_corona) / 2.0 * datos.H
    W_fuste = area_fuste * datos.gamma_hormigon
    x_fuste = datos.puntera + (datos.t_base + datos.t_corona) / 4.0

    if datos.usar_llave:
        W_dentellon = datos.ancho_llave * datos.profundidad_llave * datos.gamma_hormigon
        x_dentellon = datos.pos_llave
    else:
        W_dentellon = 0.0
        x_dentellon = 0.0

    if "suelo_sobre_talon" in globals():
        suelo_talon = suelo_sobre_talon(datos, geometria)
    else:
        suelo_talon = {
            "area_m2": L_talon * datos.altura_relleno,
            "W_ton_m": L_talon * datos.altura_relleno * datos.gamma_suelo,
            "x_centroide_m": datos.puntera + datos.t_base + L_talon / 2.0,
            "h_inicio_m": datos.altura_relleno,
            "h_fin_m": datos.altura_relleno,
        }
    W_suelo_talon = suelo_talon["W_ton_m"]
    x_suelo_talon = suelo_talon["x_centroide_m"]

    suelo_punt = suelo_sobre_puntera(datos)
    W_suelo_puntera = suelo_punt["W_ton_m"]
    x_suelo_puntera = suelo_punt["x_centroide_m"]

    M_est_zapata_puntera = W_zapata_puntera * x_zapata_puntera
    M_est_zapata_fuste = W_zapata_fuste * x_zapata_fuste
    M_est_zapata_talon = W_zapata_talon * x_zapata_talon
    M_est_zapata = M_est_zapata_puntera + M_est_zapata_fuste + M_est_zapata_talon
    M_est_fuste = W_fuste * x_fuste
    M_est_dentellon = W_dentellon * x_dentellon
    M_est_suelo_talon = W_suelo_talon * x_suelo_talon
    M_est_suelo_puntera = W_suelo_puntera * x_suelo_puntera
    M_est_PA_v = PA_v * datos.B

    V = W_zapata + W_fuste + W_dentellon + W_suelo_talon + W_suelo_puntera + PA_v
    M_est = M_est_zapata + M_est_fuste + M_est_dentellon + M_est_suelo_talon + M_est_suelo_puntera + M_est_PA_v

    h_activa = altura_relleno_en_muro(datos, geometria) if "altura_relleno_en_muro" in globals() else max(0.0, min(datos.altura_relleno, datos.H))
    brazo_PA_h = h_activa / 3.0 if h_activa > 0 else 0.0
    M_volc_PA_h = PA_h * brazo_PA_h
    M_volc = M_volc_PA_h

    M_resultante = M_est - M_volc
    x_resultante = M_resultante / V if V > 0 else float("nan")
    e = datos.B / 2.0 - x_resultante

    if abs(e) <= datos.B / 6.0:
        q_prom = V / datos.B
        q_max = q_prom * (1.0 + 6.0 * e / datos.B)
        q_min = q_prom * (1.0 - 6.0 * e / datos.B)
    else:
        B_efectivo = 3.0 * (datos.B / 2.0 - abs(e))
        q_max = 2.0 * V / B_efectivo if B_efectivo > 0 else float("inf")
        q_min = 0.0

    estado_q = "OK" if math.isfinite(q_max) and q_max <= datos.qa else "No cumple"

    return {
        "PA_ton_m": PA, "PA_h_ton_m": PA_h, "PA_v_ton_m": PA_v,
        "delta_grados": resultado_trial["delta_grados"],
        "V_total_ton_m": V, "M_est_ton_m_m": M_est, "M_volc_ton_m_m": M_volc,
        "M_resultante_ton_m_m": M_resultante, "x_resultante_m": x_resultante, "e_m": e,
        "qmax_ton_m2": q_max, "qmin_ton_m2": q_min, "q_adm_ton_m2": datos.qa, "estado_q": estado_q,

        "W_zapata_ton_m": W_zapata, "x_zapata_m": x_zapata, "M_est_zapata_ton_m_m": M_est_zapata,
        "W_zapata_puntera_ton_m": W_zapata_puntera, "x_zapata_puntera_m": x_zapata_puntera, "M_est_zapata_puntera_ton_m_m": M_est_zapata_puntera,
        "W_zapata_fuste_ton_m": W_zapata_fuste, "x_zapata_fuste_m": x_zapata_fuste, "M_est_zapata_fuste_ton_m_m": M_est_zapata_fuste,
        "W_zapata_talon_ton_m": W_zapata_talon, "x_zapata_talon_m": x_zapata_talon, "M_est_zapata_talon_ton_m_m": M_est_zapata_talon,

        "W_fuste_ton_m": W_fuste, "x_fuste_m": x_fuste, "M_est_fuste_ton_m_m": M_est_fuste,
        "W_dentellon_ton_m": W_dentellon, "x_dentellon_m": x_dentellon, "M_est_dentellon_ton_m_m": M_est_dentellon,

        "W_suelo_puntera_ton_m": W_suelo_puntera,
        "x_suelo_puntera_m": x_suelo_puntera,
        "M_est_suelo_puntera_ton_m_m": M_est_suelo_puntera,
        "h_suelo_puntera_m": suelo_punt["h_m"],
        "w_suelo_puntera_ton_m2": suelo_punt["w_ton_m2"],

        "W_suelo_talon_ton_m": W_suelo_talon,
        "x_suelo_talon_m": x_suelo_talon,
        "M_est_suelo_talon_ton_m_m": M_est_suelo_talon,
        "area_suelo_talon_m2": suelo_talon.get("area_m2", 0.0),
        "h_suelo_talon_inicio_m": suelo_talon.get("h_inicio_m", 0.0),
        "h_suelo_talon_fin_m": suelo_talon.get("h_fin_m", 0.0),

        "W_pendiente_ton_m": 0.0, "x_pendiente_m": x_suelo_talon, "M_est_pendiente_ton_m_m": 0.0,
        "M_est_PA_v_ton_m_m": M_est_PA_v,
        "brazo_PA_h_m": brazo_PA_h, "M_volc_PA_h_ton_m_m": M_volc_PA_h,
    }


def calcular_diseno_zapata_dinamico(
    datos: DatosMuro,
    numero_cunas: int = 180,
    recubrimiento_cm: float = 7.5,
    diametro_puntera_mm: float = 16.0,
    diametro_talon_mm: float = 16.0,
    separacion_max_cm: float = 30.0,
    sep_puntera_manual_cm: float | None = None,
    sep_talon_manual_cm: float | None = None
) -> dict:
    """
    Diseño preliminar de zapata por signo de momento.

    Se calcula puntera y talón por separado, pero la demanda se asigna a cara
    inferior o superior según el signo físico del momento.
    """
    presiones = calcular_presiones_contacto_servicio(datos, numero_cunas=numero_cunas)
    gamma_u = 1.50
    b_cm = 100.0
    h_cm = datos.hz * 100.0
    As_min_cara = 0.0018 * b_cm * h_cm / 2.0

    Lp = max(datos.puntera, 0.0)
    Lt = max(calcular_talon(datos), 0.0)

    q0 = presion_lineal_en_x(datos, presiones, 0.0)
    q1 = presion_lineal_en_x(datos, presiones, datos.puntera)
    q_prom_p = (q0 + q1) / 2.0 if Lp > 0 else 0.0
    w_puntera_down = suelo_sobre_puntera(datos)["w_ton_m2"] if Lp > 0 else 0.0
    w_neto_puntera = q_prom_p - w_puntera_down  # positivo = hacia arriba
    Mu_puntera_firmado = gamma_u * w_neto_puntera * Lp ** 2 / 2.0
    Vu_puntera = gamma_u * abs(w_neto_puntera) * Lp

    x_t0 = datos.puntera + datos.t_base
    x_t1 = datos.B
    q_t0 = presion_lineal_en_x(datos, presiones, x_t0)
    q_t1 = presion_lineal_en_x(datos, presiones, x_t1)
    q_prom_t = (q_t0 + q_t1) / 2.0 if Lt > 0 else 0.0

    geom = generar_puntos_muro(datos)
    if "suelo_sobre_talon" in globals():
        soil_t = suelo_sobre_talon(datos, geom)
    else:
        soil_t = {"W_ton_m": datos.gamma_suelo * datos.altura_relleno * Lt}
    w_talon_down = soil_t["W_ton_m"] / Lt if Lt > 0 else 0.0
    w_neto_talon = w_talon_down - q_prom_t  # positivo = hacia abajo
    Mu_talon_firmado = gamma_u * w_neto_talon * Lt ** 2 / 2.0
    Vu_talon = gamma_u * abs(w_neto_talon) * Lt

    M_inferior = 0.0
    M_superior = 0.0

    # Puntera: neto hacia arriba -> tracción inferior. Neto hacia abajo -> tracción superior.
    if Mu_puntera_firmado >= 0:
        M_inferior += abs(Mu_puntera_firmado)
    else:
        M_superior += abs(Mu_puntera_firmado)

    # Talón: neto hacia abajo -> tracción superior. Neto hacia arriba -> tracción inferior.
    if Mu_talon_firmado >= 0:
        M_superior += abs(Mu_talon_firmado)
    else:
        M_inferior += abs(Mu_talon_firmado)

    d_inferior_cm = h_cm - recubrimiento_cm - (diametro_puntera_mm / 10.0) / 2.0
    d_superior_cm = h_cm - recubrimiento_cm - (diametro_talon_mm / 10.0) / 2.0

    As_inferior_req, As_inferior_flex = _as_req_zapata_por_momento(M_inferior, b_cm, d_inferior_cm, datos.fc, datos.fy, As_min_cara)
    As_superior_req, As_superior_flex = _as_req_zapata_por_momento(M_superior, b_cm, d_superior_cm, datos.fc, datos.fy, As_min_cara)

    if sep_puntera_manual_cm is None:
        sep_inferior_cm, As_inferior_prov = seleccionar_separacion(area_barra_cm2(diametro_puntera_mm), As_inferior_req, separacion_max_cm)
        modo_sep_inferior = "Automática"
    else:
        sep_inferior_cm = float(sep_puntera_manual_cm)
        As_inferior_prov = _as_prov_por_separacion(diametro_puntera_mm, sep_inferior_cm)
        modo_sep_inferior = "Manual"

    if sep_talon_manual_cm is None:
        sep_superior_cm, As_superior_prov = seleccionar_separacion(area_barra_cm2(diametro_talon_mm), As_superior_req, separacion_max_cm)
        modo_sep_superior = "Automática"
    else:
        sep_superior_cm = float(sep_talon_manual_cm)
        As_superior_prov = _as_prov_por_separacion(diametro_talon_mm, sep_superior_cm)
        modo_sep_superior = "Manual"

    return {
        "presiones": presiones,
        "Lp_m": Lp, "Lt_m": Lt,
        "q_puntera_prom_ton_m2": q_prom_p,
        "w_suelo_puntera_ton_m2": w_puntera_down,
        "w_neto_puntera_ton_m2": w_neto_puntera,
        "q_talon_prom_ton_m2": q_prom_t,
        "w_suelo_talon_ton_m2": w_talon_down,
        "w_neto_talon_ton_m2": w_neto_talon,

        "Mu_puntera_ton_m_m": abs(Mu_puntera_firmado),
        "Mu_puntera_firmado_ton_m_m": Mu_puntera_firmado,
        "Vu_puntera_ton_m": Vu_puntera,
        "Mu_talon_ton_m_m": abs(Mu_talon_firmado),
        "Mu_talon_firmado_ton_m_m": Mu_talon_firmado,
        "Vu_talon_ton_m": Vu_talon,
        "Vu_talon_firmado_ton_m": gamma_u * w_neto_talon * Lt,

        "M_inferior_critico_ton_m_m": M_inferior,
        "M_superior_critico_ton_m_m": M_superior,
        "As_min_zapata_cara_cm2_m": As_min_cara,
        "As_inferior_flexion_cm2_m": As_inferior_flex,
        "As_superior_flexion_cm2_m": As_superior_flex,
        "As_inferior_req_cm2_m": As_inferior_req,
        "As_inferior_prov_cm2_m": As_inferior_prov,
        "sep_inferior_cm": sep_inferior_cm,
        "diametro_inferior_mm": diametro_puntera_mm,
        "modo_sep_inferior": modo_sep_inferior,
        "As_superior_req_cm2_m": As_superior_req,
        "As_superior_prov_cm2_m": As_superior_prov,
        "sep_superior_cm": sep_superior_cm,
        "diametro_superior_mm": diametro_talon_mm,
        "modo_sep_superior": modo_sep_superior,

        # Alias para compatibilidad con app/dibujos existentes.
        "As_puntera_req_cm2_m": As_inferior_req,
        "As_puntera_prov_cm2_m": As_inferior_prov,
        "sep_puntera_cm": sep_inferior_cm,
        "diametro_puntera_mm": diametro_puntera_mm,
        "modo_sep_puntera": modo_sep_inferior,
        "As_talon_req_cm2_m": As_superior_req,
        "As_talon_prov_cm2_m": As_superior_prov,
        "sep_talon_cm": sep_superior_cm,
        "diametro_talon_mm": diametro_talon_mm,
        "modo_sep_talon": modo_sep_superior,
    }


def calcular_diseno_zapata_definitivo(
    datos: DatosMuro,
    numero_cunas: int = 180,
    recubrimiento_cm: float = 7.5,
    diametro_puntera_mm: float = 16.0,
    diametro_talon_mm: float = 16.0,
    separacion_max_cm: float = 30.0,
    sep_puntera_manual_cm: float | None = None,
    sep_talon_manual_cm: float | None = None
) -> dict:
    res = calcular_diseno_zapata_dinamico(
        datos, numero_cunas, recubrimiento_cm, diametro_puntera_mm, diametro_talon_mm,
        separacion_max_cm, sep_puntera_manual_cm, sep_talon_manual_cm
    )
    presiones = res["presiones"]

    b_cm = 100.0
    h_cm = datos.hz * 100.0
    d_inferior_cm = h_cm - recubrimiento_cm - (diametro_puntera_mm / 10.0) / 2.0
    d_superior_cm = h_cm - recubrimiento_cm - (diametro_talon_mm / 10.0) / 2.0

    Lp = max(datos.puntera, 0.0)
    Lt = max(calcular_talon(datos), 0.0)
    gamma_u = 1.50

    d_puntera_m = d_inferior_cm / 100.0
    Lp_crit = max(Lp - d_puntera_m, 0.0)
    Vu_puntera_crit = gamma_u * abs(res["w_neto_puntera_ton_m2"]) * Lp_crit
    cortante_puntera = _cortante_o_no_aplica(Vu_puntera_crit, b_cm, d_inferior_cm, datos.fc, Lp_crit, "Lp")

    d_talon_m = d_superior_cm / 100.0
    Lt_crit = max(Lt - d_talon_m, 0.0)
    Vu_talon_crit = gamma_u * abs(res["w_neto_talon_ton_m2"]) * Lt_crit
    cortante_talon = _cortante_o_no_aplica(Vu_talon_crit, b_cm, d_superior_cm, datos.fc, Lt_crit, "Lt")

    estado_as_inferior = _estado_acero_manual(res["As_inferior_req_cm2_m"], res["As_inferior_prov_cm2_m"])
    estado_as_superior = _estado_acero_manual(res["As_superior_req_cm2_m"], res["As_superior_prov_cm2_m"])

    def falla(e):
        txt = str(e).lower()
        return txt.startswith("no cumple") or "revisar" in txt

    estados = [presiones["estado_q"], cortante_puntera["estado"], cortante_talon["estado"], estado_as_inferior, estado_as_superior]
    estado_global = "Revisar" if any(falla(e) for e in estados) else "OK"

    res.update({
        "d_puntera_cm": d_inferior_cm, "d_talon_cm": d_superior_cm,
        "d_inferior_cm": d_inferior_cm, "d_superior_cm": d_superior_cm,
        "Lp_crit_m": Lp_crit, "Lt_crit_m": Lt_crit,
        "Vu_puntera_crit_ton_m": Vu_puntera_crit,
        "Vu_talon_crit_ton_m": Vu_talon_crit,
        "cortante_puntera": cortante_puntera,
        "cortante_talon": cortante_talon,
        "estado_as_inferior": estado_as_inferior,
        "estado_as_superior": estado_as_superior,
        "estado_as_puntera": estado_as_inferior,
        "estado_as_talon": estado_as_superior,
        "ld_puntera_cm": 0.0, "ld_talon_cm": 0.0,
        "longitud_disponible_puntera_cm": 0.0, "longitud_disponible_talon_cm": 0.0,
        "estado_ld_puntera": "No se evalúa en dashboard",
        "estado_ld_talon": "No se evalúa en dashboard",
        "estado_global_zapata": estado_global,
    })
    return res


def tabla_diseno_zapata_definitivo(resultado: dict) -> pd.DataFrame:
    filas = [
        ("Estado global zapata", resultado["estado_global_zapata"], "-"),
        ("qmax", resultado["presiones"]["qmax_ton_m2"], "ton/m²"),
        ("qmin", resultado["presiones"]["qmin_ton_m2"], "ton/m²"),
        ("qa", resultado["presiones"]["q_adm_ton_m2"], "ton/m²"),
        ("Estado presión admisible", resultado["presiones"]["estado_q"], "-"),

        ("Relleno sobre puntera", resultado.get("w_suelo_puntera_ton_m2", 0.0), "ton/m²"),
        ("w neto puntera (+ arriba)", resultado.get("w_neto_puntera_ton_m2", 0.0), "ton/m²"),
        ("Mu puntera firmado", resultado.get("Mu_puntera_firmado_ton_m_m", 0.0), "ton·m/m"),
        ("w neto talón (+ abajo)", resultado.get("w_neto_talon_ton_m2", 0.0), "ton/m²"),
        ("Mu talón firmado", resultado.get("Mu_talon_firmado_ton_m_m", 0.0), "ton·m/m"),

        ("Momento crítico cara inferior", resultado["M_inferior_critico_ton_m_m"], "ton·m/m"),
        ("As flexión inferior", resultado["As_inferior_flexion_cm2_m"], "cm²/m"),
        ("As mínimo por cara", resultado["As_min_zapata_cara_cm2_m"], "cm²/m"),
        ("As inferior requerido", resultado["As_inferior_req_cm2_m"], "cm²/m"),
        ("As inferior provisto", resultado["As_inferior_prov_cm2_m"], "cm²/m"),
        ("Armado inferior zapata", _formato_armado(resultado["diametro_inferior_mm"], resultado["sep_inferior_cm"]), "mm @ cm"),
        ("Estado flexión acero inferior", resultado.get("estado_as_inferior", "No aplica"), "-"),

        ("Momento crítico cara superior", resultado["M_superior_critico_ton_m_m"], "ton·m/m"),
        ("As flexión superior", resultado["As_superior_flexion_cm2_m"], "cm²/m"),
        ("As mínimo por cara", resultado["As_min_zapata_cara_cm2_m"], "cm²/m"),
        ("As superior requerido", resultado["As_superior_req_cm2_m"], "cm²/m"),
        ("As superior provisto", resultado["As_superior_prov_cm2_m"], "cm²/m"),
        ("Armado superior zapata", _formato_armado(resultado["diametro_superior_mm"], resultado["sep_superior_cm"]), "mm @ cm"),
        ("Estado flexión acero superior", resultado.get("estado_as_superior", "No aplica"), "-"),

        ("Vu puntera crítico", resultado["Vu_puntera_crit_ton_m"], "ton/m"),
        ("φVc puntera", resultado["cortante_puntera"]["phi_Vc_ton_m"], "ton/m"),
        ("Relación Vu/φVc puntera", resultado["cortante_puntera"]["relacion"], "-"),
        ("Estado cortante puntera", resultado["cortante_puntera"]["estado"], "-"),

        ("Vu talón crítico", resultado["Vu_talon_crit_ton_m"], "ton/m"),
        ("φVc talón", resultado["cortante_talon"]["phi_Vc_ton_m"], "ton/m"),
        ("Relación Vu/φVc talón", resultado["cortante_talon"]["relacion"], "-"),
        ("Estado cortante talón", resultado["cortante_talon"]["estado"], "-"),
    ]
    return pd.DataFrame(filas, columns=["Verificación", "Valor", "Unidad"])


def tabla_momentos_estabilidad(presiones: dict) -> pd.DataFrame:
    filas = [
        ("ESTABILIZANTES", "", "", ""),
        ("Peso zapata - puntera", presiones.get("W_zapata_puntera_ton_m", 0.0), presiones.get("x_zapata_puntera_m", 0.0), presiones.get("M_est_zapata_puntera_ton_m_m", 0.0)),
        ("Peso zapata - bajo fuste", presiones.get("W_zapata_fuste_ton_m", 0.0), presiones.get("x_zapata_fuste_m", 0.0), presiones.get("M_est_zapata_fuste_ton_m_m", 0.0)),
        ("Peso zapata - talón", presiones.get("W_zapata_talon_ton_m", 0.0), presiones.get("x_zapata_talon_m", 0.0), presiones.get("M_est_zapata_talon_ton_m_m", 0.0)),
        ("Peso zapata total", presiones.get("W_zapata_ton_m", 0.0), presiones.get("x_zapata_m", 0.0), presiones.get("M_est_zapata_ton_m_m", 0.0)),
        ("Peso fuste/pantalla", presiones.get("W_fuste_ton_m", 0.0), presiones.get("x_fuste_m", 0.0), presiones.get("M_est_fuste_ton_m_m", 0.0)),
        ("Peso dentellón", presiones.get("W_dentellon_ton_m", 0.0), presiones.get("x_dentellon_m", 0.0), presiones.get("M_est_dentellon_ton_m_m", 0.0)),
        ("Peso suelo sobre puntera", presiones.get("W_suelo_puntera_ton_m", 0.0), presiones.get("x_suelo_puntera_m", 0.0), presiones.get("M_est_suelo_puntera_ton_m_m", 0.0)),
        ("Peso suelo sobre talón", presiones.get("W_suelo_talon_ton_m", 0.0), presiones.get("x_suelo_talon_m", 0.0), presiones.get("M_est_suelo_talon_ton_m_m", 0.0)),
        ("Componente vertical PA", presiones.get("PA_v_ton_m", 0.0), "", presiones.get("M_est_PA_v_ton_m_m", 0.0)),
        ("Total estabilizante", "", "", presiones.get("M_est_ton_m_m", 0.0)),
        ("DESESTABILIZANTES", "", "", ""),
        ("Componente horizontal PAh", presiones.get("PA_h_ton_m", 0.0), presiones.get("brazo_PA_h_m", 0.0), presiones.get("M_volc_PA_h_ton_m_m", 0.0)),
        ("Total desestabilizante", "", "", presiones.get("M_volc_ton_m_m", 0.0)),
        ("RESULTANTE", "", "", ""),
        ("M neto = M_est - M_volc", "", "", presiones.get("M_resultante_ton_m_m", 0.0)),
        ("x resultante", "", "", presiones.get("x_resultante_m", 0.0)),
        ("e", "", "", presiones.get("e_m", 0.0)),
    ]
    return pd.DataFrame(filas, columns=["Concepto", "Fuerza [ton/m]", "Brazo x o y [m]", "Momento [ton·m/m]"])
